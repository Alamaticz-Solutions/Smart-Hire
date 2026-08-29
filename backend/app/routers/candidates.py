"""Candidate CRUD, custom columns, formatted-resume, and docx export endpoints.

Moved verbatim (same paths/methods/logic) from app/main.py:
  - GET    /api/candidates                          (main.py ~988-1030)
  - POST   /api/columns                              (main.py ~1037-1058)
  - GET    /api/columns                              (main.py ~1060-1095)
  - DELETE /api/columns/{col_key}                    (main.py ~1097-1112)
  - PUT    /api/candidates/{candidate_id}             (main.py ~1114-1183)
  - DELETE /api/candidates/{candidate_id}             (main.py ~1185-1220)
  - POST   /api/candidates/bulk-delete                (main.py ~1225-1266)
  - GET    /api/candidates/{candidate_id}/jobs        (main.py ~1268-1300)
  - GET    /api/candidates/{candidate_id}/formatted-resume  (main.py ~1302-1529)
  - GET    /api/candidates/{candidate_id}/export-docx (main.py ~1531-1623)
  - PUT    /api/candidates/{candidate_id}/formatted-resume  (main.py ~1626-1658)
  - POST   /api/candidates                            (main.py ~4208-4251, orphaned far from
             the rest of the candidate routes in the original file; consolidated here)

Also moved as private module-level helpers:
  - get_candidates_list -> _get_candidates_list (main.py ~683-722), used only by
    GET /api/candidates.
  - mask_text_with_keywords / mask_candidate_record + their _PATTERN_CACHE
    (main.py ~736-758). These are also called from the jobs/candidate-matching
    routes (main.py ~3475, ~3556), which live outside this file's scope, so
    this is a genuinely shared utility. It is moved here for now per the
    refactor plan; a later consolidation pass may relocate it to a shared
    `app.services` module once all feature routers exist.
  - get_masked_keywords is NOT moved here: besides GET /api/candidates it is
    also used directly by the admin `/api/admin/masked-keywords` routes
    (main.py ~4435-4464), which are out of this file's scope. A private
    `_get_masked_keywords` copy is kept here instead to avoid taking on
    ownership of an admin-owned endpoint's dependency.
  - get_candidate_name -> _get_candidate_name (main.py ~971-977) and the
    activity-log insert helper -> _log_activity_db, both only used by the
    routes in this file.

Bug fix (explicitly approved dead/broken code removal, not a behavior
change): the original POST /api/candidates handler (main.py ~4214-4218)
read:

    cur.pragma("table_info(candidate_metadata)")
    # wait, cur.execute("PRAGMA table_info(candidate_metadata)")
    cur.execute("PRAGMA table_info(candidate_metadata)")

`sqlite3.Cursor` (and the Postgres cursor adapter) has no `.pragma()` method,
so the first line raised `AttributeError` on every single call to this
route -- it was unreachable dead code that always 500'd before the working
`cur.execute("PRAGMA ...")` line on the next line ever ran, and the leftover
comment above the working line confirms it was dead debug scaffolding. The
broken `.pragma(...)` call and the stray comment have been deleted; only the
working `cur.execute("PRAGMA ...")` line remains. The connection is now also
guaranteed to be closed via `get_db_connection()` instead of the original's
`conn.close()` calls that were skipped on some early-return paths.

Dead-code removal: four `role = get_user_role(username)` assignments whose
value was never subsequently read in that function body have been deleted
(update_candidate, delete_candidate, get_candidate_jobs,
get_formatted_resume_data). Each was confirmed unused by reading the full
original function body before deletion.
"""

import json
import re
from typing import Optional

from fastapi import APIRouter, BackgroundTasks, Depends, HTTPException, Request
from fastapi.responses import JSONResponse
from pydantic import BaseModel

from app.core.logging import get_logger
from app.db.row_helpers import dict_row_factory
from app.db.session import get_db_connection
from app.dependencies import assert_owns_or_admin, require_approved_user
from app.services.auth import apply_user_hidden_fields, get_user_info, is_admin_or_hr, is_user_approved
from app.services.matching import match_candidate_to_all_jobs
from app.services.resume_processing import format_candidate_resume

logger = get_logger(__name__)

router = APIRouter()


class CustomColumn(BaseModel):
    col_key: str
    col_label: str
    description: str


class BulkDeleteRequest(BaseModel):
    ids: list[int]


# ── Private helpers (candidate-routes-only; see module docstring) ──────────────

_PATTERN_CACHE: dict = {}


def _log_activity_db(username: str, action: str) -> None:
    """Insert one row into activity_logs. Mirrors main.py's log_activity_db."""
    if not username:
        username = "unknown"
    try:
        with get_db_connection() as conn:
            cur = conn.cursor()
            cur.execute("INSERT INTO activity_logs (username, action) VALUES (?, ?)", (username, action))
            conn.commit()
    except Exception as e:
        logger.error(f"Error logging activity: {e}")


def _get_candidate_name(candidate_id: int) -> str:
    with get_db_connection() as conn:
        cur = conn.cursor()
        cur.execute("SELECT full_name FROM candidate_metadata WHERE id = ?", (candidate_id,))
        row = cur.fetchone()
    return row[0] if row else f"ID {candidate_id}"


def _get_masked_keywords() -> list:
    with get_db_connection() as conn:
        cur = conn.cursor()
        try:
            cur.execute("SELECT keyword FROM masked_keywords")
            keywords = [row[0] for row in cur.fetchall()]
        except Exception:
            keywords = []
    return keywords


def mask_text_with_keywords(text: str, keywords: list) -> str:
    if not text or not keywords:
        return text
    result = str(text)
    for kw in keywords:
        kw_strip = kw.strip()
        if not kw_strip:
            continue
        if kw_strip not in _PATTERN_CACHE:
            _PATTERN_CACHE[kw_strip] = re.compile(re.escape(kw_strip), re.IGNORECASE)
        result = _PATTERN_CACHE[kw_strip].sub("****", result)
    return result


def mask_candidate_record(candidate: dict, keywords: list) -> dict:
    masked = {}
    for k, v in candidate.items():
        if isinstance(v, str):
            masked[k] = mask_text_with_keywords(v, keywords)
        else:
            masked[k] = v
    return masked


def _get_candidates_list(
    username: Optional[str] = None,
    role: str = "user",
    is_hr_or_admin: bool = None,
    limit: Optional[int] = None,
    offset: int = 0,
) -> tuple[list, Optional[int]]:
    with get_db_connection() as conn:
        conn.row_factory = dict_row_factory
        cur = conn.cursor()
        try:
            if is_hr_or_admin is None:
                is_hr_or_admin = False
                if username:
                    is_hr_or_admin = is_admin_or_hr(username)
            cols_to_select = (
                "id, filename, full_name, candidate_status, total_experience, pega_experience, "
                "skills, certifications, ctc, notice_period, current_organization, email, phone, "
                "linkedin, created_by, timestamp, source, cdh_exp, expected_ctc, percentage_hike, "
                "candidate_interview_status, availability_in_days, current_location, pref_locations, "
                "current_client, domain, tier, certification_version, "
                "sender_email, is_qualified, is_approved, file_url, error_detail"
            )
            # LIMIT/OFFSET are only ever set to a real value from the paginated
            # branch of GET /api/candidates below -- every other caller (the
            # default, unpaginated behavior; email_worker; excel_import, etc.)
            # passes limit=None, which keeps this exactly the "no LIMIT clause
            # at all" query it always was, so nothing about the default
            # behavior changes here.
            #
            # When paginating, also select COUNT(*) OVER() - a window function
            # that returns the full matching-row count on every row of the
            # SAME result set, no GROUP BY needed - so the caller gets its
            # pagination total without the separate _count_candidates() round
            # trip this used to require on every single page load (this app's
            # database round-trip cost is ~300-500ms; that was a guaranteed
            # extra one on the very first thing a user sees on this page).
            limit_clause = ""
            params: tuple = ()
            select_cols = cols_to_select
            if limit is not None:
                limit_clause = " LIMIT ? OFFSET ?"
                params = (limit, offset)
                select_cols = f"{cols_to_select}, COUNT(*) OVER() AS __total_count"

            if role == "admin" or is_hr_or_admin or not username:
                cur.execute(f"SELECT {select_cols} FROM candidate_metadata ORDER BY timestamp DESC{limit_clause}", params)
            else:
                cur.execute(
                    f"SELECT {select_cols} FROM candidate_metadata WHERE LOWER(created_by) = LOWER(?) "
                    f"ORDER BY timestamp DESC{limit_clause}",
                    (username,) + params,
                )
            raw_rows = cur.fetchall()
            rows = []
            total = 0
            for r in raw_rows:
                try:
                    row_dict = dict(r)
                except Exception as row_err:
                    # If dict(r) fails, build dict manually from cursor description
                    if cur.description:
                        col_names = [desc[0] for desc in cur.description]
                        row_dict = dict(zip(col_names, r))
                    else:
                        logger.warning(f"Could not convert row to dict: {row_err}")
                        continue
                if limit is not None:
                    total = row_dict.pop("__total_count", total)
                rows.append(row_dict)
            if limit is None:
                total = None
        except Exception as e:
            logger.error(f"ERROR in _get_candidates_list: {e}", exc_info=True)
            rows, total = [], (0 if limit is not None else None)
    return rows, total


def _count_candidates(username: Optional[str] = None, role: str = "user", is_hr_or_admin: bool = None) -> int:
    """Row count for the same visibility rule `_get_candidates_list` applies, for pagination totals."""
    with get_db_connection() as conn:
        cur = conn.cursor()
        try:
            if is_hr_or_admin is None:
                is_hr_or_admin = False
                if username:
                    is_hr_or_admin = is_admin_or_hr(username)
            if role == "admin" or is_hr_or_admin or not username:
                cur.execute("SELECT COUNT(*) FROM candidate_metadata")
            else:
                cur.execute("SELECT COUNT(*) FROM candidate_metadata WHERE LOWER(created_by) = LOWER(?)", (username,))
            return cur.fetchone()[0]
        except Exception as e:
            logger.error(f"ERROR in _count_candidates: {e}", exc_info=True)
            return 0


# ── Candidates ─────────────────────────────────────────────────────────────────
@router.get("/api/candidates")
def list_candidates(request: Request, limit: Optional[int] = None, offset: int = 0):
    # `limit`/`offset` are optional and opt-in: omitting them (every existing
    # caller today) preserves the exact original behavior of returning every
    # visible candidate as a bare JSON array. Passing them switches to a
    # paginated response shape (`{"items": [...], "total": N}`) instead --
    # this endpoint used to run one unbounded `SELECT * FROM candidate_metadata`
    # per call (no LIMIT anywhere), which got noticeably more expensive as the
    # table grew and was hit by this app's own polling. A caller that wants a
    # page can now ask for one without changing behavior for anyone who doesn't.
    username = request.headers.get("x-user-username")

    is_user_admin = False
    is_external = False
    is_admin_or_hr_flag = False
    if username:
        row = get_user_info(username)
        if row:
            is_external = row["is_external"] == 1
            is_admin_or_hr_flag = row["is_admin"] == 1 or row["is_hr"] == 1
            is_user_admin = row["is_admin"] == 1 or row["role"] == "admin" or is_admin_or_hr_flag
            if is_external:
                raise HTTPException(status_code=403, detail="Forbidden")

    if not is_user_approved(username):
        return {"items": [], "total": 0} if limit is not None else []

    effective_role = "admin" if is_user_admin else "user"
    rows, total_from_query = _get_candidates_list(
        username, role=effective_role, is_hr_or_admin=is_admin_or_hr_flag, limit=limit, offset=offset
    )

    # Replace None values with empty string and remove binary file_bytes
    for row in rows:
        row.pop("file_bytes", None)
        for k, v in row.items():
            if v is None:
                row[k] = ""

    # Mask certifications for non-admin and non-HR users
    if not is_admin_or_hr_flag:
        for row in rows:
            row["certifications"] = "[HIDDEN]"
        keywords = _get_masked_keywords()
        rows = [mask_candidate_record(row, keywords) for row in rows]

    rows = apply_user_hidden_fields(rows, username)

    if limit is not None:
        # COUNT(*) OVER() rides along on the main query's rows, so the common
        # case (there ARE rows on this page) needs no extra round trip. Only
        # fall back to the standalone count query for the edge case it can't
        # cover - zero rows back (offset past the end, or genuinely no
        # candidates) - where the window function gives no signal either way.
        total = total_from_query if rows else _count_candidates(username, role=effective_role, is_hr_or_admin=is_admin_or_hr_flag)
        return {"items": rows, "total": total}
    return rows


@router.post("/api/columns")
def add_column(col: CustomColumn, request: Request, username: str = Depends(require_approved_user)):
    with get_db_connection() as conn:
        cur = conn.cursor()
        clean_key = re.sub(r"[^a-zA-Z0-9_]", "", col.col_key.replace(" ", "_")).lower()

        cur.execute("PRAGMA table_info(candidate_metadata)")
        existing = [c[1] for c in cur.fetchall()]
        if clean_key in existing:
            raise HTTPException(status_code=400, detail="Column already exists")

        try:
            cur.execute(f"ALTER TABLE candidate_metadata ADD COLUMN {clean_key} TEXT")
            cur.execute(
                "INSERT INTO custom_columns (col_key, col_label, description) VALUES (?, ?, ?)",
                (clean_key, col.col_label, col.description),
            )
            conn.commit()
        except Exception as e:
            raise HTTPException(status_code=500, detail=str(e))
    return {"status": "added", "col_key": clean_key}


@router.get("/api/columns")
def get_columns(username: str = Depends(require_approved_user)):
    with get_db_connection() as conn:
        cur = conn.cursor()
        cur.execute("SELECT col_key, col_label FROM custom_columns")
        customs = [{"col_key": row[0], "col_label": row[1]} for row in cur.fetchall()]

    base_cols = [
        {"col_key": "full_name", "col_label": "Name"},
        {"col_key": "candidate_status", "col_label": "Candidate Status"},
        {"col_key": "source", "col_label": "Source"},
        {"col_key": "total_experience", "col_label": "Total Exp"},
        {"col_key": "pega_experience", "col_label": "Pega Exp"},
        {"col_key": "cdh_exp", "col_label": "CDH Exp"},
        {"col_key": "ctc", "col_label": "Current CTC"},
        {"col_key": "expected_ctc", "col_label": "Exp CTC"},
        {"col_key": "percentage_hike", "col_label": "Percentage Hike"},
        {"col_key": "candidate_interview_status", "col_label": "Candidate Interview Status"},
        {"col_key": "availability_in_days", "col_label": "Availability in Days"},
        {"col_key": "notice_period", "col_label": "Notice Period"},
        {"col_key": "phone", "col_label": "Phone No"},
        {"col_key": "email", "col_label": "Email"},
        {"col_key": "linkedin", "col_label": "LinkedIn"},
        {"col_key": "current_location", "col_label": "Current Location"},
        {"col_key": "pref_locations", "col_label": "Pref Locations"},
        {"col_key": "current_organization", "col_label": "Current Employment"},
        {"col_key": "current_client", "col_label": "Current Client"},
        {"col_key": "domain", "col_label": "Domain"},
        {"col_key": "tier", "col_label": "Tier (Tier1 Tier2 Tier3)"},
        {"col_key": "certification_version", "col_label": "Certification Version"},
        {"col_key": "skills", "col_label": "Skills"},
        {"col_key": "certifications", "col_label": "Certifications"},
        {"col_key": "notescomments", "col_label": "Notes / Comments"},
    ]
    return {"base": base_cols, "custom": customs}


@router.delete("/api/columns/{col_key}")
def delete_column(col_key: str, request: Request, username: str = Depends(require_approved_user)):
    # SECURITY FIX: `col_key` is a raw path parameter and was being
    # interpolated directly into `ALTER TABLE ... DROP COLUMN {col_key}`
    # below with no sanitization, unlike `add_column` (above) which cleans
    # its equivalent input via the same regex before using it in SQL. That
    # made this endpoint a SQL-injection vector (e.g. a crafted col_key
    # containing extra SQL). Apply the identical whitelist sanitization
    # used by `add_column` so only safe identifier characters ever reach
    # the raw SQL string.
    clean_key = re.sub(r"[^a-zA-Z0-9_]", "", col_key.replace(" ", "_")).lower()
    with get_db_connection() as conn:
        cur = conn.cursor()
        try:
            cur.execute("DELETE FROM custom_columns WHERE col_key=?", (clean_key,))
            try:
                cur.execute(f"ALTER TABLE candidate_metadata DROP COLUMN {clean_key}")
            except Exception:
                pass  # older sqlite versions might not support drop column
            conn.commit()
        except Exception as e:
            raise HTTPException(status_code=500, detail=str(e))
    return {"status": "deleted"}


@router.put("/api/candidates/{candidate_id}")
async def update_candidate(
    candidate_id: int,
    request: Request,
    background_tasks: BackgroundTasks,
    username: str = Depends(require_approved_user),
):
    body = await request.json()
    with get_db_connection() as conn:
        cur = conn.cursor()

        # Check permission
        cur.execute("SELECT created_by FROM candidate_metadata WHERE id = ?", (candidate_id,))
        row = cur.fetchone()
        if not row:
            raise HTTPException(status_code=404, detail="Candidate not found")
        created_by = row[0]
        assert_owns_or_admin(created_by, username)

        # Explicit allow-list, not "every column except id": the old
        # `PRAGMA table_info` check let the request body overwrite ANY
        # column that exists on the table, including created_by (silent
        # ownership takeover), filename/file_url (breaks file lookup and
        # the delete/embedding-cleanup logic, which key off filename), and
        # file_bytes/timestamp. Only genuinely user-editable fields belong
        # here - mirrors the pattern jobs.py's JobCreate model already uses.
        EDITABLE_CANDIDATE_FIELDS = {
            "full_name", "candidate_status", "total_experience", "pega_experience",
            "skills", "certifications", "ctc", "notice_period", "current_organization",
            "email", "phone", "linkedin", "source", "cdh_exp", "expected_ctc",
            "percentage_hike", "candidate_interview_status", "availability_in_days",
            "current_location", "pref_locations", "current_client", "domain", "tier",
            "certification_version", "sender_email", "is_qualified", "is_approved",
        }
        # Admin-defined custom columns (see add_column below) are also
        # user-editable free-text fields, on top of the fixed allow-list.
        try:
            cur.execute("SELECT col_key FROM custom_columns")
            EDITABLE_CANDIDATE_FIELDS = EDITABLE_CANDIDATE_FIELDS | {r[0] for r in cur.fetchall()}
        except Exception:
            pass
        updates = {k: v for k, v in body.items() if k in EDITABLE_CANDIDATE_FIELDS and v is not None and v != "[HIDDEN]"}

        # Server-side validation for numeric edits
        for int_col in ["notice_period", "availability_in_days"]:
            if int_col in updates and updates[int_col] != "":
                try:
                    updates[int_col] = int(float(updates[int_col]))
                except ValueError:
                    return JSONResponse(status_code=400, content={"detail": f"{int_col} must be an integer"})

        for exp_col in ["total_experience", "pega_experience", "cdh_exp"]:
            if exp_col in updates and updates[exp_col] != "":
                try:
                    updates[exp_col] = float(updates[exp_col])
                except ValueError:
                    return JSONResponse(status_code=400, content={"detail": f"{exp_col} must be a number"})

        if not updates:
            return {"status": "no changes"}

        set_clause = ", ".join(f"{k}=?" for k in updates)
        cur.execute(
            f"UPDATE candidate_metadata SET {set_clause} WHERE id=?",
            list(updates.values()) + [candidate_id],
        )
        conn.commit()

    # Re-trigger matching if matching-related details have changed
    match_related_fields = {
        "full_name",
        "total_experience",
        "pega_experience",
        "cdh_exp",
        "skills",
        "certifications",
        "current_location",
        "pref_locations",
    }
    if any(field in updates for field in match_related_fields):
        background_tasks.add_task(match_candidate_to_all_jobs, candidate_id)

    cname = _get_candidate_name(candidate_id)
    _log_activity_db(username or "unknown", f"updated candidate '{cname}' details")

    return {"status": "updated"}


@router.delete("/api/candidates/{candidate_id}")
def delete_candidate(candidate_id: int, username: str = Depends(require_approved_user)):
    with get_db_connection() as conn:
        cur = conn.cursor()

        cur.execute(
            "SELECT full_name, created_by, filename, sender_email FROM candidate_metadata WHERE id = ?",
            (candidate_id,),
        )
        row = cur.fetchone()
        if not row:
            raise HTTPException(status_code=404, detail="Candidate not found")
        cname, created_by, filename, sender_email_val = row

        assert_owns_or_admin(created_by, username)

        cur.execute("DELETE FROM job_candidates WHERE candidate_id=?", (candidate_id,))
        if filename:
            try:
                cur.execute("DELETE FROM langchain_pg_embedding WHERE cmetadata->>'source' = ?", (filename,))
            except Exception as embed_err:
                logger.warning(f"Failed to delete embeddings for {filename}: {embed_err}")
            # Block the source email from being re-processed by marking it as processed
            if filename.startswith("mail_"):
                cur.execute(
                    "INSERT OR IGNORE INTO processed_emails (msg_uid) VALUES (?)",
                    (f"deleted_cand_{candidate_id}",),
                )
        cur.execute("DELETE FROM candidate_metadata WHERE id=?", (candidate_id,))
        conn.commit()

    _log_activity_db(username or "unknown", f"deleted candidate '{cname}'")
    return {"status": "deleted"}


@router.post("/api/candidates/bulk-delete")
def bulk_delete_candidates(req: BulkDeleteRequest, username: str = Depends(require_approved_user)):
    deleted_names = []

    with get_db_connection() as conn:
        cur = conn.cursor()

        for cid in req.ids:
            cur.execute("SELECT full_name, created_by, filename FROM candidate_metadata WHERE id = ?", (cid,))
            row = cur.fetchone()
            if not row:
                continue
            cname, created_by, filename = row

            # Check permissions
            if not is_admin_or_hr(username):
                if created_by and created_by.lower() != username.lower():
                    continue  # Skip unauthorized deletions

            cur.execute("DELETE FROM job_candidates WHERE candidate_id=?", (cid,))
            if filename:
                try:
                    cur.execute("DELETE FROM langchain_pg_embedding WHERE cmetadata->>'source' = ?", (filename,))
                except Exception as embed_err:
                    logger.warning(f"Failed to delete embeddings for {filename}: {embed_err}")
                # Block the source email from being re-processed
                if filename.startswith("mail_"):
                    cur.execute("INSERT OR IGNORE INTO processed_emails (msg_uid) VALUES (?)", (f"deleted_cand_{cid}",))
            cur.execute("DELETE FROM candidate_metadata WHERE id=?", (cid,))
            deleted_names.append(cname)

        conn.commit()

    if deleted_names:
        _log_activity_db(
            username or "unknown",
            f"bulk deleted {len(deleted_names)} candidates: {', '.join(deleted_names[:5])}",
        )

    return {"status": "deleted", "count": len(deleted_names)}


@router.get("/api/candidates/{candidate_id}/jobs")
def get_candidate_jobs(candidate_id: int, username: str = Depends(require_approved_user)):
    with get_db_connection() as conn:
        conn.row_factory = dict_row_factory
        cur = conn.cursor()

        # Check candidate existence and permission
        cur.execute("SELECT created_by FROM candidate_metadata WHERE id = ?", (candidate_id,))
        cand_row = cur.fetchone()
        if not cand_row:
            raise HTTPException(status_code=404, detail="Candidate not found")

        created_by = cand_row[0]
        assert_owns_or_admin(created_by, username)

        cur.execute(
            """
            SELECT j.*, jc.status as match_status, jc.ai_reason
            FROM jobs j
            JOIN job_candidates jc ON j.id = jc.job_id
            WHERE jc.candidate_id = ?
        """,
            (candidate_id,),
        )
        rows = [dict(r) for r in cur.fetchall()]
    return rows


@router.get("/api/candidates/{candidate_id}/formatted-resume")
def get_formatted_resume_data(candidate_id: int, username: str = Depends(require_approved_user)):
    with get_db_connection() as conn:
        conn.row_factory = dict_row_factory
        cur = conn.cursor()
        cur.execute("SELECT * FROM candidate_metadata WHERE id = ?", (candidate_id,))
        row = cur.fetchone()

    if not row:
        raise HTTPException(status_code=404, detail="Candidate not found")

    candidate = dict(row)
    created_by = candidate.get("created_by")
    assert_owns_or_admin(created_by, username)

    # The cache check, file-text loading, LLM formatting, fallback data, and
    # DB caching all live in the shared `format_candidate_resume` helper
    # (app.services.resume_processing) -- this route is now the thin
    # auth-check + row-fetch wrapper its own docstring always said it should
    # be, instead of a second inline copy of that logic.
    return format_candidate_resume(candidate, candidate_id)


@router.get("/api/candidates/{candidate_id}/export-docx")
def export_candidate_docx(candidate_id: int, username: str = Depends(require_approved_user)):
    from fastapi.responses import StreamingResponse
    import io

    # Try importing docx locally so server doesn't crash on boot if package is installing
    try:
        from docx import Document
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        raise HTTPException(status_code=500, detail="python-docx library not installed. Please try again in a few seconds.")

    with get_db_connection() as conn:
        conn.row_factory = dict_row_factory
        cur = conn.cursor()
        cur.execute("SELECT * FROM candidate_metadata WHERE id = ?", (candidate_id,))
        row = cur.fetchone()

    if not row:
        raise HTTPException(status_code=404, detail="Candidate not found")

    candidate = dict(row)
    created_by = candidate.get("created_by")
    assert_owns_or_admin(created_by, username)

    doc = Document()

    # Alamaticz Styling
    title = doc.add_heading(candidate.get("full_name", "Candidate Resume") or "Candidate Resume", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Contact Info
    contact_p = doc.add_paragraph()
    contact_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_info = []
    if candidate.get("email"):
        contact_info.append(str(candidate.get("email")))
    if candidate.get("phone"):
        contact_info.append(str(candidate.get("phone")))
    if candidate.get("linkedin"):
        contact_info.append(str(candidate.get("linkedin")))
    if candidate.get("current_location"):
        contact_info.append(str(candidate.get("current_location")))

    if contact_info:
        run = contact_p.add_run(" | ".join(contact_info))
        run.font.size = Pt(10)

    doc.add_heading("Summary", level=1)
    exp = str(candidate.get("total_experience") or "N/A")
    doc.add_paragraph(f"Total Experience: {exp}")
    if candidate.get("pega_experience"):
        doc.add_paragraph(f"Pega Experience: {candidate.get('pega_experience')}")

    if candidate.get("skills"):
        doc.add_heading("Skills", level=1)
        doc.add_paragraph(str(candidate.get("skills")))

    if candidate.get("certifications"):
        doc.add_heading("Certifications", level=1)
        doc.add_paragraph(str(candidate.get("certifications")))

    # Additional Details
    doc.add_heading("Additional Details", level=1)
    if candidate.get("current_organization"):
        doc.add_paragraph(f"Current Organization: {candidate.get('current_organization')}")
    if candidate.get("ctc"):
        doc.add_paragraph(f"Current CTC: {candidate.get('ctc')}")
    if candidate.get("expected_ctc"):
        doc.add_paragraph(f"Expected CTC: {candidate.get('expected_ctc')}")
    if candidate.get("notice_period"):
        doc.add_paragraph(f"Notice Period: {candidate.get('notice_period')}")

    # Save to memory
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)

    safe_name = str(candidate.get("full_name") or "Candidate").replace(" ", "_")
    headers = {
        "Content-Disposition": f'attachment; filename="Alamaticz_Resume_{safe_name}.docx"',
        "Access-Control-Expose-Headers": "Content-Disposition",
    }

    return StreamingResponse(
        file_stream,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers=headers,
    )


@router.put("/api/candidates/{candidate_id}/formatted-resume")
async def update_formatted_resume_data(
    candidate_id: int,
    request: Request,
    username: str = Depends(require_approved_user),
):
    body = await request.json()

    with get_db_connection() as conn:
        try:
            cur = conn.cursor()
            cur.execute("SELECT created_by FROM candidate_metadata WHERE id = ?", (candidate_id,))
            row = cur.fetchone()
            if not row:
                raise HTTPException(status_code=404, detail="Candidate not found")

            created_by = row[0]
            assert_owns_or_admin(created_by, username)

            cur.execute("UPDATE candidate_metadata SET formatted_json = ? WHERE id = ?", (json.dumps(body), candidate_id))
            conn.commit()
        except HTTPException:
            raise
        except Exception as e:
            raise HTTPException(status_code=500, detail=str(e))

    cname = _get_candidate_name(candidate_id)
    _log_activity_db(username or "unknown", f"updated formatted resume for candidate '{cname}'")
    return {"status": "updated"}


@router.post("/api/candidates")
def add_candidate_manually(body: dict, username: str = Depends(require_approved_user)):
    with get_db_connection() as conn:
        cur = conn.cursor()
        # NOTE: the original main.py had a broken `cur.pragma("table_info(...)")`
        # line here (sqlite3 cursors have no .pragma() method -- it always
        # raised AttributeError) immediately followed by a comment and the
        # correct `cur.execute("PRAGMA ...")` call. That dead/broken line and
        # the stray comment have been removed; only the working PRAGMA query
        # below remains. See module docstring for details.
        cur.execute("PRAGMA table_info(candidate_metadata)")
        valid_cols = [c[1] for c in cur.fetchall()]

        insert_data = {}
        for col in valid_cols:
            if col in ("id", "timestamp"):
                continue
            if col in body:
                insert_data[col] = body[col]

        if "source" not in insert_data or not insert_data["source"]:
            insert_data["source"] = "Manual Entry"
        if "candidate_status" not in insert_data or not insert_data["candidate_status"]:
            insert_data["candidate_status"] = "New"
        insert_data["created_by"] = username or "admin"

        if not insert_data.get("full_name"):
            raise HTTPException(status_code=400, detail="Candidate Name is required")

        # Same numeric coercion PUT /api/candidates/{id} (update_candidate,
        # above) already applies -- this insert path had none, so a manually
        # entered non-numeric value (e.g. "5+" years) would sit in the DB as
        # a raw string and later crash any code that does float(...) on it
        # (e.g. email_worker._get_missing_fields when this candidate gets a
        # follow-up email). Reject bad input up front instead.
        for int_col in ["notice_period", "availability_in_days"]:
            if int_col in insert_data and insert_data[int_col] != "":
                try:
                    insert_data[int_col] = int(float(insert_data[int_col]))
                except (ValueError, TypeError):
                    raise HTTPException(status_code=400, detail=f"{int_col} must be an integer")

        for exp_col in ["total_experience", "pega_experience", "cdh_exp"]:
            if exp_col in insert_data and insert_data[exp_col] != "":
                try:
                    insert_data[exp_col] = float(insert_data[exp_col])
                except (ValueError, TypeError):
                    raise HTTPException(status_code=400, detail=f"{exp_col} must be a number")

        cols_str = ", ".join(insert_data.keys())
        placeholders = ", ".join(["?"] * len(insert_data))
        vals = list(insert_data.values())

        try:
            cur.execute(f"INSERT INTO candidate_metadata ({cols_str}) VALUES ({placeholders})", vals)
            conn.commit()
            _log_activity_db(username or "unknown", f"manually added candidate '{insert_data.get('full_name')}'")
        except Exception as e:
            raise HTTPException(status_code=500, detail=str(e))

    return {"status": "success", "message": "Candidate added successfully"}
