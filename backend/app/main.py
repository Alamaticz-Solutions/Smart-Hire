import os
import sys
sys.path.append(os.path.dirname(os.path.abspath(__file__)))
import json
import hashlib
import sqlite3
import shutil
import re
import gc
import threading
from typing import Optional
from datetime import datetime
import openpyxl

# Helper row factory that works for both SQLite and PostgreSQL adapters
def dict_row_factory(cursor, row):
    """Return a dict for a DB row.
    For SQLite, the default Row object can be cast to dict.
    For PostgreSQL (PGCursor), the row is already a mapping; we attempt dict conversion.
    """
    try:
        if isinstance(row, dict):
            return row
        if hasattr(row, 'keys'):
            return {k: row[k] for k in row.keys()}
        if hasattr(cursor, 'description') and cursor.description:
            return {col[0]: val for col, val in zip(cursor.description, row)}
        return dict(row)
    except Exception:
        # Fallback: return the row as is (may be a tuple for SQLite without row_factory)
        return row


def normalize_phone(phone) -> str:
    if not phone:
        return ""
    s = str(phone).strip()
    if s.endswith('.0'):
        s = s[:-2]
    digits = "".join(c for c in s if c.isdigit())
    if len(digits) >= 10:
        return digits[-10:]
    return digits

def normalize_email(email) -> str:
    if not email:
        return ""
    return str(email).strip().lower()

def get_base_email(email: str) -> str:
    if not email:
        return ""
    email_str = str(email).strip().lower()
    if "@" in email_str:
        parts = email_str.split("@")
        local_part = parts[0].split("+")[0]
        return f"{local_part}@{parts[1]}"
    return email_str

def is_similar_name(n1: str, n2: str) -> bool:
    if not n1 or not n2:
        return False
    # Clean non-alphanumeric chars and split into lowercase words
    w1 = set(re.findall(r'[a-z0-9]+', n1.lower()))
    w2 = set(re.findall(r'[a-z0-9]+', n2.lower()))
    if not w1 or not w2:
        return False
    # If they are exactly equal
    if n1.strip().lower() == n2.strip().lower():
        return True
    
    # Calculate intersection
    intersection = w1.intersection(w2)
    if not intersection:
        return False
        
    # We want at least two common words to match (e.g. "durga" and "dheeraj")
    # unless one of them is only a single word, in which case we check if it is part of the other
    shorter_len = min(len(w1), len(w2))
    if shorter_len == 1:
        word = list(w1 if len(w1) == 1 else w2)[0]
        if len(word) >= 4:
            return word in w1 or word in w2
        return False
    
    return len(intersection) >= 2

def phones_match(p1: str, p2: str) -> bool:
    # Get only the digits of both phone numbers
    d1 = "".join(c for c in str(p1) if c.isdigit())
    d2 = "".join(c for c in str(p2) if c.isdigit())
    if not d1 or not d2:
        return False
    # If they are exactly equal
    if d1 == d2:
        return True
    # If the last 10 digits are equal
    if len(d1) >= 10 and len(d2) >= 10 and d1[-10:] == d2[-10:]:
        return True
    # Check if the shorter one is contained in the longer one, as long as it's at least 9 digits
    shorter, longer = (d1, d2) if len(d1) < len(d2) else (d2, d1)
    if len(shorter) >= 9 and shorter in longer:
        return True
    return False

from fastapi import FastAPI, UploadFile, File, HTTPException, Form, Request, BackgroundTasks
from fastapi.middleware.cors import CORSMiddleware
from fastapi.staticfiles import StaticFiles
from fastapi.responses import JSONResponse, FileResponse
from starlette.exceptions import HTTPException as StarletteHTTPException
from pydantic import BaseModel
from dotenv import load_dotenv
import fitz
from langchain_core.documents import Document

class SafePyMuPDFLoader:
    def __init__(self, file_path: str):
        self.file_path = file_path

    def load(self):
        docs = []
        try:
            doc = fitz.open(self.file_path) 
            for page_num, page in enumerate(doc):
                text = page.get_text()
                metadata = {
                    "source": self.file_path,
                    "page": page_num,
                }
                docs.append(Document(page_content=text, metadata=metadata))
            doc.close()
        except Exception as e:
            print(f"Error loading PDF {self.file_path} with fitz: {e}")
            raise e
        return docs

# LangChain / AI
from langchain_community.document_loaders import Docx2txtLoader
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_community.vectorstores import PGVector
from langchain_groq import ChatGroq
from langchain_core.messages import HumanMessage
from langchain_core.runnables import RunnablePassthrough
from langchain_core.output_parsers import StrOutputParser
from langchain_core.prompts import PromptTemplate

# Load root .env first
load_dotenv()
# Explicitly load backend/.env to override/supplement configuration
_backend_env_path = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), ".env")
if os.path.exists(_backend_env_path):
    load_dotenv(_backend_env_path, override=True)

# Check and patch PostgreSQL (Mandatory)
from app.db.postgres_adapter import patch_if_configured
try:
    PG_ACTIVE = patch_if_configured()
except Exception as e:
    print(f"CRITICAL CONFIGURATION ERROR: {e}")
    import sys
    sys.exit(1)

BASE_DIR     = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
PROJECT_ROOT = os.path.dirname(BASE_DIR)

def row_to_dict(row):
    if not row: return {}
    if isinstance(row, dict): return row
    if hasattr(row, 'items') and callable(getattr(row, 'items')): 
        return dict(row.items())
    try:
        return dict(row)
    except Exception:
        return {}
# Support persistent volume directory /data (e.g. on Render)
DATA_DIR = "/data" if os.path.exists("/data") and os.access("/data", os.W_OK) else BASE_DIR


UPLOAD_DIR   = os.path.join(DATA_DIR, "static")
STATS_DB     = os.getenv("STATS_DB_PATH", os.path.join(DATA_DIR, "stats.db"))
GROQ_API_KEY = os.getenv("GROQ_API_KEY", "")

os.makedirs(UPLOAD_DIR, exist_ok=True)

os.makedirs(os.path.dirname(STATS_DB), exist_ok=True)

# ── FastAPI App ────────────────────────────────────────────────────────────────
app = FastAPI(title="Hire AI API")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

from fastapi.responses import Response
import mimetypes

@app.get("/static/{filename}")
def serve_file_from_db(filename: str):
    conn = sqlite3.connect(STATS_DB, timeout=30.0)
    cur = conn.cursor()
    
    # Query file_url and file_bytes
    cur.execute("PRAGMA table_info(candidate_metadata)")
    existing_cols = {c[1] for c in cur.fetchall()}
    
    if 'file_url' in existing_cols:
        cur.execute("SELECT file_url, file_bytes FROM candidate_metadata WHERE filename = ? LIMIT 1", (filename,))
        row = cur.fetchone()
        file_url = row[0] if row else None
        file_bytes = row[1] if row else None
    else:
        cur.execute("SELECT file_bytes FROM candidate_metadata WHERE filename = ? AND file_bytes IS NOT NULL LIMIT 1", (filename,))
        row = cur.fetchone()
        file_url = None
        file_bytes = row[0] if row else None
        
    conn.close()
    
    if file_url:
        from fastapi.responses import RedirectResponse
        return RedirectResponse(url=file_url)
        
    if not file_bytes:
        raise HTTPException(status_code=404, detail="File not found")
        
    # If returned as memoryview, convert to bytes
    if isinstance(file_bytes, memoryview):
        file_bytes = file_bytes.tobytes()
        
    mime_type, _ = mimetypes.guess_type(filename)
    if not mime_type:
        mime_type = "application/octet-stream"
        
    return Response(
        content=file_bytes,
        media_type=mime_type,
        headers={"Content-Disposition": f"inline; filename={filename}"}
    )


# Serve uploaded resumes
app.mount("/static", StaticFiles(directory=UPLOAD_DIR), name="static")

# ── Models (loaded once) ───────────────────────────────────────────────────────
_embeddings = None
_llm        = None
_active_groq_key = None
_models_loading = False
_model_lock = threading.Lock()
_processing_lock = threading.Lock()

def get_models():
    global _embeddings, _llm, _models_loading, _active_groq_key
    with _model_lock:
        _models_loading = True
        if _embeddings is None:
            try:
                _embeddings = HuggingFaceEmbeddings(model_name="sentence-transformers/all-MiniLM-L6-v2")
            except Exception as e:
                print(f"Warning: Failed to load HuggingFaceEmbeddings: {e}")
                _embeddings = None
        
        # Load env variables dynamically in case the key has changed in the file
        try:
            from dotenv import load_dotenv
            load_dotenv(override=True)
        except Exception:
            pass
        current_key = os.getenv("GROQ_API_KEY", "")
        
        if _llm is None or _active_groq_key != current_key:
            _llm = ChatGroq(temperature=0.1, model_name="llama-3.1-8b-instant", groq_api_key=current_key)
            _active_groq_key = current_key
            
        _models_loading = False
    return _embeddings, _llm

@app.on_event("startup")
def load_models_in_background():
    threading.Thread(target=get_models, daemon=True).start()
    threading.Thread(target=poll_emails_and_process, daemon=True).start()

# ── DB Helpers ─────────────────────────────────────────────────────────────────
def init_db():
    conn = sqlite3.connect(STATS_DB, timeout=30.0)
    cur  = conn.cursor()
    cur.execute('''
        CREATE TABLE IF NOT EXISTS candidate_metadata (
            id                   INTEGER PRIMARY KEY AUTOINCREMENT,
            filename             TEXT,
            full_name            TEXT,
            candidate_status     TEXT DEFAULT 'New',
            total_experience     REAL DEFAULT 0.0,
            pega_experience      REAL DEFAULT 0.0,
            skills               TEXT,
            certifications       TEXT,
            ctc                  TEXT,
            notice_period        TEXT,
            current_organization TEXT,
            email                TEXT,
            phone                TEXT,
            linkedin             TEXT,
            created_by           TEXT DEFAULT 'admin',
            timestamp            DATETIME DEFAULT CURRENT_TIMESTAMP,
            file_bytes           BYTEA
        )
    ''')
    cur.execute('''
        CREATE TABLE IF NOT EXISTS custom_columns (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            col_key TEXT UNIQUE,
            col_label TEXT,
            description TEXT
        )
    ''')
    cur.execute('''
        CREATE TABLE IF NOT EXISTS jobs (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            title TEXT NOT NULL,
            description TEXT NOT NULL,
            client_name TEXT,
            contact_name TEXT,
            client_phone TEXT,
            account_manager TEXT,
            assigned_recruiter TEXT,
            target_date TEXT,
            job_type TEXT,
            job_status TEXT,
            work_experience TEXT,
            industry TEXT,
            salary TEXT,
            required_skills TEXT,
            created_by TEXT DEFAULT 'admin',
            created_at DATETIME DEFAULT CURRENT_TIMESTAMP
        )
    ''')
    cur.execute('''
        CREATE TABLE IF NOT EXISTS job_candidates (
            job_id INTEGER,
            candidate_id INTEGER,
            ai_reason TEXT,
            status TEXT DEFAULT 'matched',
            PRIMARY KEY (job_id, candidate_id),
            FOREIGN KEY (job_id) REFERENCES jobs(id),
            FOREIGN KEY (candidate_id) REFERENCES candidate_metadata(id)
        )
    ''')
    cur.execute('''
        CREATE TABLE IF NOT EXISTS users (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            full_name TEXT,
            username TEXT UNIQUE NOT NULL,
            password_hash TEXT NOT NULL,
            role TEXT DEFAULT 'user',
            is_hr INTEGER DEFAULT 0,
            is_admin INTEGER DEFAULT 0,
            is_external INTEGER DEFAULT 0,
            hidden_fields TEXT DEFAULT '',
            email TEXT,
            mobile TEXT,
            is_approved INTEGER DEFAULT 0
        )
    ''')
    cur.execute('''
        CREATE TABLE IF NOT EXISTS change_requests (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            username TEXT NOT NULL,
            action_type TEXT NOT NULL,
            target_id TEXT,
            payload TEXT,
            description TEXT,
            status TEXT DEFAULT 'pending',
            created_at DATETIME DEFAULT CURRENT_TIMESTAMP
        )
    ''')
    try:
        cur.execute("ALTER TABLE users ADD COLUMN mobile TEXT")
        conn.commit()
    except Exception:
        conn.rollback()

    cur.execute('''
        CREATE TABLE IF NOT EXISTS activity_logs (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            username TEXT NOT NULL,
            action TEXT NOT NULL,
            timestamp DATETIME DEFAULT CURRENT_TIMESTAMP
        )
    ''')
    # Migration: fix activity_logs if it was created with old schema (user_id instead of username)
    try:
        cur.execute("PRAGMA table_info(activity_logs)")
        al_cols = [c[1] for c in cur.fetchall()]
        if 'username' not in al_cols:
            print("INFO: Migrating activity_logs table (adding username column)...")
            cur.execute("DROP TABLE IF EXISTS activity_logs")
            cur.execute('''
                CREATE TABLE activity_logs (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    username TEXT NOT NULL,
                    action TEXT NOT NULL,
                    timestamp DATETIME DEFAULT CURRENT_TIMESTAMP
                )
            ''')
            print("INFO: activity_logs table migrated successfully.")
    except Exception as e_al:
        print(f"WARNING: activity_logs migration check failed: {e_al}")

    cur.execute('''
        CREATE TABLE IF NOT EXISTS team_members (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            name TEXT UNIQUE NOT NULL,
            created_at DATETIME DEFAULT CURRENT_TIMESTAMP
        )
    ''')
    cur.execute('''
        CREATE TABLE IF NOT EXISTS job_shares (
            job_id INTEGER,
            username TEXT,
            PRIMARY KEY (job_id, username),
            FOREIGN KEY (job_id) REFERENCES jobs(id),
            FOREIGN KEY (username) REFERENCES users(username)
        )
    ''')
    cur.execute('''
        CREATE TABLE IF NOT EXISTS masked_keywords (
            keyword TEXT PRIMARY KEY
        )
    ''')
    cur.execute("SELECT COUNT(*) FROM masked_keywords")
    if cur.fetchone()[0] == 0:
        for kw in ["CSSA", "LSA"]:
            cur.execute("INSERT OR IGNORE INTO masked_keywords (keyword) VALUES (?)", (kw,))

    # Clean up any lingering processing states from an ungraceful shutdown
    # We update instead of deleting to prevent foreign key violations if the candidate was already linked to a job
    cur.execute("UPDATE candidate_metadata SET full_name = REPLACE(full_name, '⏳ Processing: ', '❌ Failed: ') WHERE full_name LIKE '%Processing%'")

    # Migration: add missing columns to candidate_metadata
    cur.execute("PRAGMA table_info(candidate_metadata)")
    existing = [c[1] for c in cur.fetchall()]
    new_cols = {
        'candidate_status': "TEXT DEFAULT 'New'",
        'source': "TEXT DEFAULT 'Resume Upload'",
        'cdh_exp': 'REAL DEFAULT 0.0',
        'expected_ctc': 'TEXT',
        'percentage_hike': 'TEXT',
        'candidate_interview_status': 'TEXT',
        'availability_in_days': 'INTEGER',
        'current_location': 'TEXT',
        'pref_locations': 'TEXT',
        'current_client': 'TEXT',
        'domain': 'TEXT',
        'tier': 'TEXT',
        'certification_version': 'TEXT',
        'current_organization': 'TEXT',
        'email': 'TEXT',
        'phone': 'TEXT',
        'linkedin': 'TEXT',
        'email_message': 'TEXT',
        'formatted_json': 'TEXT',
        'sender_email': 'TEXT',
        'is_qualified': 'INTEGER DEFAULT 0',
        'is_approved': 'INTEGER DEFAULT 1',
        'created_by': "TEXT DEFAULT 'admin'",
        'file_bytes': 'BYTEA',
        'file_url': 'TEXT'
    }
    for col, dtype in new_cols.items():
        if col not in existing:
            cur.execute(f"ALTER TABLE candidate_metadata ADD COLUMN {col} {dtype}")

    # Migration: add missing columns to jobs
    cur.execute("PRAGMA table_info(jobs)")
    existing_jobs = [c[1] for c in cur.fetchall()]
    new_jobs_cols = {
        'client_name': 'TEXT',
        'contact_name': 'TEXT',
        'client_phone': 'TEXT',
        'account_manager': 'TEXT',
        'assigned_recruiter': 'TEXT',
        'target_date': 'TEXT',
        'job_type': 'TEXT',
        'job_status': 'TEXT',
        'work_experience': 'TEXT',
        'industry': 'TEXT',
        'salary': 'TEXT',
        'required_skills': 'TEXT',
        'created_by': "TEXT DEFAULT 'admin'"
    }
    for col, dtype in new_jobs_cols.items():
        if col not in existing_jobs:
            cur.execute(f"ALTER TABLE jobs ADD COLUMN {col} {dtype}")
            
    # Migration: add missing columns to users
    # Wrapped in try/except because PostgreSQL raises an error if the column already exists
    cur.execute("PRAGMA table_info(users)")
    existing_users_cols = [c[1] for c in cur.fetchall()]
    users_migrate_cols = {
        'is_hr': 'INTEGER DEFAULT 0',
        'is_admin': 'INTEGER DEFAULT 0',
        'is_external': 'INTEGER DEFAULT 0',
        'hidden_fields': "TEXT DEFAULT ''",
        'email': 'TEXT',
        'is_approved': 'INTEGER DEFAULT 0',
        'full_name': 'TEXT',
    }
    for col, dtype in users_migrate_cols.items():
        if col not in existing_users_cols:
            try:
                cur.execute(f"ALTER TABLE users ADD COLUMN {col} {dtype}")
                conn.commit()
            except Exception:
                conn.rollback()

    # Seed default users if empty (do NOT wipe the table to preserve registered users!)
    cur.execute("SELECT COUNT(*) FROM users WHERE LOWER(username) = 'admin'")
    if cur.fetchone()[0] == 0:
        cur.execute("INSERT INTO users (full_name, username, password_hash, role, is_hr, is_admin, email, is_approved) VALUES (?, ?, ?, ?, ?, ?, ?, ?)",
                    ("System Admin", "admin", hashlib.sha256("admin".encode()).hexdigest(), "admin", 1, 1, "admin@gmail.com", 1))
    else:
        # Guarantee admin role/flags are set correctly
        try:
            cur.execute("UPDATE users SET role = 'admin', is_hr = 1, is_admin = 1, is_approved = 1 WHERE LOWER(username) = 'admin'")
        except Exception as e:
            print(f"Warning: could not update admin user flags: {e}")

    cur.execute("SELECT COUNT(*) FROM users WHERE LOWER(username) = 'user'")
    if cur.fetchone()[0] == 0:
        cur.execute("INSERT INTO users (full_name, username, password_hash, role, is_hr, is_admin, email, is_approved) VALUES (?, ?, ?, ?, ?, ?, ?, ?)",
                    ("Test User", "user", hashlib.sha256("user".encode()).hexdigest(), "user", 0, 0, "user@gmail.com", 1))

    # Seed the 4 recruiters in user management if they do not exist
    for m in ["Boopathi", "Praveen", "Harish", "Sabari"]:
        uname = m.lower()
        role = "admin" if uname == "sabari" else "user"
        is_admin_val = 1 if uname == "sabari" else 0
        is_hr_val = 1 if uname == "sabari" else 0
        cur.execute("SELECT COUNT(*) FROM users WHERE LOWER(username) = ?", (uname,))
        if cur.fetchone()[0] == 0:
            cur.execute("INSERT INTO users (full_name, username, password_hash, role, is_hr, is_admin, email, is_approved) VALUES (?, ?, ?, ?, ?, ?, ?, ?)",
                        (m, uname, hashlib.sha256(uname.encode()).hexdigest(), role, is_hr_val, is_admin_val, f"{uname}@gmail.com", 1))
        else:
            try:
                cur.execute("UPDATE users SET is_hr = ?, is_admin = ?, role = ?, is_approved = 1 WHERE LOWER(username) = ?", (is_hr_val, is_admin_val, role, uname))
            except Exception as e:
                print(f"Warning: could not update user {uname}: {e}")
        
    cur.execute("SELECT COUNT(*) FROM team_members")
    if cur.fetchone()[0] == 0:
        for m in ["Boopathi", "Praveen", "Harish", "Sabari"]:
            cur.execute("INSERT OR IGNORE INTO team_members (name) VALUES (?)", (m,))
        
    # Pre-approve seeded/default users
    try:
        cur.execute("UPDATE users SET is_approved = 1 WHERE LOWER(username) IN ('admin', 'user', 'boopathi', 'praveen', 'harish', 'sabari')")
    except Exception as e:
        print(f"Warning: could not pre-approve default users: {e}")

    # Fix emails of existing seeded/default users if they are NULL/empty
    for uname in ["admin", "user", "boopathi", "praveen", "harish", "sabari", "somasekhar9"]:
        try:
            cur.execute("UPDATE users SET email = ? WHERE LOWER(username) = ? AND (email IS NULL OR email = '')", (f"{uname}@gmail.com", uname))
        except Exception:
            pass

    # Create integrations_settings table
    cur.execute('''
    CREATE TABLE IF NOT EXISTS integrations_settings (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        email_enabled INTEGER DEFAULT 0,
        imap_host TEXT DEFAULT 'imap.gmail.com',
        imap_port INTEGER DEFAULT 993,
        smtp_host TEXT DEFAULT 'smtp.gmail.com',
        smtp_port INTEGER DEFAULT 587,
        email_user TEXT,
        email_pass TEXT,
        keywords TEXT DEFAULT 'resume,alamaticz,solution,job',
        drive_enabled INTEGER DEFAULT 0,
        reply_theme TEXT DEFAULT 'professional',
        reply_subject TEXT DEFAULT 'Re: {subject} (Ref: {ref})',
        reply_body_missing TEXT,
        reply_body_complete TEXT,
        gdrive_client_id TEXT,
        gdrive_client_secret TEXT,
        gdrive_refresh_token TEXT,
        gdrive_folder_id TEXT,
        gdrive_email TEXT,
        ms_client_id TEXT,
        ms_client_secret TEXT,
        ms_tenant_id TEXT DEFAULT 'common',
        gmail_enabled INTEGER DEFAULT 0,
        gmail_email TEXT,
        gmail_pass TEXT,
        outlook_enabled INTEGER DEFAULT 0,
        outlook_email TEXT,
        additional_emails TEXT DEFAULT '[]',
        theme_usage_counts TEXT DEFAULT '{}'
    )
    ''')

    # Check/add email template customization columns to integrations_settings (migration for existing DBs)
    cur.execute("PRAGMA table_info(integrations_settings)")
    existing_settings_cols = {c[1] for c in cur.fetchall()}
    if 'reply_theme' not in existing_settings_cols:
        cur.execute("ALTER TABLE integrations_settings ADD COLUMN reply_theme TEXT DEFAULT 'professional'")
    if 'theme_usage_counts' not in existing_settings_cols:
        cur.execute("ALTER TABLE integrations_settings ADD COLUMN theme_usage_counts TEXT DEFAULT '{}'")
    if 'reply_subject' not in existing_settings_cols:
        cur.execute("ALTER TABLE integrations_settings ADD COLUMN reply_subject TEXT DEFAULT 'Re: {subject} (Ref: {ref})'")
    if 'reply_body_missing' not in existing_settings_cols:
        cur.execute("ALTER TABLE integrations_settings ADD COLUMN reply_body_missing TEXT")
    if 'reply_body_complete' not in existing_settings_cols:
        cur.execute("ALTER TABLE integrations_settings ADD COLUMN reply_body_complete TEXT")
    if 'gdrive_client_id' not in existing_settings_cols:
        cur.execute("ALTER TABLE integrations_settings ADD COLUMN gdrive_client_id TEXT")
    if 'gdrive_client_secret' not in existing_settings_cols:
        cur.execute("ALTER TABLE integrations_settings ADD COLUMN gdrive_client_secret TEXT")
    if 'gdrive_refresh_token' not in existing_settings_cols:
        cur.execute("ALTER TABLE integrations_settings ADD COLUMN gdrive_refresh_token TEXT")
    if 'gdrive_folder_id' not in existing_settings_cols:
        cur.execute("ALTER TABLE integrations_settings ADD COLUMN gdrive_folder_id TEXT")
    if 'gdrive_email' not in existing_settings_cols:
        cur.execute("ALTER TABLE integrations_settings ADD COLUMN gdrive_email TEXT")
    if 'additional_emails' not in existing_settings_cols:
        cur.execute("ALTER TABLE integrations_settings ADD COLUMN additional_emails TEXT DEFAULT '[]'")
    if 'ms_client_id' not in existing_settings_cols:
        cur.execute("ALTER TABLE integrations_settings ADD COLUMN ms_client_id TEXT")
    if 'ms_client_secret' not in existing_settings_cols:
        cur.execute("ALTER TABLE integrations_settings ADD COLUMN ms_client_secret TEXT")
    if 'ms_tenant_id' not in existing_settings_cols:
        cur.execute("ALTER TABLE integrations_settings ADD COLUMN ms_tenant_id TEXT DEFAULT 'common'")
    if 'gmail_enabled' not in existing_settings_cols:
        cur.execute("ALTER TABLE integrations_settings ADD COLUMN gmail_enabled INTEGER DEFAULT 0")
    if 'gmail_email' not in existing_settings_cols:
        cur.execute("ALTER TABLE integrations_settings ADD COLUMN gmail_email TEXT")
    if 'gmail_pass' not in existing_settings_cols:
        cur.execute("ALTER TABLE integrations_settings ADD COLUMN gmail_pass TEXT")
    if 'outlook_enabled' not in existing_settings_cols:
        cur.execute("ALTER TABLE integrations_settings ADD COLUMN outlook_enabled INTEGER DEFAULT 0")
    if 'outlook_email' not in existing_settings_cols:
        cur.execute("ALTER TABLE integrations_settings ADD COLUMN outlook_email TEXT")
    if 'default_resume_template' not in existing_settings_cols:
        cur.execute("ALTER TABLE integrations_settings ADD COLUMN default_resume_template TEXT DEFAULT 'alamaticz'")

    # Seed integrations_settings if empty, loading from environment variables if present
    cur.execute("SELECT COUNT(*) FROM integrations_settings")
    if cur.fetchone()[0] == 0:
        env_user = os.getenv("SMTP_SENDER", "")
        env_pass = os.getenv("SMTP_PASSWORD", "")
        env_imap = os.getenv("IMAP_HOST", "imap.gmail.com")
        env_smtp = os.getenv("SMTP_HOST", "smtp.gmail.com")
        email_enabled = 1 if env_user and env_pass else 0
        cur.execute("""
        INSERT INTO integrations_settings (email_enabled, imap_host, imap_port, smtp_host, smtp_port, email_user, email_pass, keywords, drive_enabled)
        VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)
        """, (email_enabled, env_imap, 993, env_smtp, 587, env_user, env_pass, "resume,alamaticz,solution,job", 0))

    # Create processed_emails table
    cur.execute('''
    CREATE TABLE IF NOT EXISTS processed_emails (
        msg_uid TEXT PRIMARY KEY,
        processed_at DATETIME DEFAULT CURRENT_TIMESTAMP
    )
    ''')

    conn.commit()
    conn.close()

def get_candidates_list(username: Optional[str] = None, role: str = "user", is_hr_or_admin: bool = None) -> list:
    conn = sqlite3.connect(STATS_DB, timeout=30.0)
    conn.row_factory = sqlite3.Row
    cur = conn.cursor()
    try:
        if is_hr_or_admin is None:
            is_hr_or_admin = False
            if username:
                is_hr_or_admin = is_admin_or_hr(username)
        cols_to_select = (
            "id, filename, full_name, candidate_status, total_experience, pega_experience, "
            "skills, certifications, ctc, notice_period, current_organization, email, phone, "
            "linkedin, created_by, timestamp, source, cdh_exp, expected_ctc, percentage_hike, "
            "candidate_interview_status, availability_in_days, current_location, pref_locations, "
            "current_client, domain, tier, certification_version, "
            "sender_email, is_qualified, is_approved, file_url"
        )
        if role == "admin" or is_hr_or_admin or not username:
            cur.execute(f"SELECT {cols_to_select} FROM candidate_metadata ORDER BY timestamp DESC")
        else:
            cur.execute(f"SELECT {cols_to_select} FROM candidate_metadata WHERE LOWER(created_by) = LOWER(?) ORDER BY timestamp DESC", (username,))
        raw_rows = cur.fetchall()
        rows = []
        for r in raw_rows:
            try:
                rows.append(dict(r))
            except Exception as row_err:
                # If dict(r) fails, build dict manually from cursor description
                if cur.description:
                    col_names = [desc[0] for desc in cur.description]
                    rows.append(dict(zip(col_names, r)))
                else:
                    print(f"WARNING: Could not convert row to dict: {row_err}")
    except Exception as e:
        import traceback
        print(f"ERROR in get_candidates_list: {e}")
        traceback.print_exc()
        rows = []
    conn.close()
    return rows


def get_masked_keywords() -> list:
    conn = sqlite3.connect(STATS_DB, timeout=30.0)
    cur = conn.cursor()
    try:
        cur.execute("SELECT keyword FROM masked_keywords")
        keywords = [row[0] for row in cur.fetchall()]
    except Exception:
        keywords = []
    conn.close()
    return keywords

_PATTERN_CACHE = {}

def mask_text_with_keywords(text: str, keywords: list) -> str:
    if not text or not keywords:
        return text
    result = str(text)
    for kw in keywords:
        kw_strip = kw.strip()
        if not kw_strip:
            continue
        if kw_strip not in _PATTERN_CACHE:
            _PATTERN_CACHE[kw_strip] = re.compile(re.escape(kw_strip), re.IGNORECASE)
        result = _PATTERN_CACHE[kw_strip].sub("****", result)
    return result

def mask_candidate_record(candidate: dict, keywords: list) -> dict:
    masked = {}
    for k, v in candidate.items():
        if isinstance(v, str):
            masked[k] = mask_text_with_keywords(v, keywords)
        else:
            masked[k] = v
    return masked


def log_candidate(data: dict):
    conn = sqlite3.connect(STATS_DB, timeout=30.0)
    cur  = conn.cursor()
    
    cur.execute("PRAGMA table_info(candidate_metadata)")
    existing_cols = {c[1] for c in cur.fetchall()}
    
    # Filter data to only valid columns
    cols = [c for c in data.keys() if c in existing_cols and c != 'id']
    vals = []
    for c in cols:
        val = data.get(c)
        if c == 'file_bytes':
            vals.append(val)
        else:
            if val == '' or val is None:
                vals.append(None)
            elif isinstance(val, (dict, list)):
                import json
                vals.append(json.dumps(val))
            else:
                vals.append(val)
    
    existing_id = data.get('id')
            
    if existing_id is not None:
        if cols:
            set_clause = ', '.join([f"{c} = ?" for c in cols])
            cur.execute(f"UPDATE candidate_metadata SET {set_clause} WHERE id = ?", vals + [existing_id])
            new_id = existing_id
    else:
        new_id = None
        if cols:
            cur.execute(
                f"INSERT INTO candidate_metadata ({','.join(cols)}) VALUES ({','.join(['?']*len(cols))})",
                vals
            )
            new_id = cur.lastrowid
    conn.commit()
    conn.close()
    return new_id

init_db()

def log_activity_db(username: str, action: str):
    if not username:
        username = "unknown"
    try:
        conn = sqlite3.connect(STATS_DB, timeout=30.0)
        cur = conn.cursor()
        cur.execute("INSERT INTO activity_logs (username, action) VALUES (?, ?)", (username, action))
        conn.commit()
        conn.close()
    except Exception as e:
        print(f"Error logging activity: {e}")

# ═══════════════════════════════════════════════════════════════════════════════
# ENDPOINTS
# ═══════════════════════════════════════════════════════════════════════════════

# ── Health ─────────────────────────────────────────────────────────────────────
@app.get("/api/health")
def health():
    return {"status": "ok"}

class ActivityCreate(BaseModel):
    username: str
    action: str

@app.get("/api/activity")
def get_activity_logs():
    try:
        conn = sqlite3.connect(STATS_DB, timeout=30.0)
        conn.row_factory = dict_row_factory
        cur = conn.cursor()
        cur.execute("SELECT * FROM activity_logs ORDER BY timestamp DESC")
        logs = [dict(row) for row in cur.fetchall()]
        conn.close()
        return logs
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/api/activity")
def create_activity_log(req: ActivityCreate):
    log_activity_db(req.username, req.action)
    return {"status": "logged"}

@app.delete("/api/activity")
def clear_activity_logs(request: Request):
    username = request.headers.get("x-user-username")
    try:
        conn = sqlite3.connect(STATS_DB, timeout=30.0)
        cur = conn.cursor()
        cur.execute("DELETE FROM activity_logs")
        conn.commit()
        conn.close()
        log_activity_db(username or "unknown", "cleared the activity feed")
        return {"status": "cleared"}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

class TeamMemberCreate(BaseModel):
    name: str

@app.get("/api/team-members")
def list_team_members():
    try:
        conn = sqlite3.connect(STATS_DB, timeout=30.0)
        conn.row_factory = dict_row_factory
        cur = conn.cursor()
        cur.execute("SELECT * FROM team_members ORDER BY name ASC")
        members = [dict(row) for row in cur.fetchall()]
        conn.close()
        return members
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/api/team-members")
def create_team_member(req: TeamMemberCreate, request: Request):
    username = request.headers.get("x-user-username") or "admin"
    name = req.name.strip()
    if not name:
        raise HTTPException(status_code=400, detail="Name cannot be empty")
    try:
        conn = sqlite3.connect(STATS_DB, timeout=30.0)
        cur = conn.cursor()
        try:
            cur.execute("INSERT INTO team_members (name) VALUES (?)", (name,))
            conn.commit()
        except sqlite3.IntegrityError:
            raise HTTPException(status_code=400, detail=f"Team member '{name}' already exists")
        finally:
            conn.close()
        log_activity_db(username, f"added '{name}' to the recruiter persona matrix")
        return {"status": "added", "name": name}
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.delete("/api/team-members/{member_id}")
def delete_team_member(member_id: int, request: Request):
    username = request.headers.get("x-user-username") or "admin"
    try:
        conn = sqlite3.connect(STATS_DB, timeout=30.0)
        cur = conn.cursor()
        
        # Get member name first to log it
        cur.execute("SELECT name FROM team_members WHERE id = ?", (member_id,))
        row = cur.fetchone()
        if not row:
            conn.close()
            raise HTTPException(status_code=404, detail="Team member not found")
        member_name = row[0]
        
        cur.execute("DELETE FROM team_members WHERE id = ?", (member_id,))
        conn.commit()
        conn.close()
        
        log_activity_db(username, f"removed '{member_name}' from the recruiter persona matrix")
        return {"status": "deleted"}
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

def is_user_approved(username: Optional[str]) -> bool:
    if not username:
        return False
    # Seeded/default users are always approved
    if username.lower() in ("admin", "user", "boopathi", "praveen", "harish", "sabari"):
        return True
    conn = sqlite3.connect(STATS_DB, timeout=30.0)
    cur = conn.cursor()
    try:
        cur.execute("SELECT is_approved FROM users WHERE LOWER(username) = LOWER(?)", (username,))
        row = cur.fetchone()
        if row and row[0] == 1:
            return True
    except Exception:
        pass
    finally:
        conn.close()
    return False

def get_user_role(username: Optional[str]) -> str:
    if not username:
        return "user"
    conn = sqlite3.connect(STATS_DB, timeout=30.0)
    cur = conn.cursor()
    cur.execute("SELECT role, is_admin FROM users WHERE LOWER(username) = LOWER(?)", (username,))
    row = cur.fetchone()
    conn.close()
    if row:
        role, is_admin = row
        if is_admin == 1 or role == "admin":
            return "admin"
        return role
    return "user"

def create_change_request(username: str, action_type: str, target_id: Optional[str] = None, payload: Optional[str] = None, description: Optional[str] = None):
    conn = sqlite3.connect(STATS_DB, timeout=30.0)
    cur = conn.cursor()
    cur.execute("""
        INSERT INTO change_requests (username, action_type, target_id, payload, description, status)
        VALUES (?, ?, ?, ?, ?, 'pending')
    """, (username, action_type, target_id, payload, description))
    conn.commit()
    conn.close()

def get_candidate_name(candidate_id: int) -> str:
    conn = sqlite3.connect(STATS_DB, timeout=30.0)
    cur = conn.cursor()
    cur.execute("SELECT full_name FROM candidate_metadata WHERE id = ?", (candidate_id,))
    row = cur.fetchone()
    conn.close()
    return row[0] if row else f"ID {candidate_id}"

def get_job_title(job_id: int) -> str:
    conn = sqlite3.connect(STATS_DB, timeout=30.0)
    cur = conn.cursor()
    cur.execute("SELECT title FROM jobs WHERE id = ?", (job_id,))
    row = cur.fetchone()
    conn.close()
    return row[0] if row else f"ID {job_id}"

# ── Candidates ─────────────────────────────────────────────────────────────────
@app.get("/api/candidates")
def list_candidates(request: Request):
    username = request.headers.get("x-user-username")
    
    is_user_admin = False
    is_external = False
    role = "user"
    is_admin_or_hr_flag = False
    if username:
        conn = sqlite3.connect(STATS_DB, timeout=30.0)
        cur = conn.cursor()
        cur.execute("SELECT is_external, is_admin, role, is_hr FROM users WHERE LOWER(username) = LOWER(?)", (username,))
        row = cur.fetchone()
        conn.close()
        if row:
            is_external = (row[0] == 1)
            is_admin_or_hr_flag = (row[1] == 1 or row[3] == 1)
            is_user_admin = (row[1] == 1 or row[2] == "admin" or is_admin_or_hr_flag)
            role = row[2]
            if is_external:
                raise HTTPException(status_code=403, detail="Forbidden")
            
    if not is_user_approved(username):
        return []

    rows = get_candidates_list(username, role="admin" if is_user_admin else "user", is_hr_or_admin=is_admin_or_hr_flag)
    
    # Replace None values with empty string and remove binary file_bytes
    for row in rows:
        row.pop("file_bytes", None)
        for k, v in row.items():
            if v is None:
                row[k] = ""

    # Mask certifications for non-admin and non-HR users
    if not is_admin_or_hr_flag:
        for row in rows:
            row["certifications"] = "[HIDDEN]"
        keywords = get_masked_keywords()
        rows = [mask_candidate_record(row, keywords) for row in rows]
        
    rows = apply_user_hidden_fields(rows, username)
    return rows

class CustomColumn(BaseModel):
    col_key: str
    col_label: str
    description: str

@app.post("/api/columns")
def add_column(col: CustomColumn, request: Request):
    conn = sqlite3.connect(STATS_DB, timeout=30.0)
    cur = conn.cursor()
    clean_key = re.sub(r'[^a-zA-Z0-9_]', '', col.col_key.replace(' ', '_')).lower()
    
    cur.execute("PRAGMA table_info(candidate_metadata)")
    existing = [c[1] for c in cur.fetchall()]
    if clean_key in existing:
        conn.close()
        raise HTTPException(status_code=400, detail="Column already exists")
        
    try:
        cur.execute(f"ALTER TABLE candidate_metadata ADD COLUMN {clean_key} TEXT")
        cur.execute("INSERT INTO custom_columns (col_key, col_label, description) VALUES (?, ?, ?)", 
                    (clean_key, col.col_label, col.description))
        conn.commit()
    except Exception as e:
        conn.close()
        raise HTTPException(status_code=500, detail=str(e))
    conn.close()
    return {"status": "added", "col_key": clean_key}

@app.get("/api/columns")
def get_columns():
    conn = sqlite3.connect(STATS_DB, timeout=30.0)
    cur = conn.cursor()
    cur.execute("SELECT col_key, col_label FROM custom_columns")
    customs = [{"col_key": row[0], "col_label": row[1]} for row in cur.fetchall()]
    conn.close()
    
    base_cols = [
        {"col_key": "full_name", "col_label": "Name"},
        {"col_key": "candidate_status", "col_label": "Candidate Status"},
        {"col_key": "source", "col_label": "Source"},
        {"col_key": "total_experience", "col_label": "Total Exp"},
        {"col_key": "pega_experience", "col_label": "Pega Exp"},
        {"col_key": "cdh_exp", "col_label": "CDH Exp"},
        {"col_key": "ctc", "col_label": "Current CTC"},
        {"col_key": "expected_ctc", "col_label": "Exp CTC"},
        {"col_key": "percentage_hike", "col_label": "Percentage Hike"},
        {"col_key": "candidate_interview_status", "col_label": "Candidate Interview Status"},
        {"col_key": "availability_in_days", "col_label": "Availability in Days"},
        {"col_key": "notice_period", "col_label": "Notice Period"},
        {"col_key": "phone", "col_label": "Phone No"},
        {"col_key": "email", "col_label": "Email"},
        {"col_key": "linkedin", "col_label": "LinkedIn"},
        {"col_key": "current_location", "col_label": "Current Location"},
        {"col_key": "pref_locations", "col_label": "Pref Locations"},
        {"col_key": "current_organization", "col_label": "Current Employment"},
        {"col_key": "current_client", "col_label": "Current Client"},
        {"col_key": "domain", "col_label": "Domain"},
        {"col_key": "tier", "col_label": "Tier (Tier1 Tier2 Tier3)"},
        {"col_key": "certification_version", "col_label": "Certification Version"},
        {"col_key": "skills", "col_label": "Skills"},
        {"col_key": "certifications", "col_label": "Certifications"},
        {"col_key": "notescomments", "col_label": "Notes / Comments"}
    ]
    return {"base": base_cols, "custom": customs}

@app.delete("/api/columns/{col_key}")
def delete_column(col_key: str, request: Request):
    conn = sqlite3.connect(STATS_DB, timeout=30.0)
    cur = conn.cursor()
    try:
        cur.execute("DELETE FROM custom_columns WHERE col_key=?", (col_key,))
        try:
            cur.execute(f"ALTER TABLE candidate_metadata DROP COLUMN {col_key}")
        except Exception:
            pass # older sqlite versions might not support drop column
        conn.commit()
    except Exception as e:
        conn.close()
        raise HTTPException(status_code=500, detail=str(e))
    conn.close()
    return {"status": "deleted"}

@app.put("/api/candidates/{candidate_id}")
async def update_candidate(candidate_id: int, request: Request, background_tasks: BackgroundTasks):
    username = request.headers.get("x-user-username")
    if not is_user_approved(username):
        raise HTTPException(status_code=403, detail="Access denied. Your account is pending admin approval.")
    role = get_user_role(username)

    body = await request.json()
    conn = sqlite3.connect(STATS_DB, timeout=30.0)
    cur  = conn.cursor()
    
    # Check permission
    cur.execute("SELECT created_by FROM candidate_metadata WHERE id = ?", (candidate_id,))
    row = cur.fetchone()
    if not row:
        conn.close()
        raise HTTPException(status_code=404, detail="Candidate not found")
    created_by = row[0]
    if not is_admin_or_hr(username):
        if created_by and created_by.lower() != username.lower():
            conn.close()
            raise HTTPException(status_code=403, detail="Forbidden")

    cur.execute("PRAGMA table_info(candidate_metadata)")
    allowed_cols = [c[1] for c in cur.fetchall()]
    updates = {k: v for k, v in body.items() if k in allowed_cols and k != 'id' and v is not None and v != '[HIDDEN]'}
    
    # Server-side validation for numeric edits
    for int_col in ['notice_period', 'availability_in_days']:
        if int_col in updates and updates[int_col] != "":
            try:
                updates[int_col] = int(float(updates[int_col]))
            except ValueError:
                conn.close()
                return JSONResponse(status_code=400, content={"detail": f"{int_col} must be an integer"})
            
    for exp_col in ['total_experience', 'pega_experience', 'cdh_exp']:
        if exp_col in updates and updates[exp_col] != "":
            try:
                updates[exp_col] = float(updates[exp_col])
            except ValueError:
                conn.close()
                return JSONResponse(status_code=400, content={"detail": f"{exp_col} must be a number"})

    if not updates:
        conn.close()
        return {"status": "no changes"}



    set_clause = ", ".join(f"{k}=?" for k in updates)
    cur.execute(
        f"UPDATE candidate_metadata SET {set_clause} WHERE id=?",
        list(updates.values()) + [candidate_id]
    )
    conn.commit()
    conn.close()

    # Re-trigger matching if matching-related details have changed
    match_related_fields = {
        'full_name', 'total_experience', 'pega_experience', 'cdh_exp',
        'skills', 'certifications', 'current_location', 'pref_locations'
    }
    if any(field in updates for field in match_related_fields):
        background_tasks.add_task(match_candidate_to_all_jobs, candidate_id)

    cname = get_candidate_name(candidate_id)
    log_activity_db(username or "unknown", f"updated candidate '{cname}' details")

    return {"status": "updated"}

@app.delete("/api/candidates/{candidate_id}")
def delete_candidate(candidate_id: int, request: Request):
    username = request.headers.get("x-user-username")
    if not is_user_approved(username):
        raise HTTPException(status_code=403, detail="Access denied. Your account is pending admin approval.")
    role = get_user_role(username)
    
    conn = sqlite3.connect(STATS_DB, timeout=30.0)
    cur  = conn.cursor()
    
    cur.execute("SELECT full_name, created_by, filename, sender_email FROM candidate_metadata WHERE id = ?", (candidate_id,))
    row = cur.fetchone()
    if not row:
        conn.close()
        raise HTTPException(status_code=404, detail="Candidate not found")
    cname, created_by, filename, sender_email_val = row
    
    if not is_admin_or_hr(username):
        if created_by and created_by.lower() != username.lower():
            conn.close()
            raise HTTPException(status_code=403, detail="Forbidden")
            
    cur.execute("DELETE FROM job_candidates WHERE candidate_id=?", (candidate_id,))
    if filename:
        try:
            cur.execute("DELETE FROM langchain_pg_embedding WHERE cmetadata->>'source' = ?", (filename,))
        except Exception as embed_err:
            print(f"Failed to delete embeddings for {filename}: {embed_err}")
        # Block the source email from being re-processed by marking it as processed
        if filename.startswith('mail_'):
            cur.execute("INSERT OR IGNORE INTO processed_emails (msg_uid) VALUES (?)", (f"deleted_cand_{candidate_id}",))
    cur.execute("DELETE FROM candidate_metadata WHERE id=?", (candidate_id,))
    conn.commit()
    conn.close()
    log_activity_db(username or "unknown", f"deleted candidate '{cname}'")
    return {"status": "deleted"}

class BulkDeleteRequest(BaseModel):
    ids: list[int]

@app.post("/api/candidates/bulk-delete")
def bulk_delete_candidates(req: BulkDeleteRequest, request: Request):
    username = request.headers.get("x-user-username")
    if not is_user_approved(username):
        raise HTTPException(status_code=403, detail="Access denied. Your account is pending admin approval.")
    
    conn = sqlite3.connect(STATS_DB, timeout=30.0)
    cur = conn.cursor()
    
    deleted_names = []
    
    for cid in req.ids:
        cur.execute("SELECT full_name, created_by, filename FROM candidate_metadata WHERE id = ?", (cid,))
        row = cur.fetchone()
        if not row:
            continue
        cname, created_by, filename = row
        
        # Check permissions
        if not is_admin_or_hr(username):
            if created_by and created_by.lower() != username.lower():
                continue # Skip unauthorized deletions
                
        cur.execute("DELETE FROM job_candidates WHERE candidate_id=?", (cid,))
        if filename:
            try:
                cur.execute("DELETE FROM langchain_pg_embedding WHERE cmetadata->>'source' = ?", (filename,))
            except Exception as embed_err:
                print(f"Failed to delete embeddings for {filename}: {embed_err}")
            # Block the source email from being re-processed
            if filename.startswith('mail_'):
                cur.execute("INSERT OR IGNORE INTO processed_emails (msg_uid) VALUES (?)", (f"deleted_cand_{cid}",))
        cur.execute("DELETE FROM candidate_metadata WHERE id=?", (cid,))
        deleted_names.append(cname)
        
    conn.commit()
    conn.close()
    
    if deleted_names:
        log_activity_db(username or "unknown", f"bulk deleted {len(deleted_names)} candidates: {', '.join(deleted_names[:5])}")
    
    return {"status": "deleted", "count": len(deleted_names)}

@app.get("/api/candidates/{candidate_id}/jobs")
def get_candidate_jobs(candidate_id: int, request: Request):
    username = request.headers.get("x-user-username")
    if not is_user_approved(username):
        raise HTTPException(status_code=403, detail="Access denied. Your account is pending admin approval.")
    role = get_user_role(username)

    conn = sqlite3.connect(STATS_DB, timeout=30.0)
    conn.row_factory = dict_row_factory
    cur = conn.cursor()
    
    # Check candidate existence and permission
    cur.execute("SELECT created_by FROM candidate_metadata WHERE id = ?", (candidate_id,))
    cand_row = cur.fetchone()
    if not cand_row:
        conn.close()
        raise HTTPException(status_code=404, detail="Candidate not found")
        
    created_by = cand_row[0]
    if not is_admin_or_hr(username):
        if created_by and created_by.lower() != username.lower():
            conn.close()
            raise HTTPException(status_code=403, detail="Forbidden")

    cur.execute("""
        SELECT j.*, jc.status as match_status, jc.ai_reason
        FROM jobs j
        JOIN job_candidates jc ON j.id = jc.job_id
        WHERE jc.candidate_id = ?
    """, (candidate_id,))
    rows = [dict(r) for r in cur.fetchall()]
    conn.close()
    return rows

@app.get("/api/candidates/{candidate_id}/formatted-resume")
def get_formatted_resume_data(candidate_id: int, request: Request):
    username = request.headers.get("x-user-username")
    if not is_user_approved(username):
        raise HTTPException(status_code=403, detail="Access denied. Your account is pending admin approval.")
    role = get_user_role(username)

    conn = sqlite3.connect(STATS_DB, timeout=30.0)
    conn.row_factory = dict_row_factory
    cur = conn.cursor()
    cur.execute("SELECT * FROM candidate_metadata WHERE id = ?", (candidate_id,))
    row = cur.fetchone()
    conn.close()

    if not row:
        raise HTTPException(status_code=404, detail="Candidate not found")

    candidate = dict(row)
    created_by = candidate.get("created_by")
    if not is_admin_or_hr(username):
        if created_by and created_by.lower() != username.lower():
            raise HTTPException(status_code=403, detail="Forbidden")

    # Check cache first
    cached_data = None
    if candidate.get("formatted_json"):
        try:
            cached_data = json.loads(candidate.get("formatted_json"))
        except Exception:
            pass

    filename = candidate.get("filename", "")
    file_bytes = candidate.get("file_bytes")
    file_url = candidate.get("file_url")
    
    if isinstance(file_bytes, memoryview):
        file_bytes = file_bytes.tobytes()
        
    if filename and not file_bytes and file_url:
        import requests
        try:
            resp = requests.get(file_url, timeout=15)
            if resp.status_code == 200:
                file_bytes = resp.content
                print(f"INFO: Successfully fetched resume bytes from {file_url} for formatting")
            else:
                print(f"Warning: Failed to fetch resume bytes from {file_url}: status {resp.status_code}")
        except Exception as fetch_err:
            print(f"Error fetching resume bytes from {file_url}: {fetch_err}")
            
    temp_path = None
    if filename and file_bytes:
        import tempfile
        suffix = os.path.splitext(filename)[1]
        try:
            with tempfile.NamedTemporaryFile(suffix=suffix, delete=False) as tmp:
                tmp.write(file_bytes)
                temp_path = tmp.name
        except Exception as e:
            print(f"Error creating temp file for formatting: {e}")
            
    path = temp_path if temp_path else (os.path.join(UPLOAD_DIR, filename) if filename else None)
    
    # Calculate original page count if PDF
    original_page_count = 1
    if filename and path and os.path.exists(path) and filename.lower().endswith(".pdf"):
        try:
            import fitz
            doc = fitz.open(path)
            original_page_count = len(doc)
            doc.close()
        except Exception:
            pass

    if cached_data:
        if temp_path and os.path.exists(temp_path):
            try:
                os.remove(temp_path)
            except Exception:
                pass
        cached_data["original_page_count"] = original_page_count
        return cached_data
    
    # Try to load original text
    text = ""
    if filename and path and os.path.exists(path) and not filename.lower().endswith(('.xlsx', '.xls', '.csv')):
        try:
            if filename.lower().endswith(".pdf"):
                loader = SafePyMuPDFLoader(path)
            else:
                loader = Docx2txtLoader(path)
            docs = loader.load()
            text = "\n".join([d.page_content for d in docs])
        except Exception as e:
            print(f"Error loading resume file: {e}")
            
    if temp_path and os.path.exists(temp_path):
        try:
            os.remove(temp_path)
        except Exception:
            pass

    # Fallback to metadata if file text couldn't be loaded
    if not text:
        text = f"""
        Name: {candidate.get('full_name')}
        Experience: {candidate.get('total_experience')} years (Pega: {candidate.get('pega_experience')} years, CDH: {candidate.get('cdh_exp')} years)
        Current Employment: {candidate.get('current_organization')}
        Location: {candidate.get('current_location')}
        Skills: {candidate.get('skills')}
        Certifications: {candidate.get('certifications')}
        Preferred Location: {candidate.get('pref_locations')}
        CTC: {candidate.get('ctc')}
        Expected CTC: {candidate.get('expected_ctc')}
        """

    _, llm = get_models()

    prompt = f"""You are an expert resume formatter. Extract details from the candidate resume text below and structure them into the exact JSON template provided.
Make sure to fill all fields based on the candidate's actual data. If a field is not present in the resume text, leave it blank or default appropriately.

Template Structure:
{{
  "full_name": "Name of the candidate",
  "job_title": "Determine a standard job designation based on their experience and skills, e.g. 'PEGA - CERTIFIED SENIOR SYSTEM ANALYST', 'Pega Lead Business Architect', or similar. If they are a Pega certified system architect, use a format like 'PEGA - CERTIFIED SENIOR SYSTEM ARCHITECT' or 'PEGA - CERTIFIED SYSTEM ARCHITECT'",
  "profile_summary": "A professional profile summary. You MUST format this summary using the exact wording layout: '[X]+ years of experience in [domain/industry] with a proven track record of [key achievement]. Expertise in [Skill 1], [Skill 2], [Skill 3], and [Skill 4]. Experienced in [tools/platforms] and capable of [key capability]. Adept at working with [stakeholders] to deliver [outcomes].' Fill in the bracketed placeholders using the candidate's actual data. Do not leave brackets in the output, resolve them completely.",
  "domain_skills": [
    "A concise list of 4-6 specific domain or functional skill areas, e.g. 'Pega PRPC', 'Decisioning (CDH)', 'Integration Services', 'Data Modeling', 'Agile Methodologies'"
  ],
  "technical_skills": {{
    "primary": "Primary Tool/Platform: [list primary tool(s) and version(s) if known, e.g. 'Pega PRPC: v8.x, v7.x']",
    "languages": "Languages: [list programming/database languages, e.g. 'Java, SQL, XML']",
    "frontend": "Frontend: [list frontend technologies, e.g. 'HTML5, CSS3, JavaScript, React']",
    "others": "Others: [list other tools and platforms, e.g. 'Git, Jira, Azure, Maven']"
  }},
  "education": [
    {{
      "degree": "Degree name (e.g. 'B.E.', 'MCA', 'B.Tech')",
      "field": "Field of study (e.g. 'Computer Science & Engineering')",
      "school": "University or College Name",
      "years": "Year Start - Year End or Year of Completion (e.g. '2015 - 2019')"
    }}
  ],
  "certifications": [
    "List of certifications found, e.g. 'Pega Certified Senior System Architect (CSSA)', 'Certified System Architect (CSA)', 'Certified Pega Decisioning Consultant (CPDC)'"
  ],
  "work_experience": [
    {{
      "company": "Company Name",
      "dates": "Start Date - End Date (e.g. 'Jul 2019 - Present')",
      "role": "Role or Job Title",
      "bullets": [
        "Accomplishment or key responsibility bullet point (limit to 3-5 high-impact bullets per job)"
      ]
    }}
  ],
  "recognitions": [
    {{
      "date": "Month Year / Year (e.g. '2022' or 'Dec 2021')",
      "description": "Award title or recognition detail (e.g. 'Star of the Quarter for excellent project delivery')"
    }}
  ]
}}

Rules:
1. Return ONLY the valid JSON block. No explanation, no markdown backticks, no markdown blocks. Just raw valid JSON.
2. If certifications are empty, extract them from the text (do not use "[HIDDEN]").
3. All fields must be clean strings, numbers, arrays, or objects as defined in the template.
4. Formatting & Corrections: You MUST find and correct any spelling, grammatical, typographical, or indentation/alignment mistakes in the summary and experience bullets. Format the text block into professional, cohesive paragraphs and bullet lists.

Resume Text:
{text[:8000]}

JSON:"""

    try:
        resp = llm.invoke([HumanMessage(content=prompt)])
        raw = resp.content.strip()
        if "```json" in raw:
            raw = raw.split("```json")[1].split("```")[0].strip()
        elif "```" in raw:
            raw = raw.split("```")[1].split("```")[0].strip()
        start, end = raw.find('{'), raw.rfind('}')
        if start != -1 and end != -1:
            raw = raw[start:end+1]
        data = json.loads(raw)

        # Inject page count
        data["original_page_count"] = original_page_count

        # Cache the result in DB
        try:
            conn = sqlite3.connect(STATS_DB, timeout=30.0)
            cur = conn.cursor()
            cur.execute("UPDATE candidate_metadata SET formatted_json = ? WHERE id = ?", (json.dumps(data), candidate_id))
            conn.commit()
            conn.close()
        except Exception as e_cache:
            print(f"Error caching formatted resume: {e_cache}")

    except Exception as e:
        print(f"Error parsing resume via LLM: {e}")
        # Return fallback structured data
        data = {
            "full_name": candidate.get("full_name"),
            "job_title": "Pega Professional",
            "profile_summary": f"{candidate.get('total_experience', 0)} years of experience in IT with skills in {candidate.get('skills', '')}.",
            "domain_skills": [s.strip() for s in str(candidate.get("skills", "")).split(",") if s.strip()][:4],
            "technical_skills": {
                "primary": "Primary Tool/Platform: Pega",
                "languages": "Languages: Java, SQL",
                "frontend": "Frontend: HTML, CSS, JavaScript",
                "others": "Others: Git, Jira"
            },
            "education": [],
            "certifications": [c.strip() for c in str(candidate.get("certifications", "")).split(",") if c.strip()],
            "work_experience": [
                {
                    "company": candidate.get("current_organization") or "Current Employer",
                    "dates": "N/A",
                    "role": "Pega Developer",
                    "bullets": ["Contributed to application development and configuration."]
                }
            ],
            "recognitions": []
        }

    return data

@app.get("/api/candidates/{candidate_id}/export-docx")
def export_candidate_docx(candidate_id: int, request: Request):
    from fastapi.responses import StreamingResponse
    import io
    
    # Try importing docx locally so server doesn't crash on boot if package is installing
    try:
        from docx import Document
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        raise HTTPException(status_code=500, detail="python-docx library not installed. Please try again in a few seconds.")

    username = request.headers.get("x-user-username")
    if not is_user_approved(username):
        raise HTTPException(status_code=403, detail="Access denied. Your account is pending admin approval.")
    
    conn = sqlite3.connect(STATS_DB, timeout=30.0)
    conn.row_factory = dict_row_factory
    cur = conn.cursor()
    cur.execute("SELECT * FROM candidate_metadata WHERE id = ?", (candidate_id,))
    row = cur.fetchone()
    conn.close()

    if not row:
        raise HTTPException(status_code=404, detail="Candidate not found")
        
    candidate = dict(row)
    created_by = candidate.get("created_by")
    if not is_admin_or_hr(username):
        if created_by and created_by.lower() != username.lower():
            raise HTTPException(status_code=403, detail="Forbidden")
            
    doc = Document()
    
    # Alamaticz Styling
    title = doc.add_heading(candidate.get('full_name', 'Candidate Resume') or 'Candidate Resume', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Contact Info
    contact_p = doc.add_paragraph()
    contact_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_info = []
    if candidate.get('email'): contact_info.append(str(candidate.get('email')))
    if candidate.get('phone'): contact_info.append(str(candidate.get('phone')))
    if candidate.get('linkedin'): contact_info.append(str(candidate.get('linkedin')))
    if candidate.get('current_location'): contact_info.append(str(candidate.get('current_location')))
    
    if contact_info:
        run = contact_p.add_run(" | ".join(contact_info))
        run.font.size = Pt(10)
    
    doc.add_heading('Summary', level=1)
    exp = str(candidate.get('total_experience') or 'N/A')
    doc.add_paragraph(f"Total Experience: {exp}")
    if candidate.get('pega_experience'):
        doc.add_paragraph(f"Pega Experience: {candidate.get('pega_experience')}")
        
    if candidate.get('skills'):
        doc.add_heading('Skills', level=1)
        doc.add_paragraph(str(candidate.get('skills')))
        
    if candidate.get('certifications'):
        doc.add_heading('Certifications', level=1)
        doc.add_paragraph(str(candidate.get('certifications')))
        
    # Additional Details
    doc.add_heading('Additional Details', level=1)
    if candidate.get('current_organization'):
        doc.add_paragraph(f"Current Organization: {candidate.get('current_organization')}")
    if candidate.get('ctc'):
        doc.add_paragraph(f"Current CTC: {candidate.get('ctc')}")
    if candidate.get('expected_ctc'):
        doc.add_paragraph(f"Expected CTC: {candidate.get('expected_ctc')}")
    if candidate.get('notice_period'):
        doc.add_paragraph(f"Notice Period: {candidate.get('notice_period')}")
    
    # Save to memory
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    
    safe_name = str(candidate.get("full_name") or "Candidate").replace(" ", "_")
    headers = {
        'Content-Disposition': f'attachment; filename="Alamaticz_Resume_{safe_name}.docx"',
        'Access-Control-Expose-Headers': 'Content-Disposition'
    }
    
    return StreamingResponse(
        file_stream, 
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document", 
        headers=headers
    )


@app.put("/api/candidates/{candidate_id}/formatted-resume")
async def update_formatted_resume_data(candidate_id: int, request: Request):
    username = request.headers.get("x-user-username")
    if not is_user_approved(username):
        raise HTTPException(status_code=403, detail="Access denied. Your account is pending admin approval.")
    
    body = await request.json()
    
    conn = sqlite3.connect(STATS_DB, timeout=30.0)
    try:
        cur = conn.cursor()
        cur.execute("SELECT created_by FROM candidate_metadata WHERE id = ?", (candidate_id,))
        row = cur.fetchone()
        if not row:
            raise HTTPException(status_code=404, detail="Candidate not found")
            
        created_by = row[0]
        if not is_admin_or_hr(username):
            if created_by and created_by.lower() != username.lower():
                raise HTTPException(status_code=403, detail="Forbidden")
                
        cur.execute("UPDATE candidate_metadata SET formatted_json = ? WHERE id = ?", (json.dumps(body), candidate_id))
        conn.commit()
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
    finally:
        conn.close()
    
    cname = get_candidate_name(candidate_id)
    log_activity_db(username or "unknown", f"updated formatted resume for candidate '{cname}'")
    return {"status": "updated"}

# ── Upload & Extract ────────────────────────────────────────────────────────────
EXTRACT_PROMPT = """You are an expert resume and email parser. Extract EVERY piece of information from the resume and the email message (if provided) below.
Be extremely thorough — search all sections: Summary, Experience, Skills, Education, Certifications, Contact, Headers, Footers, and the email body.

Map the extracted data according to the database keys and their corresponding column headings/labels.

Return ONLY a valid JSON object with these exact keys (no extra text, no markdown, just raw JSON):

{{
  "full_name": "<Full name of the candidate. STRICT RULE: You MUST extract this ONLY from the resume document. NEVER extract it from the email body or headers. If no name is explicitly stated in the resume, leave this completely empty. Matching column heading: 'Name'>",
  "total_experience": <number — total years of professional work experience. MUST be extracted from the resume document itself (do not take this from the email body). Calculate if needed. 0 if fresher. Matching column heading: 'Total Exp'>,
  "pega_experience": <number — years of Pega PRPC/BPM experience specifically. If the email body specifies Pega experience (e.g. '5 yrs into pega'), MUST extract this value from the email body. 0 if none. Matching column heading: 'Pega Exp'>,
  "cdh_exp": <number — years of Pega Customer Decision Hub (CDH) experience specifically. If the email body specifies CDH experience, take that. Otherwise, calculate it from the resume (look at durations of CDH-related projects or roles). 0 if none. Matching column heading: 'CDH Exp'>,
  "ctc": "<Current CTC / current salary / current package. E.g. '8 LPA', '10.5 Lacs', '900000'. Look for salary, CTC, package, LPA, Lakhs. Empty if not found. Matching column heading: 'Current CTC'>",
  "expected_ctc": "<Expected CTC / expected salary / expected package. E.g. '12 LPA', '15 Lacs'. Look for expected CTC, ECTC, expected salary, expected package, expectation. Empty if not found. Matching column heading: 'Exp CTC'>",
  "percentage_hike": "<Expected percentage hike. Calculate if both current and expected CTC are present, else empty. Matching column heading: 'Percentage Hike'>",
  "candidate_interview_status": "<Default to empty string '' unless resume has recruiter notes on status. Matching column heading: 'Candidate Interview Status'>",
  "availability_in_days": <number — integer representing days until available. e.g. for 'immediate' use 0, '1 month' use 30. Return null if not found. Matching column heading: 'Availability in Days'>,
  "notice_period": "<Notice period in integer DAYS ONLY. E.g. for 'Immediate' return 0. For '1 month' return 30. ONLY output digits, no text. Matching column heading: 'Notice Period'>",
  "phone": "<Phone number with country code. Include only digits and + sign. Matching column heading: 'Phone No'>",
  "email": "<Email address. Matching column heading: 'Email'>",
  "linkedin": "<Full LinkedIn profile URL if found, else empty string. Matching column heading: 'LinkedIn'>",
  "current_location": "<Current city/location. Matching column heading: 'Current Location'>",
  "pref_locations": "<Preferred locations to work. Matching column heading: 'Pref Locations'>",
  "current_organization": "<Current or most recent employer name. Matching column heading: 'Current Employment'>",
  "current_client": "<Current client working for, if mentioned. Otherwise empty. Matching column heading: 'Current Client'>",
  "domain": "<Domain of expertise, e.g. Banking, Telecom, Healthcare. Extract from recent projects. Matching column heading: 'Domain'>",
  "tier": "<Tier1, Tier2, or Tier3 based on their college/university or companies. Guess if possible, or empty. Matching column heading: 'Tier (Tier1 Tier2 Tier3)'>",
  "certification_version": "<Version of Pega certifications, e.g. 8.6, 8.8. Empty if not found. Matching column heading: 'Certification Version'>",
  "skills": "<All technical skills comma-separated. Matching column heading: 'Skills'>",
  "certifications": "<All certifications and courses comma-separated. Matching column heading: 'Certifications'>",
  "source": "<Extract source of candidate profile if mentioned in resume, e.g. 'LinkedIn', 'Naukri', 'Resume Upload'. If none is found, return 'Resume Upload'. Matching column heading: 'Source'>"{custom_fields}
}}

Rules:
- cdh_exp, pega_experience, total_experience: Must be purely numerical floats or 0.
- notice_period: Must be strictly an integer (0, 15, 30, 60). No strings.
- availability_in_days: Integer number of days.
- If a field is not found, use an empty string "" (or null for numbers). Do not invent data.
- If both a resume and an email message are provided, extract and merge the details. The email message may contain more recent or specific details (such as current/expected CTC, notice period, or Pega/CDH experience); prioritize information from the email message in case of conflict. However, the total_experience (total years of experience) MUST always be extracted from the resume document itself (not the email body), as the resume lists the full work history.

Resume and Email Content:
{text}

JSON:"""


def pre_format_candidate_resume(candidate_id: int, text: str, filename: str):
    original_page_count = 1
    if filename:
        path = os.path.join(UPLOAD_DIR, filename)
        if os.path.exists(path) and filename.lower().endswith(".pdf"):
            try:
                import fitz
                doc = fitz.open(path)
                original_page_count = len(doc)
                doc.close()
            except Exception:
                pass

    _, llm = get_models()
    
    prompt = f"""You are an expert resume formatter. Extract details from the candidate resume text below and structure them into the exact JSON template provided.
Make sure to fill all fields based on the candidate's actual data. If a field is not present in the resume text, leave it blank or default appropriately.

Template Structure:
{{
  "full_name": "Name of the candidate",
  "job_title": "Determine a standard job designation based on their experience and skills, e.g. 'PEGA - CERTIFIED SENIOR SYSTEM ANALYST', 'Pega Lead Business Architect', or similar. If they are a Pega certified system architect, use a format like 'PEGA - CERTIFIED SENIOR SYSTEM ARCHITECT' or 'PEGA - CERTIFIED SYSTEM ARCHITECT'",
  "profile_summary": "A professional profile summary. You MUST format this summary using the exact wording layout: '[X]+ years of experience in [domain/industry] with a proven track record of [key achievement]. Expertise in [Skill 1], [Skill 2], [Skill 3], and [Skill 4]. Experienced in [tools/platforms] and capable of [key capability]. Adept at working with [stakeholders] to deliver [outcomes].' Fill in the bracketed placeholders using the candidate's actual data. Do not leave brackets in the output, resolve them completely.",
  "domain_skills": [
    "A concise list of 4-6 specific domain or functional skill areas, e.g. 'Pega PRPC', 'Decisioning (CDH)', 'Integration Services', 'Data Modeling', 'Agile Methodologies'"
  ],
  "technical_skills": {{
    "primary": "Primary Tool/Platform: [list primary tool(s) and version(s) if known, e.g. 'Pega PRPC: v8.x, v7.x']",
    "languages": "Languages: [list programming/database languages, e.g. 'Java, SQL, XML']",
    "frontend": "Frontend: [list frontend technologies, e.g. 'HTML5, CSS3, JavaScript, React']",
    "others": "Others: [list other tools and platforms, e.g. 'Git, Jira, Azure, Maven']"
  }},
  "education": [
    {{
      "degree": "Degree name (e.g. 'B.E.', 'MCA', 'B.Tech')",
      "field": "Field of study (e.g. 'Computer Science & Engineering')",
      "school": "University or College Name",
      "years": "Year Start - Year End or Year of Completion (e.g. '2015 - 2019')"
    }}
  ],
  "certifications": [
    "List of certifications found, e.g. 'Pega Certified Senior System Architect (CSSA)', 'Certified System Architect (CSA)', 'Certified Pega Decisioning Consultant (CPDC)'"
  ],
  "work_experience": [
    {{
      "company": "Company Name",
      "dates": "Start Date - End Date (e.g. 'Jul 2019 - Present')",
      "role": "Role or Job Title",
      "bullets": [
        "Accomplishment or key responsibility bullet point (limit to 3-5 high-impact bullets per job)"
      ]
    }}
  ],
  "recognitions": [
    {{
      "date": "Month Year / Year (e.g. '2022' or 'Dec 2021')",
      "description": "Award title or recognition detail (e.g. 'Star of the Quarter for excellent project delivery')"
    }}
  ]
}}

Rules:
1. Return ONLY the valid JSON block. No explanation, no markdown backticks, no markdown blocks. Just raw valid JSON.
2. If certifications are empty, extract them from the text (do not use "[HIDDEN]").
3. All fields must be clean strings, numbers, arrays, or objects as defined in the template.
4. Formatting & Corrections: You MUST find and correct any spelling, grammatical, typographical, or indentation/alignment mistakes in the summary and experience bullets. Format the text block into professional, cohesive paragraphs and bullet lists.

Resume Text:
{text[:8000]}

JSON:"""

    try:
        resp = llm.invoke([HumanMessage(content=prompt)])
        raw = resp.content.strip()
        if "```json" in raw:
            raw = raw.split("```json")[1].split("```")[0].strip()
        elif "```" in raw:
            raw = raw.split("```")[1].split("```")[0].strip()
        start, end = raw.find('{'), raw.rfind('}')
        if start != -1 and end != -1:
            raw = raw[start:end+1]
        data = json.loads(raw)

        # Inject page count
        data["original_page_count"] = original_page_count

        # Cache the result in DB
        conn = sqlite3.connect(STATS_DB, timeout=30.0)
        cur = conn.cursor()
        cur.execute("UPDATE candidate_metadata SET formatted_json = ? WHERE id = ?", (json.dumps(data), candidate_id))
        conn.commit()
        conn.close()
        print(f"INFO: Successfully pre-formatted and cached resume for candidate ID {candidate_id}")
    except Exception as e:
        print(f"Warning: Failed to pre-format candidate resume ID {candidate_id} in background: {e}")


def process_resume_logic(safe_name: str, path: str, is_approved: int = 1, username: str = "unknown", email_message: str = None, sender_email: str = None, file_url: str = None, placeholder_id: Optional[int] = None):
    try:
        _, llm = get_models()
        embeddings, _ = get_models()

        # Load document
        if safe_name.lower().endswith(".pdf"):
            loader = SafePyMuPDFLoader(path)
        else:
            loader = Docx2txtLoader(path)
        docs = loader.load()
        text = "\n".join([d.page_content for d in docs])

        # Fetch custom columns for the prompt
        conn = sqlite3.connect(STATS_DB, timeout=30.0)
        cur = conn.cursor()
        cur.execute("SELECT col_key, col_label, description FROM custom_columns")
        custom_cols = cur.fetchall()
        conn.close()
        
        custom_fields_str = ""
        if custom_cols:
            for col_key, col_label, desc in custom_cols:
                # Include the exact column heading / label and description to make extraction precise
                desc_str = f"Extract data corresponding to column heading '{col_label}'"
                if desc:
                    desc_str += f" ({desc})"
                custom_fields_str += f',\n  "{col_key}": "<{desc_str}>"'
                
        combined_text = text[:7000]
        if email_message:
            combined_text += f"\n\n=== EMAIL MESSAGE BODY ===\n{email_message}\n=========================="
            
        prompt_str = EXTRACT_PROMPT.format(text=combined_text, custom_fields=custom_fields_str)
        
        # Add a simple retry mechanism for rate limits
        max_retries = 5
        resp = None
        for attempt in range(max_retries):
            try:
                resp = llm.invoke([HumanMessage(content=prompt_str)])
                break
            except Exception as api_err:
                if ("429" in str(api_err) or "rate" in str(api_err).lower()) and attempt < max_retries - 1:
                    import time
                    time.sleep(20 + attempt * 10) # Exponential backoff for rate limits
                    continue
                raise api_err
                
        if resp is None:
            raise Exception("Failed to get response from AI model")

        raw  = resp.content.strip()

        if "```json" in raw:
            raw = raw.split("```json")[1].split("```")[0].strip()
        elif "```" in raw:
            raw = raw.split("```")[1].split("```")[0].strip()
        start, end = raw.find('{'), raw.rfind('}')
        if start != -1 and end != -1:
            raw = raw[start:end+1]

        data = json.loads(raw)

        # Safeguard against LLM placeholder hallucinations (e.g. John Doe)
        placeholder_names = {"john doe", "john doe's", "jane doe", "candidate name", "name of the candidate", "placeholder"}
        extracted_name = str(data.get('full_name', '')).strip()
        if not extracted_name or extracted_name.lower() in placeholder_names:
            data['full_name'] = ""
            
        if extracted_name.lower() in placeholder_names:
            for field in ['email', 'phone', 'linkedin', 'current_location', 'pref_locations']:
                val = str(data.get(field, '')).lower()
                if 'example.com' in val or 'johndoe' in val or '123456' in val or 'new york' in val:
                    data[field] = ""

        if username == "email_worker":
            data['source'] = 'uploaded from mail'
        if email_message:
            data['email_message'] = email_message
        if sender_email:
            data['sender_email'] = sender_email

        # -- Start Data Validation & Normalization --
        
        # Phone: Keep only digits and +
        if 'phone' in data and data['phone']:
            data['phone'] = re.sub(r'[^\d+]', '', str(data['phone']))
            
        # Email: Basic validation
        if 'email' in data and data['email']:
            if '@' not in str(data['email']):
                data['email'] = ""
                
        # Experience: Force float
        for exp_field in ['total_experience', 'pega_experience', 'cdh_exp']:
            if exp_field in data and data[exp_field] not in [None, ""]:
                try:
                    match = re.search(r'\d+(\.\d+)?', str(data[exp_field]))
                    data[exp_field] = float(match.group()) if match else 0.0
                except Exception:
                    data[exp_field] = 0.0
                    
        # Integer fields
        for num_field in ['notice_period', 'availability_in_days']:
            if num_field in data and data[num_field] not in [None, ""]:
                np_str = str(data[num_field]).lower()
                if 'immediate' in np_str:
                    data[num_field] = 0
                else:
                    try:
                        match = re.search(r'\d+', np_str)
                        val = int(match.group()) if match else ""
                        if match and 'month' in np_str:
                            val = val * 30
                        data[num_field] = val
                    except Exception:
                        data[num_field] = ""
        # -- End Data Validation & Normalization --

        # Only match if a candidate with the same filename already exists (excluding the current placeholder)
        conn = sqlite3.connect(STATS_DB, timeout=30.0)
        cur = conn.cursor()
        match_id = None
        if placeholder_id:
            cur.execute("SELECT id FROM candidate_metadata WHERE filename = ? AND id != ? LIMIT 1", (safe_name, placeholder_id))
        else:
            cur.execute("SELECT id FROM candidate_metadata WHERE filename = ? LIMIT 1", (safe_name,))
        row = cur.fetchone()
        if row:
            match_id = row[0]
                
        # Ensure file is uploaded to external storage if provider is configured and file_url not provided
        if not file_url:
            from app.services.storage import upload_to_external_storage, STORAGE_PROVIDER
            if STORAGE_PROVIDER != "local":
                url, err = upload_to_external_storage(path, safe_name)
                if not err:
                    file_url = url
                else:
                    print(f"Warning: External upload in background failed: {err}")

        if match_id:
            # Match found! Retrieve the file_bytes and file_url from the placeholder first
            cur.execute("PRAGMA table_info(candidate_metadata)")
            allowed_cols = {c[1] for c in cur.fetchall()}
            
            placeholder_bytes = None
            placeholder_url = None
            
            if 'file_url' in allowed_cols:
                if placeholder_id:
                    cur.execute("SELECT file_bytes, file_url FROM candidate_metadata WHERE id = ? LIMIT 1", (placeholder_id,))
                else:
                    cur.execute("SELECT file_bytes, file_url FROM candidate_metadata WHERE filename = ? LIMIT 1", (safe_name,))
                placeholder_row = cur.fetchone()
                if placeholder_row:
                    placeholder_bytes = placeholder_row[0]
                    placeholder_url = placeholder_row[1]
            else:
                if placeholder_id:
                    cur.execute("SELECT file_bytes FROM candidate_metadata WHERE id = ? LIMIT 1", (placeholder_id,))
                else:
                    cur.execute("SELECT file_bytes FROM candidate_metadata WHERE filename = ? LIMIT 1", (safe_name,))
                placeholder_row = cur.fetchone()
                if placeholder_row:
                    placeholder_bytes = placeholder_row[0]

            # Match found! Delete the placeholder record we created in upload_resume
            if placeholder_id:
                cur.execute("DELETE FROM candidate_metadata WHERE id = ?", (placeholder_id,))
            
            # Fetch the existing candidate metadata to merge values
            cur.execute("SELECT * FROM candidate_metadata WHERE id = ?", (match_id,))
            cur.row_factory = sqlite3.Row
            existing_row = row_to_dict(cur.fetchone())
            
            updates = {}
            for k, v in data.items():
                if k in allowed_cols and k != 'id' and k != 'filename' and k != 'file_bytes' and k != 'file_url':
                    if v is not None and v != "":
                        existing_val = existing_row.get(k)
                        if username == "email_worker" or existing_val is None or str(existing_val).strip() == "" or existing_val == 0.0 or existing_val == 0:
                            updates[k] = v
            # Always update or attach the filename to the matched candidate
            updates['filename'] = safe_name
            if placeholder_bytes and not file_url:
                updates['file_bytes'] = placeholder_bytes
            
            if file_url:
                updates['file_url'] = file_url
            elif placeholder_url:
                updates['file_url'] = placeholder_url
                
            if username == "email_worker":
                updates['source'] = 'uploaded from mail'
                if email_message:
                    updates['email_message'] = email_message
                if sender_email:
                    updates['sender_email'] = sender_email
            
            if updates:
                set_clause = ", ".join(f"{k}=?" for k in updates)
                cur.execute(f"UPDATE candidate_metadata SET {set_clause} WHERE id=?", list(updates.values()) + [match_id])
            
            candidate_id = match_id
            print(f"INFO: Auto-attached uploaded resume {safe_name} to existing candidate profile ID {match_id} ({data.get('full_name', '')})")
            conn.commit()
            conn.close()
        else:
            conn.close()
            data['filename'] = safe_name
            data['candidate_status'] = 'New'
            data['is_approved'] = is_approved
            data['created_by'] = username
            if file_url:
                data['file_url'] = file_url
            
            # If we uploaded externally, don't store bytes in DB
            if file_url:
                data['file_bytes'] = None
                
            if username == "email_worker":
                data['source'] = 'uploaded from mail'
                if email_message:
                    data['email_message'] = email_message
                if sender_email:
                    data['sender_email'] = sender_email
            if placeholder_id:
                data['id'] = placeholder_id
            candidate_id = log_candidate(data)
            
    except Exception as e:
        import traceback
        traceback.print_exc()
        error_msg = str(e)[:200]
        print(f"Error processing resume {safe_name}: {error_msg}")
        # Instead of deleting the placeholder (which makes the candidate disappear),
        # update it with a failed status so the user can see it failed and retry.
        try:
            conn = sqlite3.connect(STATS_DB, timeout=30.0)
            cur = conn.cursor()
            if placeholder_id:
                cur.execute(
                    "UPDATE candidate_metadata SET full_name = ?, candidate_status = ? WHERE id = ?",
                    (f"❌ Processing Failed: {safe_name}", "Error", placeholder_id)
                )
                conn.commit()
            conn.close()
        except Exception as db_err:
            print(f"Error updating placeholder status in DB: {db_err}")
        return

    # Add to PGVector
    try:
        # First, remove old embeddings for this resume to prevent duplicates
        try:
            conn = sqlite3.connect(STATS_DB, timeout=30.0)
            cur = conn.cursor()
            cur.execute("DELETE FROM langchain_pg_embedding WHERE cmetadata->>'source' = ?", (safe_name,))
            conn.commit()
            conn.close()
        except Exception as embed_err:
            print(f"Failed to delete old embeddings: {embed_err}")
            
        for d in docs:
            d.metadata['source'] = safe_name
        splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=100)
        chunks   = splitter.split_documents(docs)
        PGVector.from_documents(
            documents=chunks,
            embedding=embeddings,
            connection_string=os.getenv("POSTGRES_DATABASE_URL"),
            collection_name="resume_embeddings"
        )
    except Exception as e:
        pass

    # Automatically match this candidate to all active JDs in the database
    if candidate_id and is_approved == 1:
        match_candidate_to_all_jobs(candidate_id)
        # Pre-format resume in background to avoid on-the-fly LLM latency
        try:
            pre_format_candidate_resume(candidate_id, text, safe_name)
        except Exception as format_err:
            print(f"Warning: Background resume formatting failed: {format_err}")

def match_candidate_to_all_jobs(candidate_id: int):
    try:
        # Query candidate details
        conn = sqlite3.connect(STATS_DB, timeout=30.0)
        conn.row_factory = dict_row_factory
        cur = conn.cursor()
        cur.execute("SELECT * FROM candidate_metadata WHERE id = ?", (candidate_id,))
        row = cur.fetchone()
        if not row:
            conn.close()
            return
        data = dict(row)
        
        # Query all jobs (JDs)
        cand_creator = data.get('created_by')
        if cand_creator and cand_creator.lower() != "admin":
            cur.execute("SELECT id, title, description FROM jobs WHERE LOWER(created_by) = LOWER(?)", (cand_creator,))
        else:
            cur.execute("SELECT id, title, description FROM jobs")
        jobs = [dict(r) for r in cur.fetchall()]
        conn.close()

        if not jobs:
            return

        # Format JDs for the LLM
        jds_str_list = []
        for job in jobs:
            jds_str_list.append(f"JD ID: {job['id']}\nTitle: {job['title']}\nDescription: {job['description']}\n---")
        formatted_jds = "\n".join(jds_str_list)

        _, llm = get_models()

        prompt = f"""You are an expert technical recruiter. Evaluate candidate "{data.get('full_name', 'Candidate')}" against the following Job Descriptions "pin to pin".

Candidate Details:
- Skills: {data.get('skills', '')}
- Experience: {data.get('total_experience', 0)} years (Pega Experience: {data.get('pega_experience', 0)} years, CDH Experience: {data.get('cdh_exp', 0)} years)
- Certifications: {data.get('certifications', '')}
- Current Location: {data.get('current_location', '')}
- Preferred Locations: {data.get('pref_locations', '')}

Job Descriptions:
{formatted_jds}

Rules for evaluation:
1. Numeric Experience Matching: If a Job Description asks for "X+ years of experience", a candidate matches if their experience is greater than or equal to X.
   - For example: if Job Description requires "1+ years of experience in pega", then candidates with 3.0 years, 4.0 years, or 4.8 years of Pega experience all match perfectly because 3.0 >= 1.0, 4.0 >= 1.0, and 4.8 >= 1.0.
2. Certification Abbreviations:
   - CSSA is equivalent to any of: "PEGA Certified Senior System Architect", "Pega Certified Senior System Architect", "Certified Pega Senior System Architect", "Senior System Architect", or "CSSA".
   - CSA is equivalent to any of: "PEGA Certified System Architect", "Pega Certified System Architect", "Certified Pega System Architect", "System Architect", or "CSA".
   - LSA is equivalent to any of: "PEGA Certified Lead System Architect", "Pega Certified Lead System Architect", "Certified Pega Lead System Architect", "Lead System Architect", or "LSA".
3. Do not invent requirements. If the Job Description only mentions Pega experience, do NOT reject candidates for lacking CSSA or other unrelated certifications.
4. Location Matching: If the Job Description specifies a location requirement, a candidate matches if their Current Location or any of their Preferred Locations match the specified job location (e.g. if the Job Description mentions 'Chennai', a candidate with Current Location or Preferred Location 'Chennai' is a match).

For each Job Description, decide if the candidate is a good match based on the rules.
Respond with a JSON list of objects for matches only:
[
  {{"job_id": <job_id>, "reason": "<1-sentence explanation of fit based on specific experience, skills, and location>"}}
]
If there are no matches, return an empty list [].
Return ONLY the JSON block, no markdown, no other text."""

        # Add a retry mechanism for rate limits
        max_retries = 5
        resp = None
        for attempt in range(max_retries):
            try:
                resp = llm.invoke([HumanMessage(content=prompt)])
                break
            except Exception as api_err:
                if "429" in str(api_err) and attempt < max_retries - 1:
                    import time
                    time.sleep(20 + attempt * 10) # Exponential backoff for rate limits
                    continue
                raise api_err
                
        if resp is None:
            raise Exception("Failed to get response from AI model")

        raw  = resp.content.strip()

        if "```json" in raw:
            raw = raw.split("```json")[1].split("```")[0].strip()
        elif "```" in raw:
            raw = raw.split("```")[1].split("```")[0].strip()
        start, end = raw.find('['), raw.rfind(']')
        if start != -1 and end != -1:
            raw = raw[start:end+1]

        matches = json.loads(raw)
        
        conn = sqlite3.connect(STATS_DB, timeout=30.0)
        cur = conn.cursor()
        
        # Clear existing automatic matches for this candidate
        cur.execute("DELETE FROM job_candidates WHERE candidate_id = ? AND status = 'matched'", (candidate_id,))
        
        matched_any = False
        if matches:
            for match in matches:
                job_id = match.get("job_id")
                reason = match.get("reason", "Matched based on candidate details.")
                if job_id:
                    cur.execute("""
                        INSERT INTO job_candidates (job_id, candidate_id, ai_reason, status) 
                        VALUES (?, ?, ?, 'matched')
                        ON CONFLICT(job_id, candidate_id) DO UPDATE SET ai_reason = excluded.ai_reason
                    """, (job_id, candidate_id, reason))
                    matched_any = True
        
        # Automatically update qualification status
        cur.execute("SELECT COUNT(*) FROM job_candidates WHERE candidate_id = ?", (candidate_id,))
        cnt = cur.fetchone()[0]
        is_qualified = 1 if cnt > 0 else 0
        cur.execute("UPDATE candidate_metadata SET is_qualified = ? WHERE id = ?", (is_qualified, candidate_id))
        
        conn.commit()
        conn.close()
    except Exception as match_err:
        print(f"Error auto-matching candidate {candidate_id} to jobs: {match_err}")

def process_excel_file_logic(safe_name: str, path: str, username: str):
    conn = sqlite3.connect(STATS_DB, timeout=30.0)
    cur = conn.cursor()
    processed_candidate_ids = []
    try:
        wb = openpyxl.load_workbook(path, data_only=True)
        
        # Fetch custom columns
        cur.execute("SELECT col_key, col_label FROM custom_columns")
        custom_cols = {row[1].strip().lower(): row[0] for row in cur.fetchall()}
        
        # Mappings for common columns
        column_mappings = {
            'name': 'full_name',
            'full name': 'full_name',
            'candidate name': 'full_name',
            'candidate': 'full_name',
            'names': 'full_name',
            
            'phone': 'phone',
            'phone no': 'phone',
            'phone number': 'phone',
            'mobile': 'phone',
            'mobile number': 'phone',
            'mobile no': 'phone',
            'contact': 'phone',
            'contact number': 'phone',
            'contact no': 'phone',
            
            'email': 'email',
            'email id': 'email',
            'email address': 'email',
            
            'skills': 'skills',
            'key skills': 'skills',
            'technical skills': 'skills',
            
            'experience': 'total_experience',
            'exp': 'total_experience',
            'total exp': 'total_experience',
            'total experience': 'total_experience',
            'work experience': 'total_experience',
            
            'pega experience': 'pega_experience',
            'pega exp': 'pega_experience',
            
            'cdh experience': 'cdh_exp',
            'cdh exp': 'cdh_exp',
            
            'current ctc': 'ctc',
            'ctc': 'ctc',
            'salary': 'ctc',
            
            'expected ctc': 'expected_ctc',
            
            'percentage hike': 'percentage_hike',
            'hike': 'percentage_hike',
            
            'notice period': 'notice_period',
            'np': 'notice_period',
            
            'current location': 'current_location',
            'location': 'current_location',
            
            'preferred locations': 'pref_locations',
            'preferred location': 'pref_locations',
            'pref locations': 'pref_locations',
            
            'current organization': 'current_organization',
            'current employer': 'current_organization',
            'employer': 'current_organization',
            'current employment': 'current_organization',
            
            'current client': 'current_client',
            
            'domain': 'domain',
            'tier': 'tier',
            'certification version': 'certification_version',
            'certifications': 'certifications',
            
            # Additional maps for thorough Excel parsing
            'linkedin': 'linkedin',
            'linkedin url': 'linkedin',
            'linkedin profile': 'linkedin',
            'notes': 'notescomments',
            'comments': 'notescomments',
            'notes/comments': 'notescomments',
            'notescomments': 'notescomments',
            'feedback': 'notescomments',
            'candidate status': 'candidate_status',
            'candidate_status': 'candidate_status',
            'status': 'candidate_status',
            'candidate interview status': 'candidate_interview_status',
            'interview status': 'candidate_interview_status',
            'availability': 'availability_in_days',
            'availability in days': 'availability_in_days',
            'availability_in_days': 'availability_in_days',
            'source': 'source',
            'candidate source': 'source',
            'how did you find us': 'source'
        }
        
        # Load existing candidates from DB
        cur.execute("SELECT id, full_name, email, phone FROM candidate_metadata WHERE LOWER(created_by) = LOWER(?)", (username,))
        existing_candidates = [
            {
                "id": r[0],
                "full_name": r[1] or "",
                "email": r[2] or "",
                "phone": r[3] or ""
            }
            for r in cur.fetchall()
        ]
        
        # Iterate over all sheets
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            # Find the header row (iterating up to 10 rows)
            header_row_idx = None
            headers = []
            for r_idx in range(1, min(ws.max_row + 1, 11)):
                row_vals = [ws.cell(row=r_idx, column=c_idx).value for c_idx in range(1, ws.max_column + 1)]
                non_empty = [v for v in row_vals if v is not None]
                if len(non_empty) >= 2:
                    is_header = False
                    for val in non_empty:
                        val_str = str(val).strip().lower()
                        if val_str in column_mappings or val_str in custom_cols:
                            is_header = True
                            break
                    if is_header:
                        header_row_idx = r_idx
                        headers = [str(h).strip() if h is not None else "" for h in row_vals]
                        break
            
            if not header_row_idx:
                header_row_idx = 1
                headers = [str(ws.cell(row=1, column=c_idx).value).strip() if ws.cell(row=1, column=c_idx).value is not None else "" for c_idx in range(1, ws.max_column + 1)]
            
            # Map headers to DB columns
            mapped_cols = {}
            for idx, h in enumerate(headers):
                h_lower = h.lower()
                if h_lower in custom_cols:
                    mapped_cols[idx] = custom_cols[h_lower]
                elif h_lower in column_mappings:
                    mapped_cols[idx] = column_mappings[h_lower]
            
            # Process rows starting from header_row_idx + 1
            for row_idx in range(header_row_idx + 1, ws.max_row + 1):
                row_data = {}
                is_row_empty = True
                for col_idx in range(1, len(headers) + 1):
                    val = ws.cell(row=row_idx, column=col_idx).value
                    if val is not None:
                        is_row_empty = False
                    if (col_idx - 1) in mapped_cols:
                        db_col = mapped_cols[col_idx - 1]
                        row_data[db_col] = str(val).strip() if val is not None else ""
                
                if is_row_empty:
                    continue
                
                # Check that we have at least one identifying field
                if not row_data.get('full_name') and not row_data.get('email') and not row_data.get('phone'):
                    continue
                
                name = row_data.get('full_name', '')
                email = row_data.get('email', '')
                phone = row_data.get('phone', '')
                
                norm_email = normalize_email(email)
                norm_phone = normalize_phone(phone)
                
                match_id = None
                
                # Normalize numeric/experience fields
                if 'total_experience' in row_data and row_data['total_experience'] != "":
                    try:
                        row_data['total_experience'] = float(row_data['total_experience'])
                    except ValueError:
                        row_data['total_experience'] = 0.0
                if 'pega_experience' in row_data and row_data['pega_experience'] != "":
                    try:
                        row_data['pega_experience'] = float(row_data['pega_experience'])
                    except ValueError:
                        row_data['pega_experience'] = 0.0
                if 'cdh_exp' in row_data and row_data['cdh_exp'] != "":
                    try:
                        row_data['cdh_exp'] = float(row_data['cdh_exp'])
                    except ValueError:
                        row_data['cdh_exp'] = 0.0
                if 'notice_period' in row_data and row_data['notice_period'] != "":
                    try:
                        digits = "".join(c for c in row_data['notice_period'] if c.isdigit())
                        row_data['notice_period'] = int(digits) if digits else ""
                    except Exception:
                        row_data['notice_period'] = ""
                if 'availability_in_days' in row_data and row_data['availability_in_days'] != "":
                    try:
                        digits = "".join(c for c in row_data['availability_in_days'] if c.isdigit())
                        row_data['availability_in_days'] = int(digits) if digits else ""
                    except Exception:
                        row_data['availability_in_days'] = ""
                
                cur.execute("PRAGMA table_info(candidate_metadata)")
                allowed_cols = {c[1] for c in cur.fetchall()}
                db_data = {k: v for k, v in row_data.items() if k in allowed_cols and k != 'id'}
                if 'source' not in db_data or not db_data['source']:
                    db_data['source'] = 'Excel Import'
                
                if match_id:
                    # Update existing record
                    set_clause = ", ".join(f"{k}=?" for k in db_data)
                    cur.execute(f"UPDATE candidate_metadata SET {set_clause} WHERE id=?", list(db_data.values()) + [match_id])
                    candidate_id = match_id
                else:
                    # Insert new record
                    db_data['is_approved'] = 1
                    db_data['candidate_status'] = 'New'
                    db_data['filename'] = ""
                    db_data['created_by'] = username
                    cols_list = list(db_data.keys())
                    vals_list = list(db_data.values())
                    cur.execute(
                        f"INSERT INTO candidate_metadata ({','.join(cols_list)}) VALUES ({','.join(['?']*len(cols_list))})",
                        vals_list
                    )
                    candidate_id = cur.lastrowid
                    existing_candidates.append({
                        "id": candidate_id,
                        "full_name": name,
                        "email": email,
                        "phone": phone
                    })
                
                # Match candidate to all jobs later
                processed_candidate_ids.append(candidate_id)
                
    except Exception as e:
        print(f"Error parsing Excel file {safe_name}: {e}")
        if os.path.exists(path):
            try:
                os.remove(path)
                print(f"Cleaned up failed Excel file: {path}")
            except Exception as file_err:
                print(f"Error cleaning up failed Excel file {path}: {file_err}")
    finally:
        try:
            cur.execute("DELETE FROM candidate_metadata WHERE filename = ? AND full_name LIKE '⏳ Parsing Excel:%'", (safe_name,))
            conn.commit()
        except Exception as db_err:
            print(f"Error clearing Excel placeholder or committing in finally: {db_err}")
        try:
            conn.close()
        except Exception:
            pass

    for cid in processed_candidate_ids:
        try:
            match_candidate_to_all_jobs(cid)
        except Exception as match_err:
            print(f"Error matching candidate {cid}: {match_err}")

def process_excel_file(safe_name: str, path: str, username: str = "unknown"):
    try:
        with _processing_lock:
            process_excel_file_logic(safe_name, path, username)
    finally:
        try:
            if os.path.exists(path):
                os.remove(path)
        except Exception:
            pass

def process_resume(safe_name: str, path: str, is_approved: int = 1, username: str = "unknown", email_message: str = None, sender_email: str = None, file_url: str = None, placeholder_id: Optional[int] = None):
    # Use a lock to ensure only one resume is processed at a time
    # This prevents Render memory crashes (OOM), Groq Rate Limits, and SQLite database locks
    try:
        with _processing_lock:
            process_resume_logic(safe_name, path, is_approved, username, email_message, sender_email, file_url, placeholder_id)
    finally:
        try:
            if os.path.exists(path):
                os.remove(path)
        except Exception:
            pass


@app.post("/api/upload")
async def upload_resume(request: Request, background_tasks: BackgroundTasks, file: UploadFile = File(...)):
    username = request.headers.get("x-user-username")
    if not is_user_approved(username):
        raise HTTPException(status_code=403, detail="Access denied. Your account is pending admin approval.")
    is_approved = 1

    import uuid
    # Save file
    safe_name = f"{uuid.uuid4().hex}_{file.filename}"
    path = os.path.join(UPLOAD_DIR, safe_name)
    with open(path, "wb") as f:
        content = await file.read()
        f.write(content)

    ext = os.path.splitext(safe_name.lower())[1]
    if ext in ['.xlsx', '.xls', '.csv']:
        log_activity_db(username or "unknown", f"uploaded candidate excel sheet '{safe_name}'")
        background_tasks.add_task(process_excel_file, safe_name, path, username or "unknown")
        return {"status": "processing", "message": "Excel sheet uploaded and is processing in the background."}
    else:
        # Placeholder while processing in background
        # Note: We save file_bytes temporarily. The background task will upload it
        # to external storage (if configured) and clear file_bytes to save space.
        placeholder_id = log_candidate({
            "filename": safe_name, 
            "full_name": f"⏳ Processing: {safe_name}", 
            "is_approved": is_approved, 
            "created_by": username or "unknown", 
            "file_bytes": content,
            "file_url": None
        })
        
        # Process asynchronously
        log_activity_db(username or "unknown", f"uploaded resume '{safe_name}'")
        background_tasks.add_task(process_resume, safe_name, path, is_approved, username or "unknown", None, None, None, placeholder_id)

        return {"status": "processing", "message": "Resume uploaded and is processing in the background."}


@app.post("/api/jobs/parse-document")
async def parse_jd_document(request: Request, file: UploadFile = File(...)):
    username = request.headers.get("x-user-username")
    if not is_user_approved(username):
        raise HTTPException(status_code=403, detail="Access denied. Your account is pending admin approval.")
    
    # Save file temporarily in a temp folder under UPLOAD_DIR
    safe_name = file.filename
    temp_dir = os.path.join(UPLOAD_DIR, "temp_jds")
    os.makedirs(temp_dir, exist_ok=True)
    path = os.path.join(temp_dir, safe_name)
    
    try:
        with open(path, "wb") as f:
            content = await file.read()
            f.write(content)
            
        ext = os.path.splitext(safe_name.lower())[1]
        text = ""
        
        if ext == ".pdf":
            try:
                loader = SafePyMuPDFLoader(path)
                docs = loader.load()
                text = "\n".join([d.page_content for d in docs])
            except Exception as pdf_err:
                raise HTTPException(status_code=400, detail=f"Failed to parse PDF document: {str(pdf_err)}")
        elif ext in [".docx", ".doc"]:
            try:
                loader = Docx2txtLoader(path)
                docs = loader.load()
                text = "\n".join([d.page_content for d in docs])
            except Exception as docx_err:
                raise HTTPException(status_code=400, detail=f"Failed to parse Word document: {str(docx_err)}")
        else:
            raise HTTPException(status_code=400, detail="Unsupported file format. Please upload a PDF or Word (.docx) document.")
            
        if not text.strip():
            raise HTTPException(status_code=400, detail="The uploaded document appears to be empty or unreadable.")
            
        # Extract fields using LLM
        _, llm = get_models()
        
        prompt = f"""You are an expert recruitment assistant. Extract details from the following Job Description document and structure them into the exact JSON template provided.
Make sure to fill all fields based on the actual document. If a field is not present or mentioned in the document, use an empty string "" (or default appropriately).

Template Structure:
{{
  "title": "<The job title, e.g., 'Pega CSSA', 'Pega Business Analyst'>",
  "description": "<The complete or summarized job description text>",
  "client_name": "<Client name if mentioned, else empty string>",
  "client_phone": "<Client phone number if mentioned, else empty string>",
  "contact_name": "<Client contact name if mentioned, else empty string>",
  "account_manager": "<Account manager if mentioned, else empty string>",
  "assigned_recruiter": "<Assigned recruiter if mentioned, else empty string>",
  "target_date": "<Target date formatted as YYYY-MM-DD if mentioned, else empty string>",
  "job_type": "<Must be one of: 'Full time', 'Part time', 'Contract', 'Temporary'. Map appropriately. Default to 'Full time' if not specified>",
  "job_status": "<Always set to 'In-progress'>",
  "work_experience": "<Must be exactly one of: 'None', 'Fresher', '1-3 years', '3-5 years', '5+ years'. Estimate/map from the requirements>",
  "industry": "<Must be exactly one of: 'None', 'IT', 'Finance', 'Healthcare', 'Telecom', 'Other'. Map appropriately>",
  "salary": "<Salary package or CTC budget if mentioned, e.g., '10 LPA', '15 - 20 LPA', else empty string>",
  "required_skills": "<Comma-separated list of required technical skills, e.g., 'Pega, CSSA, Java'>"
}}

Rules:
1. Return ONLY the valid JSON block. No explanation, no markdown backticks, no markdown blocks. Just raw valid JSON.
2. For multiple-choice fields (job_type, work_experience, industry), you MUST choose one of the allowed options. If you cannot determine, use the default.
3. Be as accurate and thorough as possible based on the text.

Job Description Text:
{text[:8000]}

JSON:"""

        max_retries = 3
        resp = None
        for attempt in range(max_retries):
            try:
                resp = llm.invoke([HumanMessage(content=prompt)])
                break
            except Exception as api_err:
                if "429" in str(api_err) and attempt < max_retries - 1:
                    import time
                    time.sleep(3)
                    continue
                raise api_err
                
        if resp is None:
            raise Exception("Failed to get response from AI model")
            
        raw = resp.content.strip()
        if "```json" in raw:
            raw = raw.split("```json")[1].split("```")[0].strip()
        elif "```" in raw:
            raw = raw.split("```")[1].split("```")[0].strip()
            
        start, end = raw.find('{'), raw.rfind('}')
        if start != -1 and end != -1:
            raw = raw[start:end+1]
            
        data = json.loads(raw)
        
        # Clean up temporary file
        try:
            if os.path.exists(path):
                os.remove(path)
        except Exception:
            pass
            
        return data
        
    except HTTPException:
        try:
            if os.path.exists(path):
                os.remove(path)
        except Exception:
            pass
        raise
    except Exception as e:
        try:
            if os.path.exists(path):
                os.remove(path)
        except Exception:
            pass
        raise HTTPException(status_code=500, detail=f"Failed to parse and extract JD: {str(e)}")

# ── Chat ──────────────────────────────────────────────────────────────────────
class ChatRequest(BaseModel):
    message: str

CONVERSATIONAL_KW = [
    "hi", "hello", "hey", "how are you", "what's up",
    "good morning", "good afternoon", "good evening",
    "thanks", "thank you", "bye", "goodbye", "who are you",
]

def _is_conversational(msg: str) -> bool:
    m = msg.strip().lower()
    if m in CONVERSATIONAL_KW:
        return True
    # Short messages (≤4 words) with no data keywords → conversational
    data_kw = ["candidate", "resume", "experience", "pega", "skill", "cert",
               "ctc", "notice", "company", "email", "phone", "show", "list",
               "find", "give", "who", "which", "how many", "count", "year",
               "join", "immediate", "work", "hire", "select"]
    if len(m.split()) <= 4 and not any(k in m for k in data_kw):
        return True
    return False

def extract_search_filters(prompt: str, llm) -> dict:
    filter_extraction_prompt = f"""You are an HR database query assistant.
Extract search parameters from the following HR question into a JSON object.

HR Question: "{prompt}"

JSON Schema:
{{
  "name_query": "string or null (extract a candidate name if the user is asking about a specific person, e.g., 'Naresh' or 'Gopinath')",
  "skills_keywords": "array of strings or null (extract skills or certifications, e.g., ['pega', 'cssa', 'csa', 'lsa', 'clda'])",
  "min_pega_exp": "number or null (minimum years of Pega experience requested)",
  "min_total_exp": "number or null (minimum years of total experience requested)",
  "notice_period_max": "number or null (maximum notice period in days)",
  "current_location": "string or null (location name, e.g., 'Hyderabad')"
}}

Respond ONLY with the JSON object. Do not include any explanation or markdown formatting."""

    try:
        resp = llm.invoke([HumanMessage(content=filter_extraction_prompt)])
        ans = resp.content.strip()
        if ans.startswith("```json"):
            ans = ans[7:]
        if ans.endswith("```"):
            ans = ans[:-3]
        return json.loads(ans.strip())
    except Exception as e:
        print(f"Error extracting search filters: {e}")
        return {}

def query_candidates_by_filters(filters: dict, username: str, role: str) -> list:
    conn = sqlite3.connect(STATS_DB, timeout=30.0)
    conn.row_factory = dict_row_factory
    cur = conn.cursor()
    
    cols_to_select = (
        "id, filename, full_name, candidate_status, total_experience, pega_experience, "
        "skills, certifications, ctc, notice_period, current_organization, email, phone, "
        "linkedin, created_by, timestamp, source, cdh_exp, expected_ctc, percentage_hike, "
        "candidate_interview_status, availability_in_days, current_location, pref_locations, "
        "current_client, domain, tier, certification_version, "
        "sender_email, is_qualified, is_approved, file_url"
    )
    query = f"SELECT {cols_to_select} FROM candidate_metadata WHERE 1=1"
    params = []
    
    # Non-admins and non-HRs can only see their own candidates
    if not is_admin_or_hr(username) and username:
        query += " AND LOWER(created_by) = LOWER(?)"
        params.append(username)
        
    name_q = filters.get("name_query")
    if name_q:
        query += " AND (full_name LIKE ? OR current_organization LIKE ? OR email LIKE ?)"
        params.append(f"%{name_q}%")
        params.append(f"%{name_q}%")
        params.append(f"%{name_q}%")
        
    skills = filters.get("skills_keywords")
    if skills and isinstance(skills, list):
        for sk in skills:
            if sk.strip():
                query += " AND (skills LIKE ? OR certifications LIKE ?)"
                params.append(f"%{sk.strip()}%")
                params.append(f"%{sk.strip()}%")
            
    min_pega = filters.get("min_pega_exp")
    if min_pega is not None:
        try:
            query += " AND pega_experience >= ?"
            params.append(float(min_pega))
        except ValueError:
            pass
            
    min_tot = filters.get("min_total_exp")
    if min_tot is not None:
        try:
            query += " AND total_experience >= ?"
            params.append(float(min_tot))
        except ValueError:
            pass

    max_notice = filters.get("notice_period_max")
    if max_notice is not None:
        try:
            query += " AND (CAST(notice_period AS INTEGER) <= ? OR notice_period LIKE '%immediate%')"
            params.append(int(max_notice))
        except ValueError:
            pass

    loc = filters.get("current_location")
    if loc:
        query += " AND (current_location LIKE ? OR pref_locations LIKE ?)"
        params.append(f"%{loc}%")
        params.append(f"%{loc}%")

    query += " ORDER BY timestamp DESC"
    
    try:
        cur.execute(query, params)
        rows = [dict(r) for r in cur.fetchall()]
    except Exception as e:
        print(f"Error running database query: {e}")
        rows = []
    finally:
        conn.close()
    return rows


def _find_matching_rows(rows: list, names: list) -> list:
    """Return matching exact DB rows for the given candidate names."""
    lower_names = [n.strip().lower() for n in names]
    
    # Try exact match first
    matched = [r for r in rows if r.get('full_name', '').strip().lower() in lower_names]
    
    if not matched:
        # Try partial match
        matched = []
        for r in rows:
            fn_lower = r.get('full_name', '').strip().lower()
            if any((ln in fn_lower or fn_lower in ln) and ln != '' for ln in lower_names):
                matched.append(r)
                
    records = []
    for r in matched:
        records.append({
            "name":             r.get("full_name", ""),
            "total_experience": r.get("total_experience", 0),
            "pega_experience":  r.get("pega_experience", 0),
            "skills":           r.get("skills", ""),
            "certifications":   r.get("certifications", ""),
            "ctc":              r.get("ctc", ""),
            "notice_period":    r.get("notice_period", ""),
            "organization":     r.get("current_organization", ""),
            "email":            r.get("email", ""),
            "phone":            r.get("phone", ""),
        })
    return records

@app.post("/api/chat")
def chat(body: ChatRequest, request: Request):
    username = request.headers.get("x-user-username")
    if not is_user_approved(username):
        return {
            "type": "text",
            "answer": "Your account access is currently pending administrator approval. Please contact your system administrator to view and query candidate data."
        }

    conn = sqlite3.connect(STATS_DB, timeout=30.0)
    cur = conn.cursor()
    cur.execute("SELECT is_admin, role FROM users WHERE LOWER(username) = LOWER(?)", (username,))
    user_row = cur.fetchone()
    conn.close()
    is_user_admin = False
    if user_row:
        is_user_admin = (user_row[0] == 1 or user_row[1] == "admin" or is_admin_or_hr(username))

    global _embeddings, _llm, _models_loading
    if _embeddings is None or _llm is None:
        if _models_loading:
            return {"type": "text", "answer": "⏳ Hire AI is currently warming up and downloading models. This may take a few minutes depending on your internet connection. Please try again shortly!"}
        else:
            threading.Thread(target=get_models, daemon=True).start()
            return {"type": "text", "answer": "⏳ Hire AI is starting its engines... Please try again in a few seconds."}

    embeddings, llm = _embeddings, _llm
    prompt  = body.message.strip()
    p_lower = prompt.lower()

    # ── Route 1: Conversational ──────────────────────────────────────────────
    if _is_conversational(p_lower):
        resp = llm.invoke([HumanMessage(content=(
            "You are Hire AI, an intelligent HR recruitment assistant for Alamaticz Solutions. "
            "Respond warmly and professionally.\n\nUser: " + prompt
        ))])
        return {"type": "text", "answer": resp.content}

    # ── Route 2: Structured (SQLite) — always try this first ─────────────────
    try:
        # Step 2a: Extract query filters
        filters = extract_search_filters(prompt, llm)
        
        # Step 2b: Query SQLite with extracted filters
        user_role = "admin" if is_user_admin else "user"
        matched_candidates = query_candidates_by_filters(filters, username, user_role)
        
        # Fall back to all candidates if query returned nothing, but limit to top 15
        if not matched_candidates:
            all_candidates = get_candidates_list(username, role=user_role)
            matched_candidates = all_candidates[:15]
        else:
            matched_candidates = matched_candidates[:15]
            
        if matched_candidates:
            for r in matched_candidates:
                for col in ['total_experience', 'pega_experience']:
                    if col in r:
                        try:
                            r[col] = float(r[col]) if str(r[col]).strip() != "" else 0
                        except:
                            r[col] = 0

            # Build compact summary of matches
            candidate_lines = []
            for r in matched_candidates:
                line = (
                    f"Name: {r.get('full_name','?')} | "
                    f"Total Exp: {r.get('total_experience',0)} yrs | "
                    f"Pega Exp: {r.get('pega_experience',0)} yrs | "
                    f"Skills: {r.get('skills','') or 'none'} | "
                    f"Certifications: {r.get('certifications','') or 'none'} | "
                    f"CTC: {r.get('ctc','') or '?'} | "
                    f"Notice: {r.get('notice_period','') or '?'} | "
                    f"Company: {r.get('current_organization','') or '?'} | "
                    f"Email: {r.get('email','') or '?'} | "
                    f"Phone: {r.get('phone','') or '?'}"
                )
                candidate_lines.append(line)
            candidates_text = "\n".join(candidate_lines)

            total_db_count = len(get_candidates_list(username, role=user_role))

            filter_prompt = f"""You are an expert HR data analyst. I have the following candidate database:

Total candidates in database: {total_db_count}

{candidates_text}

HR question: "{prompt}"

            Instructions:
            1. Read the question carefully.
            2. If the question asks to LIST / SHOW / FIND candidates:
               - Return EXACTLY this JSON format (no other text before or after):
               MATCH_RESULT: {{"type":"table","matched_names":["Candidate Name 1", "Candidate Name 2"],"intro":"Here are the candidates:"}}
            3. If the question asks for a COUNT or YES/NO:
               - Return EXACTLY: MATCH_RESULT: {{"type":"count","answer":"There are 3 candidates matching your criteria."}}
            4. If the question is about a SPECIFIC candidate's detail:
               - Return EXACTLY: MATCH_RESULT: {{"type":"table","matched_names":["Candidate Name"],"intro":"Here are the details:"}}
            5. IMPORTANT: Use ONLY exact full_name values from the database above. Do NOT invent names.
            6. 'Pega Exp: 0 yrs' = ZERO Pega experience.
            
            Respond with ONLY the MATCH_RESULT line, nothing else."""

            resp = llm.invoke([HumanMessage(content=filter_prompt)])
            ans  = resp.content.strip()

            if "MATCH_RESULT:" in ans:
                json_str = ans.split("MATCH_RESULT:")[1].strip()
                # Extract JSON object
                s, e = json_str.find('{'), json_str.rfind('}')
                if s != -1 and e != -1:
                    result = json.loads(json_str[s:e+1])

                    if result.get("type") == "count":
                        return {"type": "text", "answer": result.get("answer", ans)}

                    if result.get("type") == "table":
                        names = result.get("matched_names", [])
                        intro = result.get("intro", "Here are the matching candidates:")
                        if names:
                            matched_rows = _find_matching_rows(matched_candidates, names)
                            if matched_rows:
                                return {"type": "table", "answer": intro, "rows": matched_rows}
                        if matched_candidates:
                            fallback_rows = _find_matching_rows(matched_candidates, [r['full_name'] for r in matched_candidates[:5]])
                            if fallback_rows:
                                return {"type": "table", "answer": intro, "rows": fallback_rows}
                        return {"type": "text", "answer": "No candidates match your query. Try a different filter."}
    except Exception as e:
        print(f"Route 2 error, falling back to RAG: {e}")
        pass  # Fall through to RAG

    # ── Route 3: RAG (ChromaDB) — for very specific resume content ────────────
    # Check if PGVector has any records
    has_vectors = False
    try:
        conn = sqlite3.connect(STATS_DB, timeout=30.0)
        cur = conn.cursor()
        cur.execute("SELECT COUNT(*) FROM langchain_pg_embedding")
        has_vectors = (cur.fetchone()[0] > 0)
        conn.close()
    except Exception:
        has_vectors = False

    if not has_vectors:
        return {"type": "text", "answer": "No resumes uploaded yet. Please upload resumes first."}

    try:
        db = PGVector(
            connection_string=os.getenv("POSTGRES_DATABASE_URL"),
            embedding_function=embeddings,
            collection_name="resume_embeddings"
        )
        retriever = db.as_retriever(search_kwargs={"k": 6})
        qa_prompt = PromptTemplate.from_template("""You are Hire AI, an expert HR recruitment assistant.
Answer using ONLY the resume context below. Be specific and structured.

Resume Context:
{context}

HR Question: {question}

Answer:""")

        def format_docs(docs):
            return "\n\n".join(d.page_content for d in docs)

        rag_chain = (
            {"context": retriever | format_docs, "question": RunnablePassthrough()}
            | qa_prompt
            | llm
            | StrOutputParser()
        )
        ans = rag_chain.invoke(prompt)
        return {"type": "text", "answer": ans}
    except Exception as e:
        return {"type": "text", "answer": f"Search error: {e}"}


# ── JD Matching ────────────────────────────────────────────────────────────────
class JDMatchRequest(BaseModel):
    job_description: str

@app.post("/api/match-jd")
def match_jd(req: JDMatchRequest, request: Request):
    jd = req.job_description.strip()
    if not jd:
        raise HTTPException(status_code=400, detail="Empty Job Description")
    
    username = request.headers.get("x-user-username") or ""
    is_user_admin = False
    
    # Check if PGVector has any records and determine user role
    has_vectors = False
    try:
        conn = sqlite3.connect(STATS_DB, timeout=30.0)
        cur = conn.cursor()
        
        # Check user role
        if username:
            cur.execute("SELECT is_admin, role FROM users WHERE LOWER(username) = LOWER(?)", (username,))
            row = cur.fetchone()
            if row:
                is_user_admin = (row[0] == 1 or row[1] == "admin" or is_admin_or_hr(username))
                
        # Check vector db
        cur.execute("SELECT COUNT(*) FROM langchain_pg_embedding")
        has_vectors = (cur.fetchone()[0] > 0)
        conn.close()
    except Exception:
        has_vectors = False

    if not has_vectors:
        return {"matches": []}

    embeddings, llm = get_models()
    db = PGVector(
        connection_string=os.getenv("POSTGRES_DATABASE_URL"),
        embedding_function=embeddings,
        collection_name="resume_embeddings"
    )
    
    try:
        docs = db.similarity_search(jd, k=15)
    except Exception:
        return {"matches": []}
    
    matched_sources = set()
    for d in docs:
        if 'source' in d.metadata:
            matched_sources.add(d.metadata['source'])
            
    if not matched_sources:
        return {"matches": []}

    conn = sqlite3.connect(STATS_DB, timeout=30.0)
    conn.row_factory = dict_row_factory
    cur = conn.cursor()
    placeholders = ",".join("?" * len(matched_sources))
    # Use tuple for psycopg2 compatibility with IN clauses
    cur.execute(f"SELECT * FROM candidate_metadata WHERE filename IN ({placeholders})", tuple(matched_sources))
    db_rows = [dict(r) for r in cur.fetchall()]
    conn.close()

    if not is_user_admin:
        # Allow rows where created_by matches username OR where created_by is null/empty (shared records)
        db_rows = [r for r in db_rows if not r.get('created_by') or r['created_by'].lower() == username.lower()]

    if not db_rows:
        return {"matches": []}

    candidate_lines = []
    for r in db_rows:
        candidate_lines.append(
            f"Name: {r.get('full_name')} | "
            f"Total Experience: {r.get('total_experience')} yrs | "
            f"Pega Experience: {r.get('pega_experience')} yrs | "
            f"CDH Experience: {r.get('cdh_exp')} yrs | "
            f"Skills: {r.get('skills')} | "
            f"Certifications: {r.get('certifications')} | "
            f"Current Location: {r.get('current_location')} | "
            f"Preferred Locations: {r.get('pref_locations')}"
        )
    
    import concurrent.futures
    import json
    
    def evaluate_batch(batch_lines):
        prompt = f"""You are an expert technical recruiter. Evaluate the following candidates against the Job Description "pin to pin".

Job Description:
{jd[:2000]}

Candidates to evaluate:
{chr(10).join(batch_lines)}

Rules for evaluation:
1. Numeric Experience Matching: If a Job Description asks for "X+ years of experience", a candidate matches if their experience is greater than or equal to X.
   - For example: if Job Description requires "1+ years of experience in pega", then candidates with 3.0 years, 4.0 years, or 4.8 years of Pega experience all match perfectly because 3.0 >= 1.0, 4.0 >= 1.0, and 4.8 >= 1.0.
2. Certification Abbreviations:
   - CSSA is equivalent to any of: "PEGA Certified Senior System Architect", "Pega Certified Senior System Architect", "Certified Pega Senior System Architect", "Senior System Architect", or "CSSA".
   - CSA is equivalent to any of: "PEGA Certified System Architect", "Pega Certified System Architect", "Certified Pega System Architect", "System Architect", or "CSA".
   - LSA is equivalent to any of: "PEGA Certified Lead System Architect", "Pega Certified Lead System Architect", "Certified Pega Lead System Architect", "Lead System Architect", or "LSA".
3. Do not invent requirements. If the Job Description only mentions Pega experience, do NOT reject candidates for lacking CSSA or other unrelated certifications.
4. Location Matching: If the Job Description specifies a location requirement, a candidate matches if their Current Location or any of their Preferred Locations match the specified job location (e.g. if the Job Description mentions 'Chennai', a candidate with Current Location or Preferred Location 'Chennai' is a match).

For EACH candidate, decide if they match based on the rules. If they match, provide a 1-sentence explanation of why they are a good fit.
Format your response exactly as a JSON list of objects:
[
  {{
    "name": "Candidate Name",
    "reason": "1-sentence explanation of why they fit based on their specific experience, skills, and location"
  }}
]
Return ONLY the raw JSON block, no markdown, no other text."""
        try:
            resp = llm.invoke([HumanMessage(content=prompt)])
            raw = resp.content.strip()
            if "```json" in raw:
                raw = raw.split("```json")[1].split("```")[0].strip()
            elif "```" in raw:
                raw = raw.split("```")[1].split("```")[0].strip()
            
            start, end = raw.find('['), raw.rfind(']')
            if start != -1 and end != -1:
                raw = raw[start:end+1]
            
            ai_reasons = json.loads(raw)
            return {str(item.get("name", "")).strip().lower(): item.get("reason", "") for item in ai_reasons}
        except Exception:
            return {}

    reason_map = {}
    batch_size = 5
    batches = [candidate_lines[i:i + batch_size] for i in range(0, len(candidate_lines), batch_size)]
    with concurrent.futures.ThreadPoolExecutor(max_workers=3) as executor:
        for res in executor.map(evaluate_batch, batches):
            reason_map.update(res)

    matches = []
    for r in db_rows:
        name = str(r.get('full_name', '')).strip().lower()
        reason = reason_map.get(name, "Matched based on resume content similarity.")
        if reason == "Matched based on resume content similarity.":
            for k, v in reason_map.items():
                if k in name or name in k:
                    reason = v
                    break
        r['ai_reason'] = reason
        matches.append(r)
        
    return {"matches": matches}


# ── Jobs & JDs ─────────────────────────────────────────────────────────────────
class JobCreate(BaseModel):
    title: str
    description: str
    client_name: Optional[str] = ""
    contact_name: Optional[str] = ""
    client_phone: Optional[str] = ""
    account_manager: Optional[str] = ""
    assigned_recruiter: Optional[str] = ""
    target_date: Optional[str] = ""
    job_type: Optional[str] = ""
    job_status: Optional[str] = ""
    work_experience: Optional[str] = ""
    industry: Optional[str] = ""
    salary: Optional[str] = ""
    required_skills: Optional[str] = ""

class JobStatusUpdate(BaseModel):
    status: Optional[str] = None
    ai_reason: Optional[str] = None

class RegisterRequest(BaseModel):
    full_name: str
    username: str
    password: str
    email: str

class LoginRequest(BaseModel):
    username: str
    password: str

class UserPermissionsUpdate(BaseModel):
    is_hr: int
    is_admin: int
    is_external: Optional[int] = 0
    hidden_fields: Optional[str] = ""
    is_approved: Optional[int] = None

def has_digit(s: str) -> bool:
    return any(c.isdigit() for c in s)

@app.get("/api/jobs")
def list_jobs(request: Request):
    username = request.headers.get("x-user-username")
    if not is_user_approved(username):
        return []
        
    conn = sqlite3.connect(STATS_DB, timeout=30.0)
    conn.row_factory = dict_row_factory
    cur = conn.cursor()
    
    # Check if requesting user is external
    is_external = False
    is_user_admin = False
    if username:
        cur.execute("SELECT is_external, is_admin, role FROM users WHERE LOWER(username) = LOWER(?)", (username,))
        row = cur.fetchone()
        if row:
            is_external = (row['is_external'] == 1)
            is_user_admin = (row['is_admin'] == 1 or row['role'] == "admin" or is_admin_or_hr(username))
            
    if is_external:
        cur.execute("""
            SELECT j.* FROM jobs j
            JOIN job_shares js ON j.id = js.job_id
            WHERE LOWER(js.username) = LOWER(?)
            ORDER BY j.id DESC
        """, (username,))
    else:
        if is_user_admin:
            cur.execute("SELECT * FROM jobs ORDER BY id DESC")
        else:
            cur.execute("SELECT * FROM jobs WHERE LOWER(created_by) = LOWER(?) ORDER BY id DESC", (username,))
        
    jobs = [dict(r) for r in cur.fetchall()]
    for job in jobs:
        job_id = job['id']
        cur.execute("SELECT status, COUNT(*) as cnt FROM job_candidates WHERE job_id = ? GROUP BY status", (job_id,))
        counts = {r['status']: r['cnt'] for r in cur.fetchall()}
        job['matched_count'] = counts.get('matched', 0)
        job['selected_count'] = counts.get('selected', 0)
        
        # Get shared usernames
        cur.execute("SELECT username FROM job_shares WHERE job_id = ?", (job_id,))
        job['shared_with'] = [r['username'] for r in cur.fetchall()]
        
    conn.close()
    return jobs

@app.post("/api/jobs")
def create_job(job: JobCreate, request: Request):
    username = request.headers.get("x-user-username") or "admin"
    if not is_user_approved(username):
        raise HTTPException(status_code=403, detail="Access denied. Your account is pending admin approval.")
    conn = sqlite3.connect(STATS_DB, timeout=30.0)
    cur = conn.cursor()
    cur.execute("""
        INSERT INTO jobs (
            title, description, client_name, contact_name, client_phone, account_manager,
            assigned_recruiter, target_date, job_type, job_status,
            work_experience, industry, salary, required_skills, created_by
        ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
    """, (
        job.title, job.description, job.client_name, job.contact_name, job.client_phone, job.account_manager,
        job.assigned_recruiter, job.target_date, job.job_type, job.job_status,
        job.work_experience, job.industry, job.salary, job.required_skills, username
    ))
    job_id = cur.lastrowid
    conn.commit()
    conn.close()
    
    try:
        match_candidates_for_job(job_id)
    except Exception as e:
        print(f"Error matching candidates for job {job_id}: {e}")
        
    conn = sqlite3.connect(STATS_DB, timeout=30.0)
    conn.row_factory = dict_row_factory
    cur = conn.cursor()
    cur.execute("SELECT * FROM jobs WHERE id = ?", (job_id,))
    updated_job = row_to_dict(cur.fetchone())
    
    cur.execute("SELECT status, COUNT(*) as cnt FROM job_candidates WHERE job_id = ? GROUP BY status", (job_id,))
    counts = {r['status']: r['cnt'] for r in cur.fetchall()}
    updated_job['matched_count'] = counts.get('matched', 0)
    updated_job['selected_count'] = counts.get('selected', 0)
    updated_job['shared_with'] = []
    username = request.headers.get("x-user-username")
    log_activity_db(username or "unknown", f"posted a Job Description for '{job.title}'")
    conn.close()
    
    return updated_job

@app.put("/api/jobs/{job_id}")
def update_job(job_id: int, job: JobCreate, request: Request):
    username = request.headers.get("x-user-username")
    if not is_user_approved(username):
        raise HTTPException(status_code=403, detail="Access denied. Your account is pending admin approval.")
    role = get_user_role(username)

    conn = sqlite3.connect(STATS_DB, timeout=30.0)
    cur = conn.cursor()
    
    cur.execute("SELECT created_by FROM jobs WHERE id = ?", (job_id,))
    row = cur.fetchone()
    if not row:
        conn.close()
        raise HTTPException(status_code=404, detail="Job not found")
    created_by = row[0]
    
    if not is_admin_or_hr(username):
        if created_by and created_by.lower() != username.lower():
            conn.close()
            raise HTTPException(status_code=403, detail="Forbidden")

    cur.execute("""
        UPDATE jobs SET 
            title = ?, description = ?, client_name = ?, contact_name = ?, client_phone = ?, account_manager = ?,
            assigned_recruiter = ?, target_date = ?, job_type = ?, job_status = ?,
            work_experience = ?, industry = ?, salary = ?, required_skills = ?
        WHERE id = ?
    """, (
        job.title, job.description, job.client_name, job.contact_name, job.client_phone, job.account_manager,
        job.assigned_recruiter, job.target_date, job.job_type, job.job_status,
        job.work_experience, job.industry, job.salary, job.required_skills,
        job_id
    ))
    if cur.rowcount == 0:
        conn.close()
        raise HTTPException(status_code=404, detail="Job not found")
    conn.commit()
    conn.close()
    
    try:
        match_candidates_for_job(job_id)
    except Exception as e:
        print(f"Error matching candidates for job {job_id}: {e}")
        
    conn = sqlite3.connect(STATS_DB, timeout=30.0)
    conn.row_factory = dict_row_factory
    cur = conn.cursor()
    cur.execute("SELECT * FROM jobs WHERE id = ?", (job_id,))
    updated_job = row_to_dict(cur.fetchone())
    
    cur.execute("SELECT status, COUNT(*) as cnt FROM job_candidates WHERE job_id = ? GROUP BY status", (job_id,))
    counts = {r['status']: r['cnt'] for r in cur.fetchall()}
    updated_job['matched_count'] = counts.get('matched', 0)
    updated_job['selected_count'] = counts.get('selected', 0)
    
    cur.execute("SELECT username FROM job_shares WHERE job_id = ?", (job_id,))
    updated_job['shared_with'] = [r['username'] for r in cur.fetchall()]
    
    conn.close()
    
    return updated_job

@app.delete("/api/jobs/{job_id}")
def delete_job(job_id: int, request: Request):
    username = request.headers.get("x-user-username")
    if not is_user_approved(username):
        raise HTTPException(status_code=403, detail="Access denied. Your account is pending admin approval.")
    role = get_user_role(username)

    conn = sqlite3.connect(STATS_DB, timeout=30.0)
    cur = conn.cursor()
    
    cur.execute("SELECT title, created_by FROM jobs WHERE id = ?", (job_id,))
    row = cur.fetchone()
    if not row:
        conn.close()
        raise HTTPException(status_code=404, detail="Job not found")
    job_title, created_by = row
    
    if not is_admin_or_hr(username):
        if created_by and created_by.lower() != username.lower():
            conn.close()
            raise HTTPException(status_code=403, detail="Forbidden")

    cur.execute("DELETE FROM job_candidates WHERE job_id = ?", (job_id,))
    cur.execute("DELETE FROM jobs WHERE id = ?", (job_id,))
    conn.commit()
    conn.close()
    log_activity_db(username or "unknown", f"deleted Job Description '{job_title}'")
    return {"message": "Job deleted"}

@app.get("/api/jobs/{job_id}/candidates")
def get_job_candidates(job_id: int, request: Request):
    username = request.headers.get("x-user-username")
    if not is_user_approved(username):
        raise HTTPException(status_code=403, detail="Access denied. Your account is pending admin approval.")
    
    conn = sqlite3.connect(STATS_DB, timeout=30.0)
    conn.row_factory = dict_row_factory
    cur = conn.cursor()

    # Get job creator
    cur.execute("SELECT created_by FROM jobs WHERE id = ?", (job_id,))
    job_row = cur.fetchone()
    if not job_row:
        conn.close()
        raise HTTPException(status_code=404, detail="Job not found")
    job_creator = job_row['created_by']

    # Get requesting user's roles
    cur.execute("SELECT is_external, is_admin, role, full_name FROM users WHERE LOWER(username) = LOWER(?)", (username,))
    user_row = cur.fetchone()
    is_external = False
    is_user_admin = False
    user_full_name = ""
    if user_row:
        is_external = (user_row['is_external'] == 1)
        is_user_admin = (user_row['is_admin'] == 1 or user_row['role'] == "admin" or is_admin_or_hr(username))
        user_full_name = user_row['full_name']

    # Enforce permission checks:
    if not is_user_admin:
        if is_external:
            # Check share
            cur.execute("SELECT 1 FROM job_shares WHERE job_id = ? AND LOWER(username) = LOWER(?)", (job_id, username))
            if not cur.fetchone():
                conn.close()
                raise HTTPException(status_code=403, detail="Forbidden")
        else:
            # Internal user: must be creator of the job
            if job_creator and job_creator.lower() != username.lower():
                conn.close()
                raise HTTPException(status_code=403, detail="Forbidden")

    cols_to_select = (
        "c.id, c.filename, c.full_name, c.candidate_status, c.total_experience, c.pega_experience, "
        "c.skills, c.certifications, c.ctc, c.notice_period, c.current_organization, c.email, c.phone, "
        "c.linkedin, c.created_by, c.timestamp, c.source, c.cdh_exp, c.expected_ctc, c.percentage_hike, "
        "c.candidate_interview_status, c.availability_in_days, c.current_location, c.pref_locations, "
        "c.current_client, c.domain, c.tier, c.certification_version, "
        "c.sender_email, c.is_qualified, c.is_approved, c.file_url"
    )
    cur.execute(f"""
        SELECT {cols_to_select}, jc.ai_reason, jc.status as job_status
        FROM candidate_metadata c
        JOIN job_candidates jc ON c.id = jc.candidate_id
        WHERE jc.job_id = ?
    """, (job_id,))
    candidates = [dict(row) for row in cur.fetchall()]
    conn.close()
    
    # Replace None values with empty string and remove binary file_bytes
    for row in candidates:
        row.pop("file_bytes", None)
        for k, v in row.items():
            if v is None:
                row[k] = ""

    if is_external:
        # Mask all columns except for ID, name, AI reason, and status for data privacy
        allowed_keys = {'id', 'full_name', 'ai_reason', 'job_status', 'candidate_status'}
        for c in candidates:
            for key in list(c.keys()):
                if key not in allowed_keys:
                    c[key] = ""

    # Mask certifications for non-admin and non-HR users
    is_user_admin_or_hr = is_admin_or_hr(username)
    if not is_user_admin_or_hr:
        for row in candidates:
            row["certifications"] = "[HIDDEN]"

    if not is_admin_or_hr(username):
        keywords = get_masked_keywords()
        candidates = [mask_candidate_record(row, keywords) for row in candidates]
        
    candidates = apply_user_hidden_fields(candidates, username)
    return candidates

@app.get("/api/jobs/{job_id}/unmatched-candidates")
def get_unmatched_candidates(job_id: int, request: Request):
    username = request.headers.get("x-user-username")
    if not is_user_approved(username):
        raise HTTPException(status_code=403, detail="Access denied. Your account is pending admin approval.")
        
    conn = sqlite3.connect(STATS_DB, timeout=30.0)
    conn.row_factory = dict_row_factory
    cur = conn.cursor()

    # Get job creator
    cur.execute("SELECT created_by FROM jobs WHERE id = ?", (job_id,))
    job_row = cur.fetchone()
    if not job_row:
        conn.close()
        raise HTTPException(status_code=404, detail="Job not found")
    job_creator = job_row['created_by']

    # Get requesting user's roles
    cur.execute("SELECT is_external, is_admin, role FROM users WHERE LOWER(username) = LOWER(?)", (username,))
    user_row = cur.fetchone()
    is_external = False
    is_user_admin = False
    if user_row:
        is_external = (user_row['is_external'] == 1)
        is_user_admin = (user_row['is_admin'] == 1 or user_row['role'] == "admin" or is_admin_or_hr(username))

    # Enforce permission checks:
    if not is_user_admin:
        if is_external:
            cur.execute("SELECT 1 FROM job_shares WHERE job_id = ? AND LOWER(username) = LOWER(?)", (job_id, username))
            if not cur.fetchone():
                conn.close()
                raise HTTPException(status_code=403, detail="Forbidden")
        else:
            if job_creator and job_creator.lower() != username.lower():
                conn.close()
                raise HTTPException(status_code=403, detail="Forbidden")

    # Filter unmatched candidates to only those owned by the user if not admin
    if is_user_admin:
        cur.execute("""
            SELECT * FROM candidate_metadata 
            WHERE id NOT IN (
                SELECT candidate_id FROM job_candidates WHERE job_id = ?
            )
            ORDER BY full_name ASC
        """, (job_id,))
    else:
        cur.execute("""
            SELECT * FROM candidate_metadata 
            WHERE LOWER(created_by) = LOWER(?)
            AND id NOT IN (
                SELECT candidate_id FROM job_candidates WHERE job_id = ?
            )
            ORDER BY full_name ASC
        """, (username, job_id))

    candidates = [dict(row) for row in cur.fetchall()]
    conn.close()

    # Replace None values with empty string
    for row in candidates:
        for k, v in row.items():
            if v is None:
                row[k] = ""

    username = request.headers.get("x-user-username")
    # Mask certifications for non-admin and non-HR users
    is_user_admin_or_hr = is_admin_or_hr(username)
    if not is_user_admin_or_hr:
        for row in candidates:
            row["certifications"] = "[HIDDEN]"

    if not is_admin_or_hr(username):
        keywords = get_masked_keywords()
        candidates = [mask_candidate_record(row, keywords) for row in candidates]
        
    candidates = apply_user_hidden_fields(candidates, username)
    return candidates

@app.post("/api/jobs/{job_id}/candidates/{candidate_id}")
def add_job_candidate(job_id: int, candidate_id: int, request: Request):
    username = request.headers.get("x-user-username")
    if not is_user_approved(username):
        raise HTTPException(status_code=403, detail="Access denied. Your account is pending admin approval.")
    role = get_user_role(username)

    conn = sqlite3.connect(STATS_DB, timeout=30.0)
    cur = conn.cursor()
    
    # Verify job creator and existence
    cur.execute("SELECT created_by FROM jobs WHERE id = ?", (job_id,))
    job_row = cur.fetchone()
    if not job_row:
        conn.close()
        raise HTTPException(status_code=404, detail="Job not found")
    job_creator = job_row[0]
    
    if not is_admin_or_hr(username):
        if job_creator and job_creator.lower() != username.lower():
            conn.close()
            raise HTTPException(status_code=403, detail="Forbidden")
            
    # Check if already mapped
    cur.execute("SELECT 1 FROM job_candidates WHERE job_id = ? AND candidate_id = ?", (job_id, candidate_id))
    if cur.fetchone():
        conn.close()
        raise HTTPException(status_code=400, detail="Candidate is already matched to this job.")
        
    cur.execute("""
        INSERT INTO job_candidates (job_id, candidate_id, ai_reason, status)
        VALUES (?, ?, 'Manually associated by recruiter.', 'matched')
    """, (job_id, candidate_id))
    
    # Fetch candidate name for logging
    cur.execute("SELECT full_name FROM candidate_metadata WHERE id = ?", (candidate_id,))
    cand_row = cur.fetchone()
    cand_name = cand_row[0] if cand_row else f"ID {candidate_id}"
    
    # Fetch job title for logging
    cur.execute("SELECT title FROM jobs WHERE id = ?", (job_id,))
    job_row = cur.fetchone()
    job_title = job_row[0] if job_row else f"ID {job_id}"
    
    conn.commit()
    conn.close()
    
    # Try to log activity
    try:
        log_activity_db("recruiter", f"manually matched candidate '{cand_name}' to job '{job_title}'")
    except Exception as e:
        print("Failed to log activity:", e)
        
    return {"message": "Candidate associated with job successfully"}

@app.put("/api/jobs/{job_id}/candidates/{candidate_id}")
def update_job_candidate_status(job_id: int, candidate_id: int, update: JobStatusUpdate, request: Request):
    username = request.headers.get("x-user-username")
    if not is_user_approved(username):
        raise HTTPException(status_code=403, detail="Access denied. Your account is pending admin approval.")
    role = get_user_role(username)

    conn = sqlite3.connect(STATS_DB, timeout=30.0)
    cur = conn.cursor()
    
    # Verify job creator and existence
    cur.execute("SELECT created_by FROM jobs WHERE id = ?", (job_id,))
    job_row = cur.fetchone()
    if not job_row:
        conn.close()
        raise HTTPException(status_code=404, detail="Job not found")
    job_creator = job_row[0]
    
    if not is_admin_or_hr(username):
        if job_creator and job_creator.lower() != username.lower():
            conn.close()
            raise HTTPException(status_code=403, detail="Forbidden")
    
    if update.status is not None:
        if update.status not in ["matched", "selected"]:
            conn.close()
            raise HTTPException(status_code=400, detail="Invalid status")
        cur.execute("UPDATE job_candidates SET status = ? WHERE job_id = ? AND candidate_id = ?", (update.status, job_id, candidate_id))
        
    if update.ai_reason is not None:
        cur.execute("UPDATE job_candidates SET ai_reason = ? WHERE job_id = ? AND candidate_id = ?", (update.ai_reason, job_id, candidate_id))
        
    conn.commit()
    conn.close()
    return {"message": "Status updated"}

@app.delete("/api/jobs/{job_id}/candidates/{candidate_id}")
def delete_job_candidate(job_id: int, candidate_id: int, request: Request):
    username = request.headers.get("x-user-username")
    if not is_user_approved(username):
        raise HTTPException(status_code=403, detail="Access denied. Your account is pending admin approval.")
    role = get_user_role(username)

    conn = sqlite3.connect(STATS_DB, timeout=30.0)
    cur = conn.cursor()
    
    # Verify job creator and existence
    cur.execute("SELECT created_by FROM jobs WHERE id = ?", (job_id,))
    job_row = cur.fetchone()
    if not job_row:
        conn.close()
        raise HTTPException(status_code=404, detail="Job not found")
    job_creator = job_row[0]
    
    if not is_admin_or_hr(username):
        if job_creator and job_creator.lower() != username.lower():
            conn.close()
            raise HTTPException(status_code=403, detail="Forbidden")
    cur.execute("DELETE FROM job_candidates WHERE job_id = ? AND candidate_id = ?", (job_id, candidate_id))
    conn.commit()
    conn.close()
    return {"message": "Candidate removed from job"}

def match_candidates_for_job(job_id: int):
    # This endpoint finds matching candidates for a job and saves them to job_candidates
    conn = sqlite3.connect(STATS_DB, timeout=30.0)
    conn.row_factory = dict_row_factory
    cur = conn.cursor()
    cur.execute("SELECT description, created_by FROM jobs WHERE id = ?", (job_id,))
    job = cur.fetchone()
    if not job:
        conn.close()
        raise HTTPException(status_code=404, detail="Job not found")
    
    jd = job['description']
    job_creator = job['created_by']
    
    # Query candidates directly from the database (only those owned by job creator unless admin or HR)
    if job_creator and not is_admin_or_hr(job_creator):
        cur.execute("SELECT * FROM candidate_metadata WHERE LOWER(created_by) = LOWER(?)", (job_creator,))
    else:
        cur.execute("SELECT * FROM candidate_metadata")
    db_rows = [dict(r) for r in cur.fetchall()]

    if not db_rows:
        conn.close()
        return {"message": "No candidates found in DB"}

    emb, llm = get_models()
    
    # Batch candidates to fit within Groq TPM limit (6000 tokens)
    batch_size = 25
    reason_map = {}
    
    for i in range(0, len(db_rows), batch_size):
        batch_rows = db_rows[i:i+batch_size]
        candidate_lines = []
        for r in batch_rows:
            candidate_lines.append(
                f"ID: {r.get('id')} | "
                f"Name: {r.get('full_name')} | "
                f"Total Experience: {r.get('total_experience')} yrs | "
                f"Pega Experience: {r.get('pega_experience')} yrs | "
                f"CDH Experience: {r.get('cdh_exp')} yrs | "
                f"Skills: {r.get('skills')} | "
                f"Certifications: {r.get('certifications')} | "
                f"Current Location: {r.get('current_location')} | "
                f"Preferred Locations: {r.get('pref_locations')}"
            )
            
        prompt = f"""You are an expert technical recruiter. Evaluate these candidates against the Job Description "pin to pin".

Job Description:
{jd[:2000]}

Candidates to evaluate:
{chr(10).join(candidate_lines)}

Rules for evaluation:
1. Numeric Experience Matching: If a Job Description asks for "X+ years of experience", a candidate matches if their experience is greater than or equal to X.
   - For example: if Job Description requires "1+ years of experience in pega", then candidates with 3.0 years, 4.0 years, or 4.8 years of Pega experience all match perfectly because 3.0 >= 1.0, 4.0 >= 1.0, and 4.8 >= 1.0.
2. Certification Abbreviations:
   - CSSA is equivalent to any of: "PEGA Certified Senior System Architect", "Pega Certified Senior System Architect", "Certified Pega Senior System Architect", "Senior System Architect", or "CSSA".
   - CSA is equivalent to any of: "PEGA Certified System Architect", "Pega Certified System Architect", "Certified Pega System Architect", "System Architect", or "CSA".
   - LSA is equivalent to any of: "PEGA Certified Lead System Architect", "Pega Certified Lead System Architect", "Certified Pega Lead System Architect", "Lead System Architect", or "LSA".
3. Do not invent requirements. If the Job Description only mentions Pega experience, do NOT reject candidates for lacking CSSA or other unrelated certifications.
4. Location Matching: If the Job Description specifies a location requirement, a candidate matches if their Current Location or any of their Preferred Locations match the specified job location (e.g. if the Job Description mentions 'Chennai', a candidate with Current Location or Preferred Location 'Chennai' is a match).

Format your response exactly as a JSON list of objects for matching candidates only:
[
  {{
    "id": <Candidate ID (integer)>,
    "reason": "1-sentence explanation of why they fit based on their specific experience, skills, and location"
  }}
]
Return ONLY the raw JSON block, no markdown, no other text."""

        max_retries = 3
        resp = None
        for attempt in range(max_retries):
            try:
                resp = llm.invoke([HumanMessage(content=prompt)])
                break
            except Exception as api_err:
                if ("429" in str(api_err) or "413" in str(api_err)) and attempt < max_retries - 1:
                    import time
                    time.sleep(2)
                    continue
                print(f"Error evaluating batch in match_candidates_for_job: {api_err}")
                break
                
        if resp is not None:
            try:
                raw = resp.content.strip()
                if "```json" in raw: raw = raw.split("```json")[1].split("```")[0].strip()
                elif "```" in raw: raw = raw.split("```")[1].split("```")[0].strip()
                start, end = raw.find('['), raw.rfind(']')
                if start != -1 and end != -1: raw = raw[start:end+1]
                ai_reasons = json.loads(raw)
                for item in ai_reasons:
                    cid = item.get("id")
                    if cid is not None:
                        reason_map[int(cid)] = item.get("reason", "")
            except Exception as parse_err:
                print(f"Error parsing batch match response: {parse_err}")
                
        # Small delay between batches to prevent rate limit issues
        import time
        time.sleep(0.5)

    # Clear old automatic matches for this job that haven't been selected
    cur.execute("DELETE FROM job_candidates WHERE job_id = ? AND status = 'matched'", (job_id,))

    matches_added = 0
    for r in db_rows:
        cid = r['id']
        reason = reason_map.get(cid)
        
        if reason:
            # Upsert into job_candidates
            cur.execute("""
                INSERT INTO job_candidates (job_id, candidate_id, ai_reason, status) 
                VALUES (?, ?, ?, 'matched')
                ON CONFLICT(job_id, candidate_id) DO UPDATE SET ai_reason = excluded.ai_reason
            """, (job_id, cid, reason))
            
            # Also mark candidate as qualified
            cur.execute("UPDATE candidate_metadata SET is_qualified = 1 WHERE id = ?", (cid,))
            matches_added += 1

    conn.commit()
    conn.close()
    return {"message": f"Successfully matched and added {matches_added} candidates to job"}

@app.post("/api/jobs/{job_id}/match")
def match_candidates_for_job_endpoint(job_id: int, request: Request):
    username = request.headers.get("x-user-username")
    if not is_user_approved(username):
        raise HTTPException(status_code=403, detail="Access denied. Your account is pending admin approval.")
    role = get_user_role(username)
    
    conn = sqlite3.connect(STATS_DB, timeout=30.0)
    cur = conn.cursor()
    cur.execute("SELECT created_by FROM jobs WHERE id = ?", (job_id,))
    row = cur.fetchone()
    if not row:
        conn.close()
        raise HTTPException(status_code=404, detail="Job not found")
    created_by = row[0]
    conn.close()
    
    if not is_admin_or_hr(username):
        if created_by and created_by.lower() != username.lower():
            raise HTTPException(status_code=403, detail="Forbidden")
            
    return match_candidates_for_job(job_id)

# ── Reset ──────────────────────────────────────────────────────────────────────
@app.post("/api/reset")
def reset_all(request: Request):
    conn = sqlite3.connect(STATS_DB, timeout=30.0)
    conn.execute("DELETE FROM candidate_metadata")
    conn.execute("DELETE FROM change_requests")
    conn.execute("DELETE FROM activity_logs")
    conn.commit()
    conn.close()
    try:
        embeddings, _ = get_models()
        db = PGVector(
            connection_string=os.getenv("POSTGRES_DATABASE_URL"),
            embedding_function=embeddings,
            collection_name="resume_embeddings"
        )
        db.delete_collection()
    except Exception:
        pass
    return {"status": "reset complete"}

# ── Auth & Admin Endpoints ───────────────────────────────────────────────────────
@app.post("/api/auth/register")
def register(req: RegisterRequest):
    if not has_digit(req.password):
        raise HTTPException(status_code=400, detail="Password must contain at least one digit")
        
    conn = sqlite3.connect(STATS_DB, timeout=30.0)
    cur = conn.cursor()
    username_exists = False
    email_limit_exceeded = False
    try:
        cur.execute("SELECT 1 FROM users WHERE LOWER(username) = LOWER(?)", (req.username,))
        if cur.fetchone():
            username_exists = True
        else:
            base_email = get_base_email(req.email)
            cur.execute("SELECT email FROM users")
            all_emails = [r[0] for r in cur.fetchall() if r[0]]
            same_email_count = sum(1 for e in all_emails if get_base_email(e) == base_email)
            if same_email_count >= 5:
                email_limit_exceeded = True
            else:
                is_approved_val = 1 if req.username.lower() in ("admin", "user", "boopathi", "praveen", "harish", "sabari") else 0
                cur.execute("INSERT INTO users (full_name, username, password_hash, email, role, is_approved) VALUES (?, ?, ?, ?, 'user', ?)",
                            (req.full_name, req.username, hashlib.sha256(req.password.encode()).hexdigest(), req.email, is_approved_val))
                if is_approved_val == 0:
                    cur.execute("""
                        INSERT INTO change_requests (username, action_type, target_id, payload, description, status)
                        VALUES (?, 'approve_user', ?, NULL, ?, 'pending')
                    """, (req.username, req.username, f"Approve access request for registered user {req.full_name} (@{req.username})"))
                conn.commit()
    except sqlite3.IntegrityError:
        username_exists = True
    except Exception as e:
        conn.close()
        raise HTTPException(status_code=500, detail=str(e))
    conn.close()
    
    if username_exists:
        raise HTTPException(status_code=400, detail="Username already exists")
    if email_limit_exceeded:
        raise HTTPException(status_code=400, detail="Maximum of 5 accounts can be created with the same email address.")
        
    return {"status": "registered", "username": req.username}

@app.post("/api/auth/login")
def login(req: LoginRequest):
    conn = sqlite3.connect(STATS_DB, timeout=30.0)
    conn.row_factory = dict_row_factory
    cur = conn.cursor()
    cur.execute("SELECT * FROM users WHERE LOWER(username) = LOWER(?) AND password_hash = ?", (req.username, hashlib.sha256(req.password.encode()).hexdigest()))
    user = cur.fetchone()
    conn.close()
    if not user:
        raise HTTPException(status_code=401, detail="Invalid username or password")
    user_dict = dict(user)
    if user_dict.get("is_approved", 0) == 0:
        raise HTTPException(status_code=403, detail="Access denied. Your account is pending admin approval.")
    return {
        "username": user_dict["username"],
        "full_name": user_dict["full_name"],
        "role": user_dict["role"],
        "is_hr": user_dict.get("is_hr", 0),
        "is_admin": user_dict.get("is_admin", 0),
        "is_external": user_dict.get("is_external", 0),
        "is_approved": user_dict.get("is_approved", 0),
        "email": user_dict.get("email", ""),
        "hidden_fields": user_dict.get("hidden_fields", "")
    }

class FirebaseSyncRequest(BaseModel):
    email: str
    full_name: str
    username: str
    mobile: Optional[str] = None

@app.post("/api/auth/firebase-sync")
def firebase_sync(req: FirebaseSyncRequest):
    conn = sqlite3.connect(STATS_DB, timeout=30.0)
    conn.row_factory = dict_row_factory
    cur = conn.cursor()
    
    cur.execute("SELECT * FROM users WHERE LOWER(username) = LOWER(?)", (req.username,))
    user = cur.fetchone()
    
    if not user:
        # Check email limit
        base_email = get_base_email(req.email)
        cur.execute("SELECT email FROM users")
        all_emails = [r[0] for r in cur.fetchall() if r[0]]
        same_email_count = sum(1 for e in all_emails if get_base_email(e) == base_email)
        if same_email_count >= 5:
            conn.close()
            raise HTTPException(status_code=400, detail="Maximum of 5 accounts can be created with the same email address.")
            
        try:
            is_approved_val = 1 if req.username.lower() in ("admin", "user", "boopathi", "praveen", "harish", "sabari") else 0
            cur.execute("INSERT INTO users (full_name, username, password_hash, email, mobile, role, is_approved) VALUES (?, ?, ?, ?, ?, 'user', ?)",
                        (req.full_name, req.username, hashlib.sha256("firebase_auth_managed".encode()).hexdigest(), req.email, req.mobile, is_approved_val))
            if is_approved_val == 0:
                cur.execute("""
                    INSERT INTO change_requests (username, action_type, target_id, payload, description, status)
                    VALUES (?, 'approve_user', ?, NULL, ?, 'pending')
                """, (req.username, req.username, f"Approve access request for registered user {req.full_name} (@{req.username})"))
            conn.commit()
            
            cur.execute("SELECT * FROM users WHERE LOWER(username) = LOWER(?)", (req.username,))
            user = cur.fetchone()
            log_activity_db(req.username, "registered an account via Firebase")
        except Exception as e:
            conn.close()
            raise HTTPException(status_code=500, detail=f"Database synchronization error: {str(e)}")
    else:
        # User exists, let's make sure email is synced if it is missing
        try:
            user_dict = dict(user)
            if (not user_dict.get("email") or user_dict.get("email") == "") and req.email:
                cur.execute("UPDATE users SET email = ? WHERE LOWER(username) = LOWER(?)", (req.email, req.username.lower()))
                conn.commit()
                # Fetch again to get updated user dict
                cur.execute("SELECT * FROM users WHERE LOWER(username) = LOWER(?)", (req.username,))
                user = cur.fetchone()
        except Exception as e:
            print(f"Failed to update user email during firebase sync: {e}")
            
    conn.close()
    user_dict = dict(user)
    if user_dict.get("is_approved", 0) == 0:
        raise HTTPException(status_code=403, detail="Access denied. Your account is pending admin approval.")
    return {
        "username": user_dict["username"],
        "full_name": user_dict["full_name"],
        "role": user_dict["role"],
        "is_hr": user_dict.get("is_hr", 0),
        "is_admin": user_dict.get("is_admin", 0),
        "is_external": user_dict.get("is_external", 0),
        "is_approved": user_dict.get("is_approved", 0),
        "email": user_dict.get("email", ""),
        "hidden_fields": user_dict.get("hidden_fields", "")
    }

@app.get("/api/auth/check-exists")
def check_user_exists(username: str, email: str):
    conn = sqlite3.connect(STATS_DB, timeout=30.0)
    cur = conn.cursor()
    
    # 1. Check if username exists
    cur.execute("SELECT 1 FROM users WHERE LOWER(username) = LOWER(?)", (username,))
    if cur.fetchone():
        conn.close()
        return {"exists": True, "reason": "Username already exists."}
        
    # 2. Check email limit (5 accounts)
    base_email = get_base_email(email)
    cur.execute("SELECT email FROM users")
    all_emails = [r[0] for r in cur.fetchall() if r[0]]
    same_email_count = sum(1 for e in all_emails if get_base_email(e) == base_email)
    if same_email_count >= 5:
        conn.close()
        return {"exists": True, "reason": "Maximum of 5 accounts can be created with the same email address."}
        
    conn.close()
    return {"exists": False}

OTP_STORE = {}

@app.get("/api/auth/get-email")
def get_email(username: str):
    conn = sqlite3.connect(STATS_DB, timeout=30.0)
    cur = conn.cursor()
    cur.execute("SELECT email FROM users WHERE LOWER(username) = LOWER(?)", (username,))
    row = cur.fetchone()
    conn.close()
    if row and row[0]:
        return {"email": row[0]}
    # Fallback
    return {"email": f"{username.lower()}@hireai.local"}

@app.get("/api/auth/status")
def get_user_status(request: Request):
    username = request.headers.get("x-user-username")
    if not username:
        raise HTTPException(status_code=401, detail="Not authenticated")
    conn = sqlite3.connect(STATS_DB, timeout=30.0)
    conn.row_factory = dict_row_factory
    cur = conn.cursor()
    cur.execute("SELECT * FROM users WHERE LOWER(username) = LOWER(?)", (username,))
    row = cur.fetchone()
    conn.close()
    if not row:
        raise HTTPException(status_code=404, detail="User not found")
    user_dict = dict(row)
    return {
        "username": user_dict["username"],
        "full_name": user_dict["full_name"],
        "role": user_dict["role"],
        "is_hr": user_dict.get("is_hr", 0),
        "is_admin": user_dict.get("is_admin", 0),
        "is_external": user_dict.get("is_external", 0),
        "is_approved": user_dict.get("is_approved", 0),
        "email": user_dict.get("email", ""),
        "hidden_fields": user_dict.get("hidden_fields", "")
    }

class ForgotPasswordRequest(BaseModel):
    mobile: str

class ResetPasswordRequest(BaseModel):
    mobile: str
    otp: str
    new_password: str

def send_otp_email(to_email: str, otp: str, raise_on_error: bool = False) -> bool:
    import smtplib
    from email.mime.text import MIMEText
    from email.mime.multipart import MIMEMultipart
    
    # Check if we have integrations settings in DB
    smtp_host = None
    smtp_port = None
    smtp_sender = None
    smtp_password = None
    
    try:
        conn = sqlite3.connect(STATS_DB, timeout=30.0)
        cur = conn.cursor()
        cur.execute("""
            SELECT smtp_host, smtp_port, email_user, email_pass, email_enabled 
            FROM integrations_settings LIMIT 1
        """)
        row = cur.fetchone()
        conn.close()
        if row and row[4]:  # If integrations table has a row and email is enabled
            smtp_host = row[0]
            smtp_port = row[1]
            smtp_sender = row[2]
            smtp_password = row[3]
    except Exception as db_err:
        print(f"[SMTP] Error fetching integration settings from DB: {db_err}")
        
    # Fallback to env variables if not set in DB
    if not smtp_host:
        smtp_host = os.getenv("SMTP_HOST", "smtp.gmail.com")
    if not smtp_port:
        try:
            smtp_port = int(os.getenv("SMTP_PORT", "587"))
        except Exception:
            smtp_port = 587
    if not smtp_sender:
        smtp_sender = os.getenv("SMTP_SENDER", "")
    if not smtp_password:
        smtp_password = os.getenv("SMTP_PASSWORD", "")
        
    if not smtp_sender or not smtp_password:
        print(f"[SMTP] Credentials not set. Skipped sending real email.")
        if raise_on_error:
            raise ValueError("SMTP email credentials are not configured. Please set them in the Connect tab or .env file.")
        return False
        
    subject = "Hire AI - Password Reset OTP"
    body = f"Hi,\n\nYour 6-digit OTP code to reset your Hire AI password is: {otp}\n\nThis code is valid for 5 minutes.\n\nBest regards,\nHire AI Team"
    
    try:
        msg = MIMEMultipart()
        msg['From'] = smtp_sender
        msg['To'] = to_email
        msg['Subject'] = subject
        msg.attach(MIMEText(body, 'plain'))
        
        server = smtplib.SMTP(smtp_host, smtp_port, timeout=10.0)
        server.starttls()
        server.login(smtp_sender, smtp_password)
        server.sendmail(smtp_sender, to_email, msg.as_string())
        server.quit()
        print(f"[SMTP] Successfully sent OTP email to {to_email}")
        return True
    except Exception as e:
        print(f"[SMTP] Failed to send email to {to_email}: {str(e)}")
        if raise_on_error:
            raise RuntimeError(f"Failed to send email: {str(e)}")
        return False

@app.post("/api/auth/forgot-password/request")
def request_otp(req: ForgotPasswordRequest):
    import random
    import time
    mobile = req.mobile.strip()
    
    conn = sqlite3.connect(STATS_DB, timeout=30.0)
    cur = conn.cursor()
    cur.execute("SELECT mobile, username FROM users")
    users = cur.fetchall()
    matching_users = [u for u in users if u[0] and str(u[0]).strip() == mobile]
    conn.close()
    
    if not matching_users:
        raise HTTPException(status_code=404, detail="No registered account found with this mobile number.")
        
    username = matching_users[0][1]
    
    otp = f"{random.randint(100000, 999999)}"
    OTP_STORE[mobile] = {
        "otp": otp,
        "expires_at": time.time() + 300.0,
        "username": username
    }
    
    print(f"========================================")
    print(f"[OTP SIMULATION] Password reset OTP for mobile {mobile}: {otp}")
    print(f"========================================")
    
    msg = "Simulated OTP code returned for demonstration."
    res_data = {"message": msg, "otp": otp}
        
    return res_data

@app.post("/api/auth/forgot-password/reset")
def reset_password(req: ResetPasswordRequest):
    import time
    mobile = req.mobile.strip()
    otp_code = req.otp.strip()
    new_pass = req.new_password
    
    if not has_digit(new_pass):
        raise HTTPException(status_code=400, detail="Password must contain at least one digit")
        
    if mobile not in OTP_STORE:
        raise HTTPException(status_code=400, detail="No active password reset request found for this mobile number.")
        
    stored = OTP_STORE[mobile]
    if time.time() > stored["expires_at"]:
        del OTP_STORE[mobile]
        raise HTTPException(status_code=400, detail="OTP has expired. Please request a new one.")
        
    if stored["otp"] != otp_code:
        raise HTTPException(status_code=400, detail="Invalid OTP code. Please try again.")
        
    username = stored["username"]
    del OTP_STORE[mobile]
    
    conn = sqlite3.connect(STATS_DB, timeout=30.0)
    cur = conn.cursor()
    new_hash = hashlib.sha256(new_pass.encode()).hexdigest()
    cur.execute("UPDATE users SET password_hash = ? WHERE username = ?", (new_hash, username))
    conn.commit()
    conn.close()
    
    return {
        "status": "success",
        "message": "Password reset successfully. You can now log in with your new password."
    }

@app.post("/api/candidates")
def add_candidate_manually(request: Request, body: dict):
    username = request.headers.get("x-user-username")
    if not is_user_approved(username):
        raise HTTPException(status_code=403, detail="Access denied. Your account is pending admin approval.")
        
    conn = sqlite3.connect(STATS_DB, timeout=30.0)
    cur = conn.cursor()
    cur.pragma("table_info(candidate_metadata)")
    # wait, cur.execute("PRAGMA table_info(candidate_metadata)")
    cur.execute("PRAGMA table_info(candidate_metadata)")
    valid_cols = [c[1] for c in cur.fetchall()]
    
    insert_data = {}
    for col in valid_cols:
        if col in ('id', 'timestamp'):
            continue
        if col in body:
            insert_data[col] = body[col]
            
    if 'source' not in insert_data or not insert_data['source']:
        insert_data['source'] = 'Manual Entry'
    if 'candidate_status' not in insert_data or not insert_data['candidate_status']:
        insert_data['candidate_status'] = 'New'
    insert_data['created_by'] = username or "admin"
        
    if not insert_data.get('full_name'):
        conn.close()
        raise HTTPException(status_code=400, detail="Candidate Name is required")
        
    cols_str = ", ".join(insert_data.keys())
    placeholders = ", ".join(["?"] * len(insert_data))
    vals = list(insert_data.values())
    
    try:
        cur.execute(f"INSERT INTO candidate_metadata ({cols_str}) VALUES ({placeholders})", vals)
        conn.commit()
        log_activity_db(username or "unknown", f"manually added candidate '{insert_data.get('full_name')}'")
    except Exception as e:
        conn.close()
        raise HTTPException(status_code=500, detail=str(e))
        
    conn.close()
    return {"status": "success", "message": "Candidate added successfully"}

@app.get("/api/admin/requests")
def list_change_requests(request: Request):
    username = request.headers.get("x-user-username")
    role = get_user_role(username)
    if role != "admin":
        raise HTTPException(status_code=403, detail="Forbidden")
        
    conn = sqlite3.connect(STATS_DB, timeout=30.0)
    conn.row_factory = dict_row_factory
    cur = conn.cursor()
    cur.execute("SELECT * FROM change_requests ORDER BY created_at DESC")
    rows = [dict(r) for r in cur.fetchall()]
    conn.close()
    return rows

@app.get("/api/admin/users")
def list_users(request: Request):
    username = request.headers.get("x-user-username")
    if not is_admin_or_hr(username):
        raise HTTPException(status_code=403, detail="Forbidden")
        
        
    conn = sqlite3.connect(STATS_DB, timeout=30.0)
    conn.row_factory = dict_row_factory
    cur = conn.cursor()
    cur.execute("SELECT id, full_name, username, role, is_hr, is_admin, is_external, hidden_fields, is_approved, email FROM users")
    rows = [dict(r) for r in cur.fetchall()]
    conn.close()
    return rows

def is_admin_or_hr(username: str) -> bool:
    if not username:
        return False
    conn = sqlite3.connect(STATS_DB, timeout=30.0)
    cur = conn.cursor()
    cur.execute("SELECT is_admin, is_hr FROM users WHERE LOWER(username) = LOWER(?)", (username,))
    row = cur.fetchone()
    conn.close()
    if row:
        is_admin, is_hr = row
        return is_admin == 1 or is_hr == 1
    return False

def get_user_hidden_fields(username: str) -> list[str]:
    if not username:
        return []
    conn = sqlite3.connect(STATS_DB, timeout=30.0)
    cur = conn.cursor()
    try:
        cur.execute("SELECT hidden_fields FROM users WHERE LOWER(username) = LOWER(?)", (username,))
        row = cur.fetchone()
        if row and row[0]:
            conn.close()
            return [f.strip().lower() for f in row[0].split(",") if f.strip()]
    except Exception as e:
        pass
    conn.close()
    return []

def apply_user_hidden_fields(rows: list[dict], username: str) -> list[dict]:
    hidden = get_user_hidden_fields(username)
    if not hidden:
        return rows
    for r in rows:
        for field in hidden:
            if field in r:
                r[field] = "[HIDDEN]"
    return rows

@app.put("/api/admin/users/{user_id}/permissions")
def update_user_permissions_endpoint(user_id: int, body: UserPermissionsUpdate, request: Request):
    username = request.headers.get("x-user-username")
    role = get_user_role(username)
    if role != "admin":
        raise HTTPException(status_code=403, detail="Forbidden")
        
    conn = sqlite3.connect(STATS_DB, timeout=30.0)
    cur = conn.cursor()
    
    is_external_val = body.is_external if body.is_external is not None else 0
    is_hr_val = body.is_hr
    is_admin_val = body.is_admin
    hidden_fields_val = body.hidden_fields if body.hidden_fields is not None else ""
    
    if is_external_val == 1:
        is_hr_val = 0
        is_admin_val = 0
        
    new_role = "admin" if is_admin_val == 1 else "user"
    if body.is_approved is not None:
        cur.execute("UPDATE users SET is_hr = ?, is_admin = ?, is_external = ?, role = ?, hidden_fields = ?, is_approved = ? WHERE id = ?", 
                    (is_hr_val, is_admin_val, is_external_val, new_role, hidden_fields_val, body.is_approved, user_id))
    else:
        cur.execute("UPDATE users SET is_hr = ?, is_admin = ?, is_external = ?, role = ?, hidden_fields = ? WHERE id = ?", 
                    (is_hr_val, is_admin_val, is_external_val, new_role, hidden_fields_val, user_id))
    conn.commit()
    conn.close()
    return {"status": "updated"}

class JobShareRequest(BaseModel):
    usernames: list[str]

@app.post("/api/jobs/{job_id}/share")
def share_job(job_id: int, req: JobShareRequest, request: Request):
    username = request.headers.get("x-user-username")
    if not is_user_approved(username):
        raise HTTPException(status_code=403, detail="Access denied. Your account is pending admin approval.")
    if not is_admin_or_hr(username):
        raise HTTPException(status_code=403, detail="Forbidden")
        
    conn = sqlite3.connect(STATS_DB, timeout=30.0)
    cur = conn.cursor()
    
    # Verify job creator and existence
    cur.execute("SELECT created_by FROM jobs WHERE id = ?", (job_id,))
    job_row = cur.fetchone()
    if not job_row:
        conn.close()
        raise HTTPException(status_code=404, detail="Job not found")
        
    created_by = job_row[0]
    role = get_user_role(username)
    if role != "admin":
        if created_by and created_by.lower() != username.lower():
            conn.close()
            raise HTTPException(status_code=403, detail="Forbidden")
    
    # Verify job exists
    cur.execute("SELECT 1 FROM jobs WHERE id = ?", (job_id,))
    if not cur.fetchone():
        conn.close()
        raise HTTPException(status_code=404, detail="Job not found")
        
    # Clear existing shares
    cur.execute("DELETE FROM job_shares WHERE job_id = ?", (job_id,))
    
    # Insert new shares
    for u in req.usernames:
        cur.execute("INSERT INTO job_shares (job_id, username) VALUES (?, ?)", (job_id, u))
        
    conn.commit()
    conn.close()
    
    # Log activity
    job_title = get_job_title(job_id)
    log_activity_db(username or "unknown", f"shared Job Description '{job_title}' with {len(req.usernames)} external users")
    
    return {"status": "shared"}

@app.get("/api/jobs/{job_id}/shares")
def get_job_shares(job_id: int, request: Request):
    username = request.headers.get("x-user-username")
    if not is_user_approved(username):
        raise HTTPException(status_code=403, detail="Access denied. Your account is pending admin approval.")
    if not is_admin_or_hr(username):
        raise HTTPException(status_code=403, detail="Forbidden")
        
    conn = sqlite3.connect(STATS_DB, timeout=30.0)
    cur = conn.cursor()
    
    # Verify job creator and existence
    cur.execute("SELECT created_by FROM jobs WHERE id = ?", (job_id,))
    job_row = cur.fetchone()
    if not job_row:
        conn.close()
        raise HTTPException(status_code=404, detail="Job not found")
        
    created_by = job_row[0]
    role = get_user_role(username)
    if role != "admin":
        if created_by and created_by.lower() != username.lower():
            conn.close()
            raise HTTPException(status_code=403, detail="Forbidden")
    cur.execute("SELECT username FROM job_shares WHERE job_id = ?", (job_id,))
    rows = cur.fetchall()
    conn.close()
    
    return [r[0] for r in rows]

class MaskedKeywordCreate(BaseModel):
    keyword: str

@app.get("/api/admin/masked-keywords")
def get_admin_masked_keywords(request: Request):
    username = request.headers.get("x-user-username")
    role = get_user_role(username)
    if role != "admin":
        raise HTTPException(status_code=403, detail="Forbidden")
    return get_masked_keywords()

@app.post("/api/admin/masked-keywords")
def add_admin_masked_keyword(req: MaskedKeywordCreate, request: Request):
    username = request.headers.get("x-user-username")
    role = get_user_role(username)
    if role != "admin":
        raise HTTPException(status_code=403, detail="Forbidden")
    kw = req.keyword.strip()
    if not kw:
        raise HTTPException(status_code=400, detail="Keyword cannot be empty")
    conn = sqlite3.connect(STATS_DB, timeout=30.0)
    cur = conn.cursor()
    try:
        cur.execute("INSERT OR IGNORE INTO masked_keywords (keyword) VALUES (?)", (kw,))
        conn.commit()
    except Exception as e:
        conn.close()
        raise HTTPException(status_code=500, detail=str(e))
    conn.close()
    log_activity_db(username, f"added masked keyword '{kw}'")
    return {"status": "added", "keyword": kw}

@app.delete("/api/admin/masked-keywords/{keyword}")
def delete_admin_masked_keyword(keyword: str, request: Request):
    username = request.headers.get("x-user-username")
    role = get_user_role(username)
    if role != "admin":
        raise HTTPException(status_code=403, detail="Forbidden")
    kw = keyword.strip()
    conn = sqlite3.connect(STATS_DB, timeout=30.0)
    cur = conn.cursor()
    try:
        cur.execute("DELETE FROM masked_keywords WHERE keyword = ?", (kw,))
        conn.commit()
    except Exception as e:
        conn.close()
        raise HTTPException(status_code=500, detail=str(e))
    conn.close()
    log_activity_db(username, f"deleted masked keyword '{kw}'")
    return {"status": "deleted", "keyword": kw}

@app.delete("/api/admin/users/{user_id}")
def delete_user_endpoint(user_id: int, request: Request):
    username = request.headers.get("x-user-username")
    role = get_user_role(username)
    if role != "admin":
        raise HTTPException(status_code=403, detail="Forbidden")
        
    conn = sqlite3.connect(STATS_DB, timeout=30.0)
    cur = conn.cursor()
    # Prevent self-deletion and get email/fullname
    cur.execute("SELECT username, email, full_name FROM users WHERE id = ?", (user_id,))
    row = cur.fetchone()
    if not row:
        conn.close()
        raise HTTPException(status_code=404, detail="User not found")
        
    deleted_username = row[0]
    deleted_email = row[1]
    deleted_fullname = row[2]
    
    if deleted_username.lower() == username.lower():
        conn.close()
        raise HTTPException(status_code=400, detail="You cannot delete yourself.")
        
    # We will delete from users table at the end to prevent foreign key violations
    
    # Delete from change_requests (where requester or target is this user)
    cur.execute("DELETE FROM change_requests WHERE LOWER(username) = LOWER(?) OR LOWER(target_id) = LOWER(?)", (deleted_username, deleted_username))
    
    # Delete from job_shares
    cur.execute("DELETE FROM job_shares WHERE LOWER(username) = LOWER(?)", (deleted_username,))
    
    # Delete from team_members (by full name and by username, if exists)
    cur.execute("DELETE FROM team_members WHERE LOWER(name) = LOWER(?) OR LOWER(name) = LOWER(?)", (deleted_fullname, deleted_username))
    
    # Delete from activity_logs
    cur.execute("DELETE FROM activity_logs WHERE LOWER(username) = LOWER(?)", (deleted_username,))
    
    # Delete candidate records and their resume files owned/created by this user
    cur.execute("SELECT id, filename FROM candidate_metadata WHERE LOWER(created_by) = LOWER(?)", (deleted_username,))
    candidates = cur.fetchall()
    if candidates:
        candidate_ids = [c[0] for c in candidates]
        placeholders = ",".join("?" for _ in candidate_ids)
        
        # Delete resume files from disk
        for c_id, fname in candidates:
            if fname:
                fpath = os.path.join(UPLOAD_DIR, fname)
                if os.path.exists(fpath):
                    try:
                        os.remove(fpath)
                    except Exception as e:
                        print(f"Error removing resume file {fpath}: {e}")
                        
        # Delete from DB
        cur.execute(f"DELETE FROM job_candidates WHERE candidate_id IN ({placeholders})", candidate_ids)
        cur.execute(f"DELETE FROM candidate_metadata WHERE id IN ({placeholders})", candidate_ids)
        
    # Delete job records created by this user
    cur.execute("SELECT id FROM jobs WHERE LOWER(created_by) = LOWER(?)", (deleted_username,))
    job_ids = [r[0] for r in cur.fetchall()]
    if job_ids:
        placeholders = ",".join("?" for _ in job_ids)
        cur.execute(f"DELETE FROM job_candidates WHERE job_id IN ({placeholders})", job_ids)
        cur.execute(f"DELETE FROM job_shares WHERE job_id IN ({placeholders})", job_ids)
        cur.execute(f"DELETE FROM jobs WHERE id IN ({placeholders})", job_ids)
        
    # Delete from users table at the end
    cur.execute("DELETE FROM users WHERE id = ?", (user_id,))
    
    conn.commit()
    conn.close()
    
    log_activity_db(username, f"completely deleted user '{deleted_username}' from system")
    return {"status": "deleted"}

@app.post("/api/admin/requests/{request_id}/approve")
def approve_change_request(request_id: int, request: Request, background_tasks: BackgroundTasks):
    username = request.headers.get("x-user-username")
    role = get_user_role(username)
    if role != "admin":
        raise HTTPException(status_code=403, detail="Forbidden")
        
    conn = sqlite3.connect(STATS_DB, timeout=30.0)
    conn.row_factory = dict_row_factory
    cur = conn.cursor()
    cur.execute("SELECT * FROM change_requests WHERE id = ?", (request_id,))
    req_row = cur.fetchone()
    if not req_row:
        conn.close()
        raise HTTPException(status_code=404, detail="Request not found")
        
    req_data = dict(req_row)
    if req_data["status"] != "pending":
        conn.close()
        raise HTTPException(status_code=400, detail="Request is already resolved")
        
    action_type = req_data["action_type"]
    target_id = req_data["target_id"]
    payload = req_data["payload"]
    
    try:
        if action_type == "add_column":
            col_data = json.loads(payload)
            clean_key = re.sub(r'[^a-zA-Z0-9_]', '', col_data["col_key"].replace(' ', '_')).lower()
            cur.execute("PRAGMA table_info(candidate_metadata)")
            existing = [c[1] for c in cur.fetchall()]
            if clean_key not in existing:
                cur.execute(f"ALTER TABLE candidate_metadata ADD COLUMN {clean_key} TEXT")
                cur.execute("INSERT OR IGNORE INTO custom_columns (col_key, col_label, description) VALUES (?, ?, ?)", 
                            (clean_key, col_data["col_label"], col_data["description"]))
                
        elif action_type == "delete_column":
            col_key = target_id
            cur.execute("DELETE FROM custom_columns WHERE col_key=?", (col_key,))
            try:
                cur.execute(f"ALTER TABLE candidate_metadata DROP COLUMN {col_key}")
            except Exception:
                pass
                
        elif action_type == "update_candidate":
            candidate_id = int(target_id)
            updates = json.loads(payload)
            set_clause = ", ".join(f"{k}=?" for k in updates)
            cur.execute(
                f"UPDATE candidate_metadata SET {set_clause} WHERE id=?",
                list(updates.values()) + [candidate_id]
            )
            match_related_fields = {
                'full_name', 'total_experience', 'pega_experience', 'cdh_exp',
                'skills', 'certifications', 'current_location', 'pref_locations'
            }
            if any(field in updates for field in match_related_fields):
                background_tasks.add_task(match_candidate_to_all_jobs, candidate_id)
                
        elif action_type == "delete_candidate":
            candidate_id = int(target_id)
            cur.execute("DELETE FROM job_candidates WHERE candidate_id=?", (candidate_id,))
            cur.execute("DELETE FROM candidate_metadata WHERE id=?", (candidate_id,))
            
        elif action_type == "approve_resume":
            candidate_id = int(target_id)
            cur.execute("UPDATE candidate_metadata SET is_approved = 1 WHERE id = ?", (candidate_id,))
            background_tasks.add_task(match_candidate_to_all_jobs, candidate_id)
            
        elif action_type == "create_job":
            job_data = json.loads(payload)
            cur.execute("INSERT INTO jobs (title, description) VALUES (?, ?)", (job_data["title"], job_data["description"]))
            job_id = cur.lastrowid
            background_tasks.add_task(match_candidates_for_job, job_id)
            
        elif action_type == "update_job":
            job_id = int(target_id)
            job_data = json.loads(payload)
            cur.execute("UPDATE jobs SET title = ?, description = ? WHERE id = ?", (job_data["title"], job_data["description"], job_id))
            background_tasks.add_task(match_candidates_for_job, job_id)
            
        elif action_type == "delete_job":
            job_id = int(target_id)
            cur.execute("DELETE FROM job_candidates WHERE job_id = ?", (job_id,))
            cur.execute("DELETE FROM jobs WHERE id = ?", (job_id,))
            
        elif action_type == "update_job_candidate":
            job_id, candidate_id = map(int, target_id.split(":"))
            update_data = json.loads(payload)
            if "status" in update_data and update_data["status"] is not None:
                cur.execute("UPDATE job_candidates SET status = ? WHERE job_id = ? AND candidate_id = ?", 
                            (update_data["status"], job_id, candidate_id))
            if "ai_reason" in update_data and update_data["ai_reason"] is not None:
                cur.execute("UPDATE job_candidates SET ai_reason = ? WHERE job_id = ? AND candidate_id = ?", 
                            (update_data["ai_reason"], job_id, candidate_id))
                                
        elif action_type == "delete_job_candidate":
            job_id, candidate_id = map(int, target_id.split(":"))
            cur.execute("DELETE FROM job_candidates WHERE job_id = ? AND candidate_id = ?", (job_id, candidate_id))
            
        elif action_type == "match_job":
            job_id = int(target_id)
            background_tasks.add_task(match_candidates_for_job, job_id)
            
        elif action_type == "reset_all":
            cur.execute("DELETE FROM candidate_metadata")
            try:
                embeddings, _ = get_models()
                db = PGVector(
                    connection_string=os.getenv("POSTGRES_DATABASE_URL"),
                    embedding_function=embeddings,
                    collection_name="resume_embeddings"
                )
                db.delete_collection()
            except Exception:
                pass
                
        elif action_type == "approve_user":
            target_username = target_id
            cur.execute("UPDATE users SET is_approved = 1 WHERE LOWER(username) = LOWER(?)", (target_username,))
                
        cur.execute("UPDATE change_requests SET status = 'approved' WHERE id = ?", (request_id,))
        conn.commit()
    except Exception as e:
        conn.rollback()
        conn.close()
        raise HTTPException(status_code=500, detail=f"Failed to execute request: {str(e)}")
        
    conn.close()
    return {"status": "approved"}

@app.post("/api/admin/requests/{request_id}/reject")
def reject_change_request(request_id: int, request: Request):
    username = request.headers.get("x-user-username")
    role = get_user_role(username)
    if role != "admin":
        raise HTTPException(status_code=403, detail="Forbidden")
        
    conn = sqlite3.connect(STATS_DB, timeout=30.0)
    cur = conn.cursor()
    cur.execute("SELECT status, action_type, target_id FROM change_requests WHERE id = ?", (request_id,))
    row = cur.fetchone()
    if not row:
        conn.close()
        raise HTTPException(status_code=404, detail="Request not found")
    status, action_type, target_id = row
    if status != "pending":
        conn.close()
        raise HTTPException(status_code=400, detail="Request is already resolved")
        
    if action_type == "approve_user":
        cur.execute("DELETE FROM users WHERE LOWER(username) = LOWER(?)", (target_id,))
        
    cur.execute("UPDATE change_requests SET status = 'rejected' WHERE id = ?", (request_id,))
    conn.commit()
    conn.close()
    return {"status": "rejected"}

# THEME PRESETS FOR CONFIRMATION REPLY EMAILS
THEME_TEMPLATES = {
    'professional': {
        'subject': 'Re: {subject} (Ref: {ref})',
        'body_missing': (
            "Dear {candidate_name},\n\n"
            "Thank you for your interest in Alamaticz Solutions and for submitting your application.\n\n"
            "We appreciate the time you have taken to apply for this opportunity. To help us evaluate your profile further, "
            "kindly share the following details:\n\n"
            "{missing_fields}\n\n"
            "Once we receive the above information, our recruitment team will review your profile and get back to you regarding the next steps in the selection process.\n\n"
            "We look forward to hearing from you.\n\n"
            "Best regards,\n\n"
            "HR Team\n"
            "Alamaticz Solutions"
        ),
        'body_complete': (
            "Dear {candidate_name},\n\n"
            "Thank you for your interest in Alamaticz Solutions and for submitting your application.\n\n"
            "We appreciate the time you have taken to apply for this opportunity. Our recruitment team will review your profile and get back to you regarding the next steps in the selection process.\n\n"
            "Best regards,\n\n"
            "HR Team\n"
            "Alamaticz Solutions"
        )
    },
    'creative': {
        'subject': 'Excited to connect! Re: {subject} (Ref: {ref})',
        'body_missing': (
            "Hi {candidate_name}!\n\n"
            "Thanks for reaching out and sharing your resume with us. We love connecting with talented people!\n\n"
            "We are eager to dive into your application, but we are missing a few details. Could you please share the following with us?\n\n"
            "{missing_fields}\n\n"
            "As soon as we get these details, we'll review your profile and get back to you regarding the next steps.\n\n"
            "Can't wait to hear back from you!\n\n"
            "Cheers,\n\n"
            "The Talent Team\n"
            "Alamaticz Solutions"
        ),
        'body_complete': (
            "Hi {candidate_name}!\n\n"
            "Thanks for reaching out and sharing your application with us!\n\n"
            "We have everything we need. Our team is already looking over your profile, and we'll be in touch soon with the next steps.\n\n"
            "Have a fantastic day!\n\n"
            "Cheers,\n\n"
            "The Talent Team\n"
            "Alamaticz Solutions"
        )
    },
    'warm': {
        'subject': 'Thank you for applying! Re: {subject} (Ref: {ref})',
        'body_missing': (
            "Hello {candidate_name},\n\n"
            "We hope you are having a wonderful day! Thank you so much for taking the time to apply to our team.\n\n"
            "To help us get a better picture of your experience and fit for the role, could you please help us with these remaining details?\n\n"
            "{missing_fields}\n\n"
            "We really appreciate your support and look forward to reviewing your application as soon as we receive this.\n\n"
            "Wishing you all the best,\n\n"
            "Your Friends at HR\n"
            "Alamaticz Solutions"
        ),
        'body_complete': (
            "Hello {candidate_name},\n\n"
            "We hope you are doing well! Thank you so much for sending over your application.\n\n"
            "This is just a quick note to let you know we've received all your information. Our team will review everything carefully and get back to you soon.\n\n"
            "Take care,\n\n"
            "Your Friends at HR\n"
            "Alamaticz Solutions"
        )
    }
}

# ── Integrations Settings Endpoints ───────────────────────────────────────────
class IntegrationSettingsRequest(BaseModel):
    email_enabled: int
    imap_host: str
    imap_port: int
    smtp_host: str
    smtp_port: int
    email_user: str
    email_pass: str
    keywords: str
    drive_enabled: int
    reply_theme: Optional[str] = "professional"
    reply_subject: Optional[str] = "Re: {subject} (Ref: {ref})"
    reply_body_missing: Optional[str] = None
    reply_body_complete: Optional[str] = None
    gdrive_client_id: Optional[str] = None
    gdrive_client_secret: Optional[str] = None
    gdrive_refresh_token: Optional[str] = None
    gdrive_folder_id: Optional[str] = None
    gdrive_email: Optional[str] = None
    ms_client_id: Optional[str] = None
    ms_client_secret: Optional[str] = None
    ms_tenant_id: Optional[str] = "common"
    gmail_enabled: Optional[int] = 0
    gmail_email: Optional[str] = None
    gmail_pass: Optional[str] = None
    outlook_enabled: Optional[int] = 0
    outlook_email: Optional[str] = None
    additional_emails: Optional[str] = "[]"
    theme_usage_counts: Optional[str] = "{}"
    default_resume_template: Optional[str] = "alamaticz"

@app.get("/api/integrations")
def get_integrations_settings(request: Request):
    username = request.headers.get("x-user-username")
    role = get_user_role(username)
    if role != "admin":
        raise HTTPException(status_code=403, detail="Forbidden")
        
    conn = sqlite3.connect(STATS_DB, timeout=30.0)
    cur = conn.cursor()
    cur.execute("""
        SELECT email_enabled, imap_host, imap_port, smtp_host, smtp_port, 
               email_user, email_pass, keywords, drive_enabled,
               reply_theme, reply_subject, reply_body_missing, reply_body_complete,
               gdrive_client_id, gdrive_client_secret, gdrive_refresh_token, gdrive_folder_id, gdrive_email,
               ms_client_id, ms_client_secret, ms_tenant_id,
               additional_emails, theme_usage_counts,
               gmail_enabled, gmail_email, gmail_pass, outlook_enabled, outlook_email, default_resume_template
        FROM integrations_settings LIMIT 1
    """)
    row = cur.fetchone()
    conn.close()
    
    if not row:
        return {
            "email_enabled": 0, "imap_host": "imap.gmail.com", "imap_port": 993,
            "smtp_host": "smtp.gmail.com", "smtp_port": 587, "email_user": "",
            "email_pass": "", "keywords": "resume,alamaticz,solution,job", "drive_enabled": 0,
            "reply_theme": "professional",
            "reply_subject": "Re: {subject} (Ref: {ref})",
            "reply_body_missing": THEME_TEMPLATES["professional"]["body_missing"],
            "reply_body_complete": THEME_TEMPLATES["professional"]["body_complete"],
            "gdrive_client_id": "",
            "gdrive_client_secret": "",
            "gdrive_refresh_token": "",
            "gdrive_folder_id": "",
            "gdrive_email": "",
            "ms_client_id": "",
            "ms_client_secret": "",
            "ms_tenant_id": "common",
            "gmail_enabled": 0,
            "gmail_email": "",
            "gmail_pass": "",
            "outlook_enabled": 0,
            "outlook_email": "",
            "default_resume_template": "alamaticz",
            "additional_emails": "[]",
            "theme_usage_counts": "{}"
        }
        
    masked_pass = ""
    if row[6]:
        masked_pass = "****"
        
    reply_theme = row[9] or "professional"
    reply_subject = row[10]
    reply_body_missing = row[11]
    reply_body_complete = row[12]
    
    # Load default templates if preset is not custom
    if reply_theme in THEME_TEMPLATES:
        if reply_theme != 'custom':
            reply_subject = THEME_TEMPLATES[reply_theme]['subject']
            reply_body_missing = THEME_TEMPLATES[reply_theme]['body_missing']
            reply_body_complete = THEME_TEMPLATES[reply_theme]['body_complete']
        else:
            if not reply_subject: reply_subject = THEME_TEMPLATES['professional']['subject']
            if not reply_body_missing: reply_body_missing = THEME_TEMPLATES['professional']['body_missing']
            if not reply_body_complete: reply_body_complete = THEME_TEMPLATES['professional']['body_complete']
    else:
        if not reply_subject: reply_subject = THEME_TEMPLATES['professional']['subject']
        if not reply_body_missing: reply_body_missing = THEME_TEMPLATES['professional']['body_missing']
        if not reply_body_complete: reply_body_complete = THEME_TEMPLATES['professional']['body_complete']
        
    masked_gdrive_secret = ""
    if row[14]:
        masked_gdrive_secret = "****"
        
    masked_gdrive_refresh = ""
    if row[15]:
        masked_gdrive_refresh = "****"
        
    masked_ms_secret = ""
    if row[19]:
        masked_ms_secret = "****"
        
    ms_tenant_id = row[20] or "common"
        
    additional_emails_raw = row[21] if len(row) > 21 else "[]"
    masked_additional_emails = "[]"
    if additional_emails_raw:
        try:
            add_emails = json.loads(additional_emails_raw)
            for item in add_emails:
                if item.get("email_pass"):
                    item["email_pass"] = "****"
            masked_additional_emails = json.dumps(add_emails)
        except Exception:
            pass

    theme_usage_counts = row[22] if len(row) > 22 else "{}"

    gmail_enabled = row[23] if len(row) > 23 else 0
    gmail_email = row[24] if len(row) > 24 else ""
    gmail_pass_masked = "****" if len(row) > 25 and row[25] else ""
    outlook_enabled = row[26] if len(row) > 26 else 0
    outlook_email = row[27] if len(row) > 27 else ""

    return {
        "email_enabled": row[0],
        "imap_host": row[1] or "imap.gmail.com",
        "imap_port": row[2] or 993,
        "smtp_host": row[3] or "smtp.gmail.com",
        "smtp_port": row[4] or 587,
        "email_user": row[5] or "",
        "email_pass": masked_pass,
        "keywords": row[7] or "resume,alamaticz,solution,job",
        "drive_enabled": row[8],
        "reply_theme": reply_theme,
        "reply_subject": reply_subject,
        "reply_body_missing": reply_body_missing,
        "reply_body_complete": reply_body_complete,
        "gdrive_client_id": row[13] or "",
        "gdrive_client_secret": masked_gdrive_secret,
        "gdrive_refresh_token": masked_gdrive_refresh,
        "gdrive_folder_id": row[16] or "",
        "gdrive_email": row[17] or "",
        "ms_client_id": row[18] or "",
        "ms_client_secret": masked_ms_secret,
        "ms_tenant_id": ms_tenant_id,
        "gmail_enabled": gmail_enabled,
        "gmail_email": gmail_email,
        "gmail_pass": gmail_pass_masked,
        "outlook_enabled": outlook_enabled,
        "outlook_email": outlook_email,
        "additional_emails": masked_additional_emails,
        "theme_usage_counts": theme_usage_counts,
        "default_resume_template": row[28] if len(row) > 28 and row[28] else "alamaticz"
    }

@app.post("/api/integrations")
def save_integrations_settings(settings: IntegrationSettingsRequest, request: Request):
    username = request.headers.get("x-user-username")
    role = get_user_role(username)
    if role != "admin":
        raise HTTPException(status_code=403, detail="Forbidden")
        
    conn = sqlite3.connect(STATS_DB, timeout=30.0)
    cur = conn.cursor()
    
    cur.execute("SELECT id, email_pass, gdrive_client_secret, gdrive_refresh_token, ms_client_secret, ms_tenant_id, additional_emails, theme_usage_counts, gmail_pass FROM integrations_settings LIMIT 1")
    row = cur.fetchone()
    
    final_pass = settings.email_pass
    if final_pass == "****" and row:
        final_pass = row[1]
        
    final_gmail_pass = settings.gmail_pass
    if final_gmail_pass == "****" and row and len(row) > 8:
        final_gmail_pass = row[8]
        
    final_gdrive_secret = settings.gdrive_client_secret
    if final_gdrive_secret == "****" and row:
        final_gdrive_secret = row[2]
        
    final_gdrive_refresh = settings.gdrive_refresh_token
    if final_gdrive_refresh == "****" and row:
        final_gdrive_refresh = row[3]
        
    final_ms_secret = settings.ms_client_secret
    if final_ms_secret == "****" and row:
        final_ms_secret = row[4]
        
    final_ms_tenant = settings.ms_tenant_id
        
    # Handle additional_emails passwords restoration
    try:
        new_list = json.loads(settings.additional_emails or "[]")
    except Exception:
        new_list = []

    old_passwords = {}
    if row and len(row) > 6 and row[6]:
        try:
            old_list = json.loads(row[6])
            for item in old_list:
                if "email_user" in item and "email_pass" in item:
                    old_passwords[item["email_user"].lower()] = item["email_pass"]
        except Exception:
            pass

    for item in new_list:
        if item.get("email_pass") == "****" and item.get("email_user"):
            item["email_pass"] = old_passwords.get(item["email_user"].lower(), "")
            
    final_additional_emails = json.dumps(new_list)

    # Load and increment theme usage count
    current_counts = {}
    if row and len(row) > 7 and row[7]:
        try:
            current_counts = json.loads(row[7])
        except Exception:
            current_counts = {}
    selected_theme = settings.reply_theme or "professional"
    current_counts[selected_theme] = current_counts.get(selected_theme, 0) + 1
    final_theme_usage = json.dumps(current_counts)

    if row:
        cur.execute("""
        UPDATE integrations_settings SET
            email_enabled = ?, imap_host = ?, imap_port = ?,
            smtp_host = ?, smtp_port = ?, email_user = ?,
            email_pass = ?, keywords = ?, drive_enabled = ?,
            reply_theme = ?, reply_subject = ?,
            reply_body_missing = ?, reply_body_complete = ?,
            gdrive_client_id = ?, gdrive_client_secret = ?,
            gdrive_refresh_token = ?, gdrive_folder_id = ?,
            gdrive_email = ?, ms_client_id = ?, ms_client_secret = ?, 
            ms_tenant_id = ?, additional_emails = ?,
            theme_usage_counts = ?,
            gmail_enabled = ?, gmail_email = ?, gmail_pass = ?,
            outlook_enabled = ?, outlook_email = ?,
            default_resume_template = ?
        WHERE id = ?
        """, (settings.email_enabled, settings.imap_host, settings.imap_port,
              settings.smtp_host, settings.smtp_port, settings.email_user,
              final_pass, settings.keywords, settings.drive_enabled,
              settings.reply_theme, settings.reply_subject,
              settings.reply_body_missing, settings.reply_body_complete,
              settings.gdrive_client_id, final_gdrive_secret,
              final_gdrive_refresh, settings.gdrive_folder_id,
              settings.gdrive_email, settings.ms_client_id, final_ms_secret, 
              final_ms_tenant, final_additional_emails, final_theme_usage,
              settings.gmail_enabled, settings.gmail_email, final_gmail_pass,
              settings.outlook_enabled, settings.outlook_email, settings.default_resume_template, row[0]))
    else:
        cur.execute("""
        INSERT INTO integrations_settings (
            email_enabled, imap_host, imap_port, smtp_host, smtp_port, email_user, email_pass, keywords, drive_enabled,
            reply_theme, reply_subject, reply_body_missing, reply_body_complete,
            gdrive_client_id, gdrive_client_secret, gdrive_refresh_token, gdrive_folder_id, gdrive_email, 
            ms_client_id, ms_client_secret, ms_tenant_id, additional_emails,
            theme_usage_counts, gmail_enabled, gmail_email, gmail_pass, outlook_enabled, outlook_email, default_resume_template
        ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
        """, (settings.email_enabled, settings.imap_host, settings.imap_port,
              settings.smtp_host, settings.smtp_port, settings.email_user,
              final_pass, settings.keywords, settings.drive_enabled,
              settings.reply_theme, settings.reply_subject,
              settings.reply_body_missing, settings.reply_body_complete,
              settings.gdrive_client_id, final_gdrive_secret,
              settings.gdrive_refresh_token, settings.gdrive_folder_id,
              settings.gdrive_email, settings.ms_client_id, final_ms_secret, 
              final_ms_tenant, final_additional_emails, final_theme_usage,
              settings.gmail_enabled, settings.gmail_email, final_gmail_pass,
              settings.outlook_enabled, settings.outlook_email, settings.default_resume_template))

    if settings.email_enabled == 0:
        try:
            cur.execute("SELECT id, filename FROM candidate_metadata WHERE source = 'uploaded from mail'")
            email_candidates = cur.fetchall()
            for cand_id, filename in email_candidates:
                if filename:
                    file_path = os.path.join(UPLOAD_DIR, filename)
                    if os.path.exists(file_path):
                        try:
                            os.remove(file_path)
                            print(f"Deleted email candidate file: {file_path}")
                        except Exception as file_err:
                            print(f"Error deleting file {file_path}: {file_err}")
                
                cur.execute("DELETE FROM job_candidates WHERE candidate_id = ?", (cand_id,))
                cur.execute("DELETE FROM candidate_metadata WHERE id = ?", (cand_id,))
            print(f"Deleted {len(email_candidates)} email-imported candidates because email sync was disabled.")
        except Exception as delete_err:
            print(f"Error purging email candidates: {delete_err}")

    conn.commit()
    conn.close()
    
    log_activity_db(username, "updated integrations settings")
    return {"status": "saved"}

class GDriveExchangeRequest(BaseModel):
    client_id: str
    client_secret: str
    code: str

@app.post("/api/integrations/gdrive/exchange")
def exchange_gdrive_code(payload: GDriveExchangeRequest, request: Request):
    username = request.headers.get("x-user-username")
    role = get_user_role(username)
    if role != "admin":
        raise HTTPException(status_code=403, detail="Forbidden")
        
    import requests
    
    client_id = payload.client_id.strip('"\' ')
    client_secret = payload.client_secret.strip('"\' ')
    code = payload.code.strip()
    
    # Extract code if full URL was pasted
    if "code=" in code:
        import urllib.parse
        try:
            parsed = urllib.parse.urlparse(code)
            code = urllib.parse.parse_qs(parsed.query)["code"][0]
        except Exception:
            pass
            
    token_url = "https://oauth2.googleapis.com/token"
    token_data = {
        "code": code,
        "client_id": client_id,
        "client_secret": client_secret,
        "redirect_uri": "http://localhost",
        "grant_type": "authorization_code"
    }
    
    try:
        resp = requests.post(token_url, data=token_data, timeout=15)
        if resp.status_code != 200:
            raise HTTPException(status_code=400, detail=f"Token exchange failed: {resp.text}")
            
        resp_data = resp.json()
        refresh_token = resp_data.get("refresh_token")
        access_token = resp_data.get("access_token")
        
        if not refresh_token:
            raise HTTPException(
                status_code=400,
                detail="Google did not return a refresh token. Google only sends the refresh token on the FIRST authorization. Please go to Google Account Settings -> Security -> Third-party apps with account access, remove 'Hire AI' access, and authorize again."
            )
            
        # Try to retrieve Google account email using drive/v3/about
        email = "Unknown Google Account"
        if access_token:
            headers = {
                "Authorization": f"Bearer {access_token}"
            }
            about_resp = requests.get("https://www.googleapis.com/drive/v3/about?fields=user", headers=headers, timeout=10)
            if about_resp.status_code == 200:
                user_info = about_resp.json().get("user", {})
                email = user_info.get("emailAddress", "Unknown Google Account")
                
        return {
            "refresh_token": refresh_token,
            "email": email
        }
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to communicate with Google API: {str(e)}")

class TestMailboxRequest(BaseModel):
    imap_host: str
    imap_port: int
    email_user: str
    email_pass: str

@app.post("/api/integrations/test")
def test_mailbox_connection(settings: TestMailboxRequest, request: Request):
    username = request.headers.get("x-user-username")
    role = get_user_role(username)
    if role != "admin":
        raise HTTPException(status_code=403, detail="Forbidden")
        
    email_user = settings.email_user
    email_pass = settings.email_pass
    imap_host = settings.imap_host
    imap_port = settings.imap_port
    
    # If password is masked, restore it from the DB
    if email_pass == "****":
        conn = sqlite3.connect(STATS_DB, timeout=30.0)
        cur = conn.cursor()
        cur.execute("SELECT email_user, email_pass, additional_emails FROM integrations_settings LIMIT 1")
        row = cur.fetchone()
        conn.close()
        
        if row:
            # Check primary
            if row[0] and row[0].lower() == email_user.lower():
                email_pass = row[1]
            else:
                # Check additional
                try:
                    add_emails = json.loads(row[2] or "[]")
                    for item in add_emails:
                        if item.get("email_user") and item["email_user"].lower() == email_user.lower():
                            email_pass = item.get("email_pass", "")
                            break
                except Exception:
                    pass
                    
    if 'office365' in imap_host.lower() or 'outlook' in imap_host.lower():
        # Use Microsoft Graph API
        conn = sqlite3.connect(STATS_DB, timeout=30.0)
        cur = conn.cursor()
        cur.execute("SELECT ms_client_id, ms_client_secret, ms_tenant_id FROM integrations_settings LIMIT 1")
        ms_row = cur.fetchone()
        conn.close()
        
        if not ms_row or not ms_row[0] or not ms_row[1] or not ms_row[2]:
            return {"status": "error", "message": "Microsoft Graph credentials are not configured in Primary Mailbox."}
            
        ms_client_id, ms_client_secret, ms_tenant_id = ms_row
        
        try:
            import requests
            token_url = f"https://login.microsoftonline.com/{ms_tenant_id}/oauth2/v2.0/token"
            data = {
                "client_id": ms_client_id,
                "client_secret": ms_client_secret,
                "scope": "https://graph.microsoft.com/.default",
                "grant_type": "client_credentials"
            }
            token_res = requests.post(token_url, data=data)
            token_data = token_res.json()
            if "access_token" not in token_data:
                return {"status": "error", "message": f"Failed to get Microsoft token: {token_data.get('error_description', 'Unknown error')}"}
                
            access_token = token_data["access_token"]
            headers = {"Authorization": f"Bearer {access_token}"}
            msg_url = f"https://graph.microsoft.com/v1.0/users/{email_user}/mailFolders/Inbox/messages?$top=1"
            msg_res = requests.get(msg_url, headers=headers)
            
            if msg_res.status_code == 200:
                return {"status": "connected", "message": f"Successfully connected to {email_user} via Microsoft Graph API!"}
            else:
                err = msg_res.json().get('error', {})
                return {"status": "error", "message": f"Graph API Error: {err.get('message', 'Unknown')}"}
        except Exception as e:
            return {"status": "error", "message": f"Graph API Connection failed: {str(e)}"}
    else:
        if email_pass == "****" or not email_pass:
            return {"status": "error", "message": "Password is not configured."}
            
        import imaplib
        try:
            mail = imaplib.IMAP4_SSL(imap_host, imap_port, timeout=10)
            mail.login(email_user, email_pass)
            mail.logout()
            return {"status": "connected", "message": f"Successfully connected to {email_user}!"}
        except Exception as e:
            err_msg = str(e)
            if "Application-specific password required" in err_msg:
                return {"status": "error", "message": "Authentication failed: Application-specific password required."}
            return {"status": "error", "message": f"Connection failed: {err_msg}"}

@app.get("/api/integrations/status")
def test_integrations_connection(request: Request):
    username = request.headers.get("x-user-username")
    role = get_user_role(username)
    if role != "admin":
        raise HTTPException(status_code=403, detail="Forbidden")
        
    conn = sqlite3.connect(STATS_DB, timeout=30.0)
    conn.row_factory = sqlite3.Row
    cur = conn.cursor()
    cur.execute("SELECT * FROM integrations_settings LIMIT 1")
    row = cur.fetchone()
    conn.close()
    
    if not row:
        return {"status": "disabled", "message": "Email integration is disabled."}
        
    settings = dict(row)
    if not settings.get("email_enabled"):
        return {"status": "disabled", "message": "Email integration is disabled."}
        
    outlook_enabled = settings.get("outlook_enabled", 0)
    outlook_email = settings.get("outlook_email", "")
    gmail_enabled = settings.get("gmail_enabled", 0)
    gmail_email = settings.get("gmail_email", "")
    gmail_pass = settings.get("gmail_pass", "")
    ms_client_id = settings.get("ms_client_id", "")
    ms_client_secret = settings.get("ms_client_secret", "")
    ms_tenant_id = settings.get("ms_tenant_id", "common")
    imap_host = settings.get("imap_host", "")
    imap_port = settings.get("imap_port", 993)
    email_user = settings.get("email_user", "")
    email_pass = settings.get("email_pass", "")
    
    if outlook_enabled and outlook_email:
        if not ms_client_id or not ms_client_secret or not ms_tenant_id:
            return {"status": "error", "message": "Microsoft Graph credentials are not fully configured."}
        
        try:
            import requests
            token_url = f"https://login.microsoftonline.com/{ms_tenant_id}/oauth2/v2.0/token"
            data = {
                "client_id": ms_client_id,
                "client_secret": ms_client_secret,
                "scope": "https://graph.microsoft.com/.default",
                "grant_type": "client_credentials"
            }
            token_res = requests.post(token_url, data=data)
            token_data = token_res.json()
            if "access_token" not in token_data:
                return {"status": "error", "message": f"Failed to get Microsoft token: {token_data.get('error_description', 'Unknown error')}"}
                
            access_token = token_data["access_token"]
            headers = {"Authorization": f"Bearer {access_token}"}
            msg_url = f"https://graph.microsoft.com/v1.0/users/{outlook_email}/mailFolders/Inbox/messages?$top=1"
            msg_res = requests.get(msg_url, headers=headers)
            
            if msg_res.status_code == 200:
                return {"status": "connected", "message": f"Successfully connected to {outlook_email} via Microsoft Graph API!"}
            else:
                err = msg_res.json().get('error', {})
                return {"status": "error", "message": f"Graph API Error: {err.get('message', 'Unknown')}"}
        except Exception as e:
            return {"status": "error", "message": f"Graph API Connection failed: {str(e)}"}
            
    elif gmail_enabled and gmail_email:
        if not gmail_pass:
            return {"status": "error", "message": "Gmail app password is not configured."}
        import imaplib
        try:
            mail = imaplib.IMAP4_SSL("imap.gmail.com", 993, timeout=10)
            mail.login(gmail_email, gmail_pass)
            mail.logout()
            return {"status": "connected", "message": f"Successfully connected to {gmail_email}!"}
        except Exception as e:
            err_msg = str(e)
            if "Application-specific password required" in err_msg:
                return {"status": "error", "message": "Authentication failed: Application-specific password required."}
            return {"status": "error", "message": f"Connection failed: {err_msg}"}
            
    else:
        if not imap_host or not email_user:
            return {"status": "unconfigured", "message": "Credentials are not fully configured. Enable an integration."}
            
        if 'office365' in imap_host.lower() or 'outlook' in imap_host.lower():
            if not ms_client_id or not ms_client_secret or not ms_tenant_id:
                return {"status": "error", "message": "Microsoft Graph credentials are not configured in Primary Mailbox."}
            
            try:
                import requests
                token_url = f"https://login.microsoftonline.com/{ms_tenant_id}/oauth2/v2.0/token"
                data = {
                    "client_id": ms_client_id,
                    "client_secret": ms_client_secret,
                    "scope": "https://graph.microsoft.com/.default",
                    "grant_type": "client_credentials"
                }
                token_res = requests.post(token_url, data=data)
                token_data = token_res.json()
                if "access_token" not in token_data:
                    return {"status": "error", "message": f"Failed to get Microsoft token: {token_data.get('error_description', 'Unknown error')}"}
                    
                access_token = token_data["access_token"]
                headers = {"Authorization": f"Bearer {access_token}"}
                msg_url = f"https://graph.microsoft.com/v1.0/users/{email_user}/mailFolders/Inbox/messages?$top=1"
                msg_res = requests.get(msg_url, headers=headers)
                
                if msg_res.status_code == 200:
                    return {"status": "connected", "message": f"Successfully connected to {email_user} via Microsoft Graph API!"}
                else:
                    err = msg_res.json().get('error', {})
                    return {"status": "error", "message": f"Graph API Error: {err.get('message', 'Unknown')}"}
            except Exception as e:
                return {"status": "error", "message": f"Graph API Connection failed: {str(e)}"}
        else:
            if not email_pass:
                return {"status": "error", "message": "Password is not configured."}
            import imaplib
            try:
                mail = imaplib.IMAP4_SSL(imap_host, imap_port, timeout=10)
                mail.login(email_user, email_pass)
                mail.logout()
                return {"status": "connected", "message": "Successfully connected to Mailbox!"}
            except Exception as e:
                err_msg = str(e)
                if "Application-specific password required" in err_msg:
                    return {"status": "error", "message": "Authentication failed: Application-specific password required."}
                return {"status": "error", "message": f"Connection failed: {err_msg}"}

def send_email_via_graph(ms_client_id, ms_client_secret, ms_tenant_id, email_user, recipient_email, subject, body):
    import requests
    try:
        token_url = f"https://login.microsoftonline.com/{ms_tenant_id}/oauth2/v2.0/token"
        data = {
            "client_id": ms_client_id,
            "client_secret": ms_client_secret,
            "scope": "https://graph.microsoft.com/.default",
            "grant_type": "client_credentials"
        }
        token_res = requests.post(token_url, data=data)
        access_token = token_res.json().get("access_token")
        if not access_token:
            print("ERROR: Failed to get Graph API token for sending email.")
            return False

        headers = {
            "Authorization": f"Bearer {access_token}",
            "Content-Type": "application/json"
        }
        
        email_data = {
            "message": {
                "subject": subject,
                "body": {
                    "contentType": "Text",
                    "content": body
                },
                "toRecipients": [
                    {
                        "emailAddress": {
                            "address": recipient_email
                        }
                    }
                ]
            },
            "saveToSentItems": "true"
        }
        
        send_url = f"https://graph.microsoft.com/v1.0/users/{email_user}/sendMail"
        res = requests.post(send_url, headers=headers, json=email_data)
        
        if res.status_code == 202:
            return True
        else:
            print(f"ERROR: Graph API sendMail failed: {res.text}")
            return False
    except Exception as e:
        print(f"ERROR: Exception in send_email_via_graph: {e}")
        return False

class TestEmailTemplateRequest(BaseModel):
    recipient_email: str
    preview_type: str
    subject_template: str
    body_template: str

@app.post("/api/settings/test-email-template")
def test_email_template_endpoint(request_data: TestEmailTemplateRequest, request: Request):
    username = request.headers.get("x-user-username")
    role = get_user_role(username)
    if role != "admin":
        raise HTTPException(status_code=403, detail="Forbidden")

    conn = sqlite3.connect(STATS_DB, timeout=30.0)
    cur = conn.cursor()
    cur.execute("SELECT imap_host, smtp_host, smtp_port, email_user, email_pass, ms_client_id, ms_client_secret, ms_tenant_id FROM integrations_settings LIMIT 1")
    row = cur.fetchone()
    conn.close()

    if not row:
        return {"status": "error", "message": "Email integration is not configured."}

    imap_host = row[0]
    smtp_host = row[1]
    smtp_port = row[2] or 587
    email_user = row[3]
    email_pass = row[4]
    ms_client_id = row[5]
    ms_client_secret = row[6]
    ms_tenant_id = row[7]

    if not email_user:
        return {"status": "error", "message": "Email user is not configured."}

    full_name = "Jane Doe"
    missing_fields_str = "Phone, Current Location, Expected CTC" if request_data.preview_type == "missing" else ""
    
    subject = request_data.subject_template.replace('{{full_name}}', full_name)
    body = request_data.body_template.replace('{{full_name}}', full_name).replace('{{missing_fields}}', missing_fields_str)

    try:
        if imap_host and ('office365' in imap_host.lower() or 'outlook' in imap_host.lower()):
            if not ms_client_id or not ms_client_secret or not ms_tenant_id:
                return {"status": "error", "message": "Microsoft Graph credentials are not fully configured."}
            
            success = send_email_via_graph(ms_client_id, ms_client_secret, ms_tenant_id, email_user, request_data.recipient_email, subject, body)
            if success:
                return {"status": "success", "message": f"Test email sent to {request_data.recipient_email}"}
            else:
                return {"status": "error", "message": "Failed to send email via Microsoft Graph API."}
        else:
            import smtplib
            from email.mime.text import MIMEText
            from email.mime.multipart import MIMEMultipart

            if not email_pass:
                return {"status": "error", "message": "SMTP password is not configured."}

            msg = MIMEMultipart()
            msg['From'] = email_user
            msg['To'] = request_data.recipient_email
            msg['Subject'] = subject
            msg.attach(MIMEText(body, 'plain'))

            server = smtplib.SMTP(smtp_host, smtp_port, timeout=10.0)
            server.starttls()
            server.login(email_user, email_pass)
            server.sendmail(email_user, request_data.recipient_email, msg.as_string())
            server.quit()
            
            return {"status": "success", "message": f"Test email sent to {request_data.recipient_email}"}
    except Exception as e:
        return {"status": "error", "message": f"Failed to send email: {str(e)}"}

def process_single_mailbox(email_user, email_pass, imap_host, imap_port, smtp_host, smtp_port, keywords_str, reply_theme, reply_subject, reply_body_missing, reply_body_complete, ms_client_id=None, ms_client_secret=None, ms_tenant_id=None):
    print(f'INFO: Starting to process mailbox for {email_user}')
    import imaplib
    import email
    from email.header import decode_header
    import hashlib
    import smtplib
    from email.mime.text import MIMEText
    from email.mime.multipart import MIMEMultipart

    # Parse keywords
    keywords = [k.strip().lower() for k in keywords_str.split(",") if k.strip()]
    if not keywords:
        keywords = ["resume", "alamaticz", "solution", "job"]

    # Connect IMAP or Graph API
    raw_emails_to_process = []
    
    if 'office365' in imap_host.lower() or 'outlook' in imap_host.lower():
        conn = sqlite3.connect(STATS_DB, timeout=30.0)
        cur = conn.cursor()
        cur.execute("SELECT ms_client_id, ms_client_secret, ms_tenant_id FROM integrations_settings LIMIT 1")
        ms_row = cur.fetchone()
        conn.close()
        if ms_row and ms_row[0] and ms_row[1] and ms_row[2]:
            ms_client_id, ms_client_secret, ms_tenant_id = ms_row
            try:
                import requests
                token_url = f"https://login.microsoftonline.com/{ms_tenant_id}/oauth2/v2.0/token"
                data = {
                    "client_id": ms_client_id,
                    "client_secret": ms_client_secret,
                    "scope": "https://graph.microsoft.com/.default",
                    "grant_type": "client_credentials"
                }
                token_res = requests.post(token_url, data=data)
                access_token = token_res.json().get("access_token")
                if access_token:
                    headers = {"Authorization": f"Bearer {access_token}"}
                    msg_url = f"https://graph.microsoft.com/v1.0/users/{email_user}/mailFolders/Inbox/messages?$top=50&$select=id,internetMessageId"
                    msg_res = requests.get(msg_url, headers=headers)
                    if msg_res.status_code == 200:
                        messages = msg_res.json().get('value', [])
                        messages.reverse() # chronological
                        
                        conn = sqlite3.connect(STATS_DB, timeout=30.0)
                        cur = conn.cursor()
                        
                        for m in messages:
                            m_id = m['id']
                            internet_msg_id = m.get('internetMessageId', "")
                            if internet_msg_id:
                                import re
                                match = re.search(r'<([^>]+)>', internet_msg_id)
                                if match:
                                    internet_msg_id = match.group(1)
                                
                            effective_msg_id = internet_msg_id if internet_msg_id else m_id
                            cur.execute("SELECT 1 FROM processed_emails WHERE msg_uid = ?", (effective_msg_id,))
                            exists = cur.fetchone()
                            if exists:
                                continue
                                    
                            mime_url = f"https://graph.microsoft.com/v1.0/users/{email_user}/messages/{m_id}/$value"
                            mime_res = requests.get(mime_url, headers=headers)
                            if mime_res.status_code == 200:
                                raw_emails_to_process.append((mime_res.content, effective_msg_id))
                                
                        conn.close()
            except Exception as e:
                print(f"Graph API Error: {e}")
    else:
        mail = None
        try:
            mail = imaplib.IMAP4_SSL(imap_host, imap_port or 993, timeout=15)
            mail.login(email_user, email_pass)
            status_select, select_data = mail.select("inbox")

            unseen_nums = []
            status_unseen, response_unseen = mail.search(None, "UNSEEN")
            if status_unseen == "OK" and response_unseen[0]:
                unseen_nums = response_unseen[0].split()

            total_msgs = int(select_data[0]) if status_select == "OK" and select_data[0] else 0
            
            recent_nums = []
            if total_msgs > 0:
                recent_nums = [str(i).encode() for i in range(max(1, total_msgs - 49), total_msgs + 1)]

            msg_nums_set = set(unseen_nums + recent_nums)
            msg_nums = sorted(list(msg_nums_set), key=lambda x: int(x))
            
            conn = sqlite3.connect(STATS_DB, timeout=30.0)
            cur = conn.cursor()
            
            for num in msg_nums:
                try:
                    status, msg_data = mail.fetch(num, '(BODY.PEEK[HEADER.FIELDS (MESSAGE-ID DATE FROM SUBJECT)])')
                    if status == "OK" and msg_data[0]:
                        header_data = msg_data[0][1].decode('utf-8', errors='ignore')
                        msg_id = ""
                        for line in header_data.split('\n'):
                            if line.lower().startswith('message-id:'):
                                msg_id = line.split(':', 1)[1].strip()
                                import re
                                match = re.search(r'<([^>]+)>', msg_id)
                                if match:
                                    msg_id = match.group(1)
                                break
                        
                        if not msg_id:
                            import hashlib
                            msg_id = "hash_" + hashlib.md5(header_data.encode('utf-8')).hexdigest()
                            
                        cur.execute("SELECT 1 FROM processed_emails WHERE msg_uid = ?", (msg_id,))
                        if cur.fetchone():
                            continue
                                
                        status_full, msg_data_full = mail.fetch(num, '(BODY.PEEK[])')
                        if status_full == "OK" and msg_data_full[0]:
                            raw_emails_to_process.append((msg_data_full[0][1], msg_id))
                except Exception as e:
                    print(f"Error fetching IMAP email: {e}")
                    
            conn.close()
        except Exception as e:
            print(f"IMAP Error: {e}")
        finally:
            if mail:
                try:
                    mail.logout()
                except:
                    pass

    # Process all fetched emails
    for item in raw_emails_to_process:
        try:
            if isinstance(item, tuple):
                raw_email, effective_msg_id = item
            else:
                raw_email = item
                effective_msg_id = None
                
            msg = email.message_from_bytes(raw_email)
            
            # Extract Message-ID
            if effective_msg_id:
                msg_id = effective_msg_id
            else:
                msg_id = msg.get("Message-ID", "")
                if msg_id:
                    import re
                    match = re.search(r'<([^>]+)>', msg_id)
                    if match:
                        msg_id = match.group(1)
                
                if not msg_id:
                    # Fallback: hash of headers to make a pseudo ID
                    subject_header = str(msg.get("Subject", ""))
                    date_header = str(msg.get("Date", ""))
                    from_header = str(msg.get("From", ""))
                    header_text2 = f"Subject: {subject_header}\nDate: {date_header}\nFrom: {from_header}\n\n"
                    msg_id = hashlib.md5(header_text2.encode('utf-8', errors='ignore')).hexdigest()

            # Check if already processed
            conn = sqlite3.connect(STATS_DB, timeout=30.0)
            cur = conn.cursor()
            cur.execute("SELECT 1 FROM processed_emails WHERE msg_uid = ?", (msg_id,))
            exists = cur.fetchone()
            conn.close()

            if exists:
                continue

            # Extract Subject, From, Body
            subject_header = msg.get("Subject", "")
            decoded_subject = ""
            for part, encoding in decode_header(subject_header):
                if isinstance(part, bytes):
                    decoded_subject += part.decode(encoding or "utf-8", errors="ignore")
                else:
                    decoded_subject += str(part)
            
            from_header = msg.get("From", "")
            decoded_from = ""
            for part, encoding in decode_header(from_header):
                if isinstance(part, bytes):
                    decoded_from += part.decode(encoding or "utf-8", errors="ignore")
                else:
                    decoded_from += str(part)

            # --- PREVENT INFINITE BOUNCE LOOPS ---
            subject_lower = decoded_subject.lower()
            from_lower = decoded_from.lower()
            is_bounce = False
            
            if "undeliverable:" in subject_lower or "delivery status notification" in subject_lower or "delivery failure" in subject_lower or "returned mail" in subject_lower:
                is_bounce = True
            if "postmaster@" in from_lower or "mailer-daemon@" in from_lower or "noreply@" in from_lower or "no-reply@" in from_lower:
                is_bounce = True
                
            if is_bounce:
                conn = sqlite3.connect(STATS_DB, timeout=30.0)
                cur = conn.cursor()
                cur.execute("INSERT OR IGNORE INTO processed_emails (msg_uid) VALUES (?)", (msg_id,))
                conn.commit()
                conn.close()
                print(f"INFO: Skipping bounce/automated message from {decoded_from} with subject '{decoded_subject}' to prevent infinite auto-reply loop.")
                continue

            # Extract plain text body
            body_text = ""
            attachments = []
            if msg.is_multipart():
                for part in msg.walk():
                    content_type = part.get_content_type()
                    content_disposition = str(part.get("Content-Disposition"))
                    if content_type == "text/plain" and "attachment" not in content_disposition:
                        payload = part.get_payload(decode=True)
                        if payload:
                            body_text += payload.decode(part.get_content_charset() or "utf-8", errors="ignore")
                    elif "attachment" in content_disposition or part.get_filename():
                        filename = part.get_filename()
                        if filename:
                            decoded_filename = ""
                            for filename_part, encoding in decode_header(filename):
                                if isinstance(filename_part, bytes):
                                    decoded_filename += filename_part.decode(encoding or "utf-8", errors="ignore")
                                else:
                                    decoded_filename += str(filename_part)
                            if decoded_filename:
                                attachments.append((decoded_filename, part.get_payload(decode=True)))
            else:
                payload = msg.get_payload(decode=True)
                if payload:
                    body_text = payload.decode(msg.get_content_charset() or "utf-8", errors="ignore")
            
            # Check if this email is from an existing candidate (via Ref tag in Subject, or matching sender email)
            import email.utils
            sender_name, sender_email = email.utils.parseaddr(decoded_from)
            matched_candidate_row = None

            # 1. Search by reference ID in subject
            match_ref = re.search(r'Ref:\s*CAND-(\d+)', decoded_subject, re.IGNORECASE)
            if match_ref:
                try:
                    ref_candidate_id = int(match_ref.group(1))
                    conn = sqlite3.connect(STATS_DB, timeout=30.0)
                    cur = conn.cursor()
                    cur.execute("SELECT * FROM candidate_metadata WHERE id = ?", (ref_candidate_id,))
                    matched_candidate_row = cur.fetchone()
                    conn.close()
                except Exception as e_ref:
                    print(f"ERROR matching by Ref ID in subject: {e_ref}")
            # 2. Fall back to search by sender email address
            if not matched_candidate_row and sender_email:
                try:
                    conn = sqlite3.connect(STATS_DB, timeout=30.0)
                    cur = conn.cursor()
                    cur.execute("SELECT * FROM candidate_metadata WHERE LOWER(sender_email) = ? OR LOWER(email) = ? ORDER BY id DESC LIMIT 1", (sender_email.lower(), sender_email.lower()))
                    matched_candidate_row = cur.fetchone()
                    conn.close()
                except Exception as e_email:
                    print(f"ERROR matching by sender email: {e_email}")

            if matched_candidate_row:
                existing_candidate = dict(matched_candidate_row)
                print(f"INFO: Identified follow-up email from existing candidate ID {existing_candidate['id']} ({existing_candidate['full_name']})")
            
                # Check if a new resume file is attached and download/index it
                new_filename = None
                new_file_bytes = None
                resume_text = ""
                if attachments:
                    has_resume = any(
                        fname.lower().endswith((".pdf", ".docx"))
                        for fname, _ in attachments if fname
                    )
                    if has_resume:
                        for fname, content in attachments:
                            if not content:
                                continue
                            f_lower = fname.lower()
                            if f_lower.endswith(".pdf") or f_lower.endswith(".docx"):
                                safe_name = re.sub(r'[^a-zA-Z0-9._-]', '_', fname)
                                safe_name = f"mail_{hashlib.md5(msg_id.encode()).hexdigest()[:8]}_{safe_name}"
                                fpath = os.path.join(UPLOAD_DIR, safe_name)
                                with open(fpath, "wb") as f_out:
                                    f_out.write(content)
                                new_filename = safe_name
                                new_file_bytes = content
                            
                                # Index in PGVector and extract text
                                try:
                                    if safe_name.lower().endswith(".pdf"):
                                        loader = SafePyMuPDFLoader(fpath)
                                    else:
                                        loader = Docx2txtLoader(fpath)
                                    docs = loader.load()
                                    resume_text = "\n".join([d.page_content for d in docs])
                                    embeddings, _ = get_models()
                                    for d in docs:
                                        d.metadata['source'] = safe_name
                                    splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=100)
                                    chunks = splitter.split_documents(docs)
                                    PGVector.from_documents(
                                        documents=chunks,
                                        embedding=embeddings,
                                        connection_string=os.getenv("POSTGRES_DATABASE_URL"),
                                        collection_name="resume_embeddings"
                                    )
                                except Exception as pg_err:
                                    print(f"Error adding updated resume to PGVector: {pg_err}")
                                finally:
                                    try:
                                        if os.path.exists(fpath):
                                            os.remove(fpath)
                                    except Exception:
                                        pass
                                break

                # Use EXTRACT_PROMPT to parse the combined new resume content and/or email text
                combined_text = resume_text[:7000] if resume_text else ""
                if body_text:
                    combined_text += f"\n\n=== EMAIL MESSAGE BODY ===\n{body_text}\n=========================="

                conn = sqlite3.connect(STATS_DB, timeout=30.0)
                cur = conn.cursor()
                cur.execute("SELECT col_key, col_label, description FROM custom_columns")
                custom_cols = cur.fetchall()
                conn.close()
            
                custom_fields_str = ""
                if custom_cols:
                    for col_key, col_label, desc in custom_cols:
                        desc_str = f"Extract data corresponding to column heading '{col_label}'"
                        if desc:
                            desc_str += f" ({desc})"
                        custom_fields_str += f',\n  "{col_key}": "<{desc_str}>"'
                    
                prompt_str = EXTRACT_PROMPT.format(text=combined_text, custom_fields=custom_fields_str)
            
                _, llm = get_models()
                parsed_data = {}
                try:
                    # Add a simple retry mechanism for rate limits
                    max_retries = 3
                    resp = None
                    for attempt in range(max_retries):
                        try:
                            resp = llm.invoke([HumanMessage(content=prompt_str)])
                            break
                        except Exception as api_err:
                            if "429" in str(api_err) and attempt < max_retries - 1:
                                import time
                                time.sleep(3)
                                continue
                            raise api_err
                        
                    if resp is not None:
                        raw_resp = resp.content.strip()
                        if "```json" in raw_resp:
                            raw_resp = raw_resp.split("```json")[1].split("```")[0].strip()
                        elif "```" in raw_resp:
                            raw_resp = raw_resp.split("```")[1].split("```")[0].strip()
                        start_idx, end_idx = raw_resp.find('{'), raw_resp.rfind('}')
                        if start_idx != -1 and end_idx != -1:
                            raw_resp = raw_resp[start_idx:end_idx+1]
                        parsed_data = json.loads(raw_resp)
                except Exception as llm_err:
                    print(f"ERROR parsing follow-up email/resume via LLM: {llm_err}")

                # Normalization & Validation
                # Phone: Keep only digits and +
                if 'phone' in parsed_data and parsed_data['phone']:
                    parsed_data['phone'] = re.sub(r'[^\d+]', '', str(parsed_data['phone']))
                
                # Email: Basic validation
                if 'email' in parsed_data and parsed_data['email']:
                    if '@' not in str(parsed_data['email']):
                        parsed_data['email'] = ""
                    
                # Experience: Force float
                for exp_field in ['total_experience', 'pega_experience', 'cdh_exp']:
                    if exp_field in parsed_data and parsed_data[exp_field] not in [None, ""]:
                        try:
                            match = re.search(r'\d+(\.\d+)?', str(parsed_data[exp_field]))
                            parsed_data[exp_field] = float(match.group()) if match else 0.0
                        except Exception:
                            parsed_data[exp_field] = 0.0
                        
                # Integer fields
                for num_field in ['notice_period', 'availability_in_days']:
                    if num_field in parsed_data and parsed_data[num_field] not in [None, ""]:
                        np_str = str(parsed_data[num_field]).lower()
                        if 'immediate' in np_str:
                            parsed_data[num_field] = 0
                        else:
                            try:
                                match = re.search(r'\d+', np_str)
                                val = int(match.group()) if match else ""
                                if match and 'month' in np_str:
                                    val = val * 30
                                parsed_data[num_field] = val
                            except Exception:
                                parsed_data[num_field] = ""

                # Update existing row with all dynamic columns
                updates = {}
                conn = sqlite3.connect(STATS_DB, timeout=30.0)
                cur = conn.cursor()
                cur.execute("PRAGMA table_info(candidate_metadata)")
                allowed_cols = {c[1] for c in cur.fetchall()}
                conn.close()

                for k, v in parsed_data.items():
                    if k in allowed_cols and k not in ('id', 'filename', 'full_name'):
                        if v is not None and str(v).strip() not in ("", "null", "None"):
                            updates[k] = v

                if new_filename:
                    updates['filename'] = new_filename
                    if new_file_bytes:
                        updates['file_bytes'] = new_file_bytes
                if sender_email:
                    updates['sender_email'] = sender_email
                updates['source'] = 'uploaded from mail'

                # Merge email message body
                old_msg = existing_candidate.get('email_message') or ""
                new_msg = f"{old_msg}\n\n=== FOLLOW-UP EMAIL MESSAGE ===\n{body_text}".strip()
                updates['email_message'] = new_msg
            
                # Reset formatted resume cache
                updates['formatted_json'] = None

                if updates:
                    try:
                        conn = sqlite3.connect(STATS_DB, timeout=30.0)
                        cur = conn.cursor()
                        set_clause = ", ".join(f"{col}=?" for col in updates)
                        cur.execute(f"UPDATE candidate_metadata SET {set_clause} WHERE id=?", list(updates.values()) + [existing_candidate['id']])
                        conn.commit()
                        conn.close()
                        print(f"INFO: Successfully updated candidate ID {existing_candidate['id']} with follow-up details.")
                    except Exception as db_update_err:
                        print(f"ERROR updating candidate from follow-up email: {db_update_err}")

                # Recheck missing fields on updated profile
                try:
                    conn = sqlite3.connect(STATS_DB, timeout=30.0)
                    cur = conn.cursor()
                    cur.execute("SELECT * FROM candidate_metadata WHERE id=?", (existing_candidate['id'],))
                    updated_candidate = row_to_dict(cur.fetchone())
                    conn.close()

                    candidate_name = updated_candidate.get('full_name') or sender_name or 'Candidate'
                    missing_fields = []
                    total_exp = updated_candidate.get('total_experience')
                    if total_exp is None or str(total_exp).strip() == "" or float(total_exp) == 0.0:
                        missing_fields.append("Total years of experience")
                
                    pega_exp = updated_candidate.get('pega_experience')
                    cdh_exp = updated_candidate.get('cdh_exp')
                    has_pega = pega_exp is not None and str(pega_exp).strip() != "" and float(pega_exp) > 0.0
                    has_cdh = cdh_exp is not None and str(cdh_exp).strip() != "" and float(cdh_exp) > 0.0
                    if not has_pega and not has_cdh:
                        missing_fields.append("Relevant experience for this role")
                
                    ctc = updated_candidate.get('ctc')
                    if not ctc or str(ctc).strip() in ("", "—", "-", "None", "null"):
                        missing_fields.append("Current CTC")
                
                    expected_ctc = updated_candidate.get('expected_ctc')
                    if not expected_ctc or str(expected_ctc).strip() in ("", "—", "-", "None", "null"):
                        missing_fields.append("Expected CTC")
                
                    notice_period = updated_candidate.get('notice_period')
                    if notice_period is None or str(notice_period).strip() in ("", "—", "-", "None", "null"):
                        missing_fields.append("Notice period / Earliest joining date")
                
                    current_location = updated_candidate.get('current_location')
                    if not current_location or str(current_location).strip() in ("", "—", "-", "None", "null"):
                        missing_fields.append("Current location")
                
                    pref_locations = updated_candidate.get('pref_locations')
                    if not pref_locations or str(pref_locations).strip() in ("", "—", "-", "None", "null"):
                        missing_fields.append("Preferred work location(s)")
                
                    linkedin = updated_candidate.get('linkedin')
                    if not linkedin or str(linkedin).strip() in ("", "—", "-", "None", "null"):
                        missing_fields.append("LinkedIn profile URL")

                    if missing_fields:
                        # Resolve templates
                        theme_name = reply_theme or 'professional'
                        subj_tpl = reply_subject
                        body_missing_tpl = reply_body_missing
                        body_complete_tpl = reply_body_complete

                        if theme_name in THEME_TEMPLATES:
                            if theme_name != 'custom':
                                subj_tpl = THEME_TEMPLATES[theme_name]['subject']
                                body_missing_tpl = THEME_TEMPLATES[theme_name]['body_missing']
                                body_complete_tpl = THEME_TEMPLATES[theme_name]['body_complete']
                            else:
                                if not subj_tpl: subj_tpl = THEME_TEMPLATES['professional']['subject']
                                if not body_missing_tpl: body_missing_tpl = THEME_TEMPLATES['professional']['body_missing']
                                if not body_complete_tpl: body_complete_tpl = THEME_TEMPLATES['professional']['body_complete']
                        else:
                            if not subj_tpl: subj_tpl = THEME_TEMPLATES['professional']['subject']
                            if not body_missing_tpl: body_missing_tpl = THEME_TEMPLATES['professional']['body_missing']
                            if not body_complete_tpl: body_complete_tpl = THEME_TEMPLATES['professional']['body_complete']

                        missing_list_str = "\n".join(f"* {field}" for field in missing_fields)
                        body_reply = body_missing_tpl.replace("{candidate_name}", candidate_name).replace("{missing_fields}", missing_list_str)
                    else:
                        # Resolve templates
                        theme_name = reply_theme or 'professional'
                        subj_tpl = reply_subject
                        body_missing_tpl = reply_body_missing
                        body_complete_tpl = reply_body_complete

                        if theme_name in THEME_TEMPLATES:
                            if theme_name != 'custom':
                                subj_tpl = THEME_TEMPLATES[theme_name]['subject']
                                body_missing_tpl = THEME_TEMPLATES[theme_name]['body_missing']
                                body_complete_tpl = THEME_TEMPLATES[theme_name]['body_complete']
                            else:
                                if not subj_tpl: subj_tpl = THEME_TEMPLATES['professional']['subject']
                                if not body_missing_tpl: body_missing_tpl = THEME_TEMPLATES['professional']['body_missing']
                                if not body_complete_tpl: body_complete_tpl = THEME_TEMPLATES['professional']['body_complete']
                        else:
                            if not subj_tpl: subj_tpl = THEME_TEMPLATES['professional']['subject']
                            if not body_missing_tpl: body_missing_tpl = THEME_TEMPLATES['professional']['body_missing']
                            if not body_complete_tpl: body_complete_tpl = THEME_TEMPLATES['professional']['body_complete']

                        body_reply = body_complete_tpl.replace("{candidate_name}", candidate_name)

                    candidate_email = updated_candidate.get('email')
                    recipient_email = candidate_email if candidate_email and str(candidate_email).strip() not in ("", "null", "None") else sender_email

                    ack_msg = MIMEMultipart()
                    ack_msg['From'] = email_user
                    ack_msg['To'] = recipient_email
                    ack_msg['Subject'] = subj_tpl.replace("{subject}", decoded_subject).replace("{ref}", f"CAND-{updated_candidate['id']}")
                    ack_msg.attach(MIMEText(body_reply, 'plain'))

                    if recipient_email.lower() == email_user.lower():
                        print(f"INFO: Skipping auto-reply to self ({recipient_email}) to prevent infinite loop.")
                    else:
                        is_outlook = 'office365' in imap_host.lower() or 'outlook' in imap_host.lower()
                        if is_outlook and ms_client_id and ms_client_secret and ms_tenant_id:
                            print(f"INFO: Sending follow-up acknowledgment via Graph API to {recipient_email}")
                            subject = subj_tpl.replace("{subject}", decoded_subject).replace("{ref}", f"CAND-{updated_candidate['id']}")
                            send_email_via_graph(ms_client_id, ms_client_secret, ms_tenant_id, email_user, recipient_email, subject, body_reply)
                        else:
                            s_server = smtplib.SMTP(smtp_host, smtp_port or 587, timeout=10.0)
                            s_server.starttls()
                            s_server.login(email_user, email_pass)
                            s_server.sendmail(email_user, recipient_email, ack_msg.as_string())
                            s_server.quit()
                            print(f"INFO: Sent follow-up acknowledgment to {recipient_email}")
                except Exception as reply_err:
                    print(f"ERROR sending follow-up email reply: {reply_err}")

                # Mark email as read and processed
                try:
                    conn = sqlite3.connect(STATS_DB, timeout=30.0)
                    cur = conn.cursor()
                    cur.execute("INSERT INTO processed_emails (msg_uid) VALUES (?)", (msg_id,))
                    conn.commit()
                    conn.close()
                except Exception as mark_err:
                    print(f"ERROR marking email as processed: {mark_err}")
                
                continue

            # Check if the email contains any resume attachments (.pdf or .docx)
            has_resume_attachment = False
            if attachments:
                has_resume_attachment = any(
                    fname.lower().endswith((".pdf", ".docx"))
                    for fname, _ in attachments if fname
                )

            # Always match if there is a resume attachment, otherwise fall back to keyword match
            match_found = has_resume_attachment
            if not match_found:
                attachment_names = " ".join([fname for fname, _ in attachments if fname]) if attachments else ""
                search_content = f"{decoded_subject} {body_text} {attachment_names}".lower()
                for kw in keywords:
                    if kw in search_content:
                        match_found = True
                        break

            processed_attachment = False
            if match_found and attachments:
                # Only proceed if there is at least one resume attachment (.pdf or .docx)
                has_resume = any(
                    fname.lower().endswith(".pdf") or fname.lower().endswith(".docx")
                    for fname, _ in attachments if fname
                )
                if has_resume:
                    # Process attachments
                    for fname, content in attachments:
                        if not content:
                            continue
                        f_lower = fname.lower()
                        if f_lower.endswith(".pdf") or f_lower.endswith(".docx"):
                            safe_name = re.sub(r'[^a-zA-Z0-9._-]', '_', fname)
                            safe_name = f"mail_{hashlib.md5(msg_id.encode()).hexdigest()[:8]}_{safe_name}"
                            fpath = os.path.join(UPLOAD_DIR, safe_name)
                            with open(fpath, "wb") as f_out:
                                f_out.write(content)
                        
                            try:
                                process_resume(safe_name, fpath, is_approved=1, username="email_worker", email_message=body_text, sender_email=sender_email)
                                processed_attachment = True
                            except Exception as e_proc:
                                print(f"ERROR: Failed processing resume {safe_name} from email: {e_proc}")
            
                if has_resume_attachment and not processed_attachment:
                    print(f"WARNING: Email {msg_id} has resume attachment, but processing failed. Skipping DB log to allow retry.")
                    continue
                
                if processed_attachment:
                    # Mark email as read on server
                    if mail:
                        try:
                            mail.store(num, '+FLAGS', '\\Seen')
                        except:
                            pass
                
                    # Send auto acknowledgment
                    try:
                        import email.utils
                        sender_name, sender_email = email.utils.parseaddr(from_header)
                        if sender_email:
                            # Fetch parsed candidate metadata to find missing fields
                            candidate = None
                            try:
                                conn = sqlite3.connect(STATS_DB, timeout=30.0)
                                cur = conn.cursor()
                                cur.execute("SELECT * FROM candidate_metadata WHERE filename = ? ORDER BY id DESC LIMIT 1", (safe_name,))
                                candidate_row = cur.fetchone()
                                if candidate_row:
                                    candidate = dict(candidate_row)
                                conn.close()
                            except Exception as db_err:
                                print(f"ERROR: Failed to fetch candidate metadata for auto-acknowledgement: {db_err}")

                            candidate_name = 'Candidate'
                            missing_fields = []
                            if candidate:
                                candidate_name = candidate.get('full_name') or sender_name or 'Candidate'
                            
                                # 1. Total years of experience
                                total_exp = candidate.get('total_experience')
                                if total_exp is None or str(total_exp).strip() == "" or float(total_exp) == 0.0:
                                    missing_fields.append("Total years of experience")
                                
                                # 2. Relevant experience for this role
                                pega_exp = candidate.get('pega_experience')
                                cdh_exp = candidate.get('cdh_exp')
                                has_pega = pega_exp is not None and str(pega_exp).strip() != "" and float(pega_exp) > 0.0
                                has_cdh = cdh_exp is not None and str(cdh_exp).strip() != "" and float(cdh_exp) > 0.0
                                if not has_pega and not has_cdh:
                                    missing_fields.append("Relevant experience for this role")
                                
                                # 3. Current CTC
                                ctc = candidate.get('ctc')
                                if not ctc or str(ctc).strip() in ("", "—", "-", "None", "null"):
                                    missing_fields.append("Current CTC")
                                
                                # 4. Expected CTC
                                expected_ctc = candidate.get('expected_ctc')
                                if not expected_ctc or str(expected_ctc).strip() in ("", "—", "-", "None", "null"):
                                    missing_fields.append("Expected CTC")
                                
                                # 5. Notice period / Earliest joining date
                                notice_period = candidate.get('notice_period')
                                if notice_period is None or str(notice_period).strip() in ("", "—", "-", "None", "null"):
                                    missing_fields.append("Notice period / Earliest joining date")
                                
                                # 6. Current location
                                current_location = candidate.get('current_location')
                                if not current_location or str(current_location).strip() in ("", "—", "-", "None", "null"):
                                    missing_fields.append("Current location")
                                
                                # 7. Preferred work location(s)
                                pref_locations = candidate.get('pref_locations')
                                if not pref_locations or str(pref_locations).strip() in ("", "—", "-", "None", "null"):
                                    missing_fields.append("Preferred work location(s)")
                                
                                # 8. LinkedIn profile URL
                                linkedin = candidate.get('linkedin')
                                if not linkedin or str(linkedin).strip() in ("", "—", "-", "None", "null"):
                                    missing_fields.append("LinkedIn profile URL")
                            else:
                                candidate_name = sender_name or 'Candidate'
                                missing_fields = [
                                    "Total years of experience",
                                    "Relevant experience for this role",
                                    "Current CTC",
                                    "Expected CTC",
                                    "Notice period / Earliest joining date",
                                    "Current location",
                                    "Preferred work location(s)",
                                    "LinkedIn profile URL"
                                ]

                            # Resolve templates
                            theme_name = reply_theme or 'professional'
                            subj_tpl = reply_subject
                            body_missing_tpl = reply_body_missing
                            body_complete_tpl = reply_body_complete

                            if theme_name in THEME_TEMPLATES:
                                if theme_name != 'custom':
                                    subj_tpl = THEME_TEMPLATES[theme_name]['subject']
                                    body_missing_tpl = THEME_TEMPLATES[theme_name]['body_missing']
                                    body_complete_tpl = THEME_TEMPLATES[theme_name]['body_complete']
                                else:
                                    if not subj_tpl: subj_tpl = THEME_TEMPLATES['professional']['subject']
                                    if not body_missing_tpl: body_missing_tpl = THEME_TEMPLATES['professional']['body_missing']
                                    if not body_complete_tpl: body_complete_tpl = THEME_TEMPLATES['professional']['body_complete']
                            else:
                                if not subj_tpl: subj_tpl = THEME_TEMPLATES['professional']['subject']
                                if not body_missing_tpl: body_missing_tpl = THEME_TEMPLATES['professional']['body_missing']
                                if not body_complete_tpl: body_complete_tpl = THEME_TEMPLATES['professional']['body_complete']

                            if missing_fields:
                                missing_list_str = "\n".join(f"* {field}" for field in missing_fields)
                                body_reply = body_missing_tpl.replace("{candidate_name}", candidate_name).replace("{missing_fields}", missing_list_str)
                            else:
                                body_reply = body_complete_tpl.replace("{candidate_name}", candidate_name)

                            candidate_email = candidate.get('email') if candidate else None
                            recipient_email = candidate_email if candidate_email and str(candidate_email).strip() not in ("", "null", "None") else sender_email

                            ack_msg = MIMEMultipart()
                            ack_msg['From'] = email_user
                            ack_msg['To'] = recipient_email
                            cand_id = candidate.get('id') if candidate else 0
                            ref_str = f"CAND-{cand_id}" if cand_id else "NEW"
                            ack_msg['Subject'] = subj_tpl.replace("{subject}", decoded_subject).replace("{ref}", ref_str)
                            ack_msg.attach(MIMEText(body_reply, 'plain'))
                        
                            if recipient_email.lower() == email_user.lower():
                                print(f"INFO: Skipping auto-reply to self ({recipient_email}) to prevent infinite loop.")
                            else:
                                is_outlook = 'office365' in imap_host.lower() or 'outlook' in imap_host.lower()
                                if is_outlook and ms_client_id and ms_client_secret and ms_tenant_id:
                                    print(f"INFO: Sending auto-acknowledgement via Graph API to {recipient_email}")
                                    cand_id = candidate.get('id') if candidate else 0
                                    ref_str = f"CAND-{cand_id}" if cand_id else "NEW"
                                    subject = subj_tpl.replace("{subject}", decoded_subject).replace("{ref}", ref_str)
                                    send_email_via_graph(ms_client_id, ms_client_secret, ms_tenant_id, email_user, recipient_email, subject, body_reply)
                                else:
                                    s_server = smtplib.SMTP(smtp_host, smtp_port or 587, timeout=10.0)
                                    s_server.starttls()
                                    s_server.login(email_user, email_pass)
                                    s_server.sendmail(email_user, recipient_email, ack_msg.as_string())
                                    s_server.quit()
                                    print(f"INFO: Sent auto-acknowledgement to {recipient_email}")
                    except Exception as smtp_err:
                        print(f"ERROR: Failed sending auto-acknowledgement: {smtp_err}")

            # Insert to processed_emails
            conn = sqlite3.connect(STATS_DB, timeout=30.0)
            cur = conn.cursor()
            cur.execute("INSERT OR IGNORE INTO processed_emails (msg_uid) VALUES (?)", (msg_id,))
            conn.commit()
            conn.close()

        except Exception as msg_err:
            print(f"ERROR: Failed processing single email: {msg_err}")

def poll_emails_and_process():
    import time
    print("INFO: Starting background email worker thread...")
    while True:
        try:
            # Fetch settings from DB (robust: handle missing columns gracefully)
            conn = sqlite3.connect(STATS_DB, timeout=30.0)
            conn.row_factory = sqlite3.Row
            cur = conn.cursor()
            cur.execute("SELECT * FROM integrations_settings LIMIT 1")
            row = cur.fetchone()
            conn.close()

            if not row:
                time.sleep(30)
                continue

            settings = row_to_dict(row)
            email_enabled = settings.get("email_enabled", 0)
            imap_host = settings.get("imap_host", "")
            imap_port = settings.get("imap_port", 993)
            smtp_host = settings.get("smtp_host", "smtp.gmail.com")
            smtp_port = settings.get("smtp_port", 587)
            email_user = settings.get("email_user", "")
            email_pass = settings.get("email_pass", "")
            keywords_str = settings.get("keywords", "resume,alamaticz,solution,job")
            reply_theme = settings.get("reply_theme", "professional")
            reply_subject = settings.get("reply_subject", "Re: {subject} (Ref: {ref})")
            reply_body_missing = settings.get("reply_body_missing", None)
            reply_body_complete = settings.get("reply_body_complete", None)
            additional_emails_raw = settings.get("additional_emails", None)
            ms_client_id = settings.get("ms_client_id", "")
            ms_client_secret = settings.get("ms_client_secret", "")
            ms_tenant_id = settings.get("ms_tenant_id", "common")

            gmail_enabled = settings.get("gmail_enabled", 0)
            gmail_email = settings.get("gmail_email", "")
            gmail_pass = settings.get("gmail_pass", "")
            outlook_enabled = settings.get("outlook_enabled", 0)
            outlook_email = settings.get("outlook_email", "")

            # 1. Process legacy primary email if enabled
            if email_enabled and imap_host and email_user and email_pass:
                try:
                    process_single_mailbox(
                        email_user, email_pass, imap_host, imap_port, 
                        smtp_host, smtp_port, keywords_str,
                        reply_theme, reply_subject, reply_body_missing, reply_body_complete,
                        ms_client_id, ms_client_secret, ms_tenant_id
                    )
                except Exception as prim_err:
                    print(f"ERROR processing legacy primary mailbox {email_user}: {prim_err}")

            # 1b. Process Gmail if enabled
            if gmail_enabled and gmail_email and gmail_pass:
                try:
                    process_single_mailbox(
                        gmail_email, gmail_pass, "imap.gmail.com", 993, 
                        "smtp.gmail.com", 587, keywords_str,
                        reply_theme, reply_subject, reply_body_missing, reply_body_complete,
                        None, None, None
                    )
                except Exception as gmail_err:
                    print(f"ERROR processing Gmail mailbox {gmail_email}: {gmail_err}")

            # 1c. Process Outlook if enabled
            if outlook_enabled and outlook_email and ms_client_id and ms_client_secret and ms_tenant_id:
                try:
                    process_single_mailbox(
                        outlook_email, "dummy_pass", "outlook.office365.com", 993, 
                        "smtp.office365.com", 587, keywords_str,
                        reply_theme, reply_subject, reply_body_missing, reply_body_complete,
                        ms_client_id, ms_client_secret, ms_tenant_id
                    )
                except Exception as outlook_err:
                    print(f"ERROR processing Outlook mailbox {outlook_email}: {outlook_err}")

            # 2. Process additional emails if configured
            if additional_emails_raw:
                try:
                    add_emails = json.loads(additional_emails_raw)
                    for item in add_emails:
                        if item.get("email_enabled") and item.get("email_user") and item.get("email_pass"):
                            try:
                                process_single_mailbox(
                                    item.get("email_user"),
                                    item.get("email_pass"),
                                    item.get("imap_host", "imap.gmail.com"),
                                    item.get("imap_port", 993),
                                    item.get("smtp_host", "smtp.gmail.com"),
                                    item.get("smtp_port", 587),
                                    keywords_str,
                                    reply_theme,
                                    reply_subject,
                                    reply_body_missing,
                                    reply_body_complete,
                                    ms_client_id, 
                                    ms_client_secret, 
                                    ms_tenant_id
                                )
                            except Exception as add_err:
                                print(f"ERROR processing additional mailbox {item.get('email_user')}: {add_err}")
                except Exception as json_err:
                    print(f"ERROR parsing additional_emails from DB: {json_err}")

        except Exception as conn_err:
            print(f"ERROR: Email background loop connection error: {conn_err}")
        
        # Poll every 1 minute as requested
        time.sleep(60)

# ── Serve React Frontend ───────────────────────────────────────────────────────
FRONTEND_DIST = os.path.join(PROJECT_ROOT, "frontend", "dist")
if os.path.exists(FRONTEND_DIST):
    app.mount("/", StaticFiles(directory=FRONTEND_DIST, html=True), name="frontend")

@app.exception_handler(StarletteHTTPException)
async def catch_all_spa_routes(request: Request, exc: StarletteHTTPException):
    # If a 404 occurs and it's not an API call, serve the React index.html for client-side routing
    if exc.status_code == 404 and not request.url.path.startswith("/api/"):
        index_file = os.path.join(FRONTEND_DIST, "index.html")
        if os.path.exists(index_file):
            return FileResponse(index_file)
    return JSONResponse(status_code=exc.status_code, content={"detail": exc.detail})

